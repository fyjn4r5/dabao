#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# 1.py - 批量生成部门信件（终极稳定版 + AAA 替换为隐藏 UUID + 附带内容功能）
import os
import sys
import json
import tkinter as tk
from tkinter import messagebox, filedialog, colorchooser, PhotoImage, ttk
import datetime
import threading
import subprocess
import re
import uuid
import zipfile
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tkinter.font as tkfont
import shutil
import time
# ==================== 基础配置 ====================
BASE_DIR = os.path.dirname(os.path.abspath(sys.executable)) if getattr(sys, "frozen", False) else os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(BASE_DIR, "departments_config.json")
DEFAULT_CONTENT = "请在这里输入要写的信件统一内容。\n1、默认导出是宋体、小四。\n2、支持字体颜色修改。\n3、支持添加字体背景颜色。\n4、支持记忆使用的颜色（适合单色的颜色）。\n5、支持字体改为黑体。\n6、支持恢复宋体。\n7、支持粘贴、框选、复制。\n8、支持按住 Ctrl 配合鼠标滚轮滚动即可调整字体大小【仅支持阅读，不支持导出】\n9、支持包含有 AAA 讲自动替换为唯一 UUID【适合配合“跟进助手”这款软件使用】\n10、支持把写好的内容导出与导入【极少使用】"
DEFAULT_SIGNATURE = "自己的代码"
DEFAULT_HINT = ""
KEYWORD = "AAA"
KEYWORD_LOWER = KEYWORD.lower()
UUID_REGEX = re.compile(r'[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}')

class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip = None
        widget.bind("<Enter>", self.enter)
        widget.bind("<Leave>", self.leave)
    def enter(self, event=None):
        if self.tip or not self.text:
            return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5
        self.tip = tk.Toplevel(self.widget)
        self.tip.wm_overrideredirect(True)
        self.tip.wm_geometry(f"+{x}+{y}")
        lbl = tk.Label(self.tip, text=self.text, bg="#FFFFDD", fg="#333",
                       relief="solid", bd=1, font=("微软雅黑", 9), wraplength=300)
        lbl.pack()
    def leave(self, event=None):
        if self.tip:
            self.tip.destroy()
            self.tip = None

MAX_GROUP_WIDTH = 20
# ==================== 自动换行容器 ====================
class FlowFrame(tk.Frame):
    """子控件超出宽度自动折行"""
    def __init__(self, parent, **kwargs):
        super().__init__(parent, **kwargs)
        self._reracking = False
        self.bind("<Configure>", self._on_configure)
    def _on_configure(self, event):
        self.after_idle(self._rerack)
    def _rerack(self):
        if self._reracking:
            return
        self._reracking = True
        try:
            width = self.winfo_width()
            children = self.winfo_children()
            if not children or width <= 1:
                return
            for w in children:
                w.grid_forget()
            row = col = 0
            row_width = 0
            for w in children:
                w_width = w.winfo_reqwidth() + 20
                if row_width + w_width > width:
                    row += 1
                    col = 0
                    row_width = 0
                w.grid(row=row, column=col, sticky='w', padx=10, pady=1)
                col += 1
                row_width += w_width
        finally:
            self._reracking = False
# ==================== 行号画布 ====================
class LineNumbers(tk.Canvas):
    def __init__(self, parent, **kwargs):
        super().__init__(parent, **kwargs)
        self.text_widget = None
    def attach(self, text_widget):
        self.text_widget = text_widget
        self.text_widget.bind('<Configure>', self.redraw)
        self.text_widget.bind('<KeyRelease>', self.redraw)
        self.text_widget.bind('<MouseWheel>', self.redraw)
        self.text_widget.bind('<Button-4>', self.redraw)
        self.text_widget.bind('<Button-5>', self.redraw)
        self.text_widget.bind('<FocusIn>', self.redraw)
        self.text_widget.bind('<FocusOut>', self.redraw)
        self.redraw()
    def redraw(self, *args):
        if self.text_widget is None:
            return
        self.delete("all")
        i = self.text_widget.index("@0,0")
        while True:
            dline = self.text_widget.dlineinfo(i)
            if dline is None:
                break
            y = dline[1]
            linenum = str(i).split('.')[0]
            font_obj = tkfont.Font(font=self.text_widget.cget('font'))
            font_size = font_obj.actual()['size']
            self.create_text(6, y, anchor="nw", text=linenum, fill="#666666",
                             font=("Consolas", min(12, font_size - 2)))
            i = self.text_widget.index(f"{i}+1line")
# ==================== 隐藏 UUID 辅助函数 ====================
def add_vanish_to_run(run):
    """为 run 添加 <w:vanish/>，实现“显示/隐藏编辑标记”时可见"""
    rPr = run._r.get_or_add_rPr()
    vanish = OxmlElement('w:vanish')
    rPr.append(vanish)
# ==================== 外部 Word 替换 AAA 为隐藏 UUID ====================

# ==================== 主类 ====================
class LetterGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("批量生成信件_260324")
        self.root.geometry("950x650")
        self.center_root_window()
        self.icon_path = None
        try:
            if getattr(sys, 'frozen', False):
                base_path = sys._MEIPASS
            else:
                base_path = BASE_DIR
            for p in [os.path.join(base_path, "1.ico"), os.path.join(base_path, "icons", "1.ico")]:
                if os.path.exists(p):
                    self.icon_path = p
                    self.root.iconbitmap(self.icon_path)
                    break
        except Exception as e:
            print("主窗口图标设置失败:", e)
        self.group_vars = {}
        self.member_vars = {}
        self.last_fg_color = None
        self.last_bg_color = None
        self.text_font_size = 15
        self.text_font = tkfont.Font(family="宋体", size=self.text_font_size)
        self.default_save_dir = BASE_DIR
        self.custom_save_path = None
        self.attach_content_paths = []  # 附带内容路径列表
        self.generate_txt_var = tk.BooleanVar(value=False)
        self.generate_docx_var = tk.BooleanVar(value=True)
        # ==================== 设置区 ====================
        self.settings_frame = tk.Frame(root, relief="groove", bd=1)
        self.settings_frame.pack(fill=tk.X, padx=10, pady=(10,5))
        self.settings_header = tk.Frame(self.settings_frame)
        self.settings_header.pack(fill=tk.X)
        self.settings_label = tk.Label(self.settings_header, text="设置", font=("微软雅黑", 11, "bold"), cursor="hand2")
        self.settings_label.pack(side=tk.LEFT, padx=8, pady=4)
        self.toggle_icon = tk.Label(self.settings_header, text="▲", width=3, cursor="hand2", font=("Segoe UI Symbol", 10))
        self.toggle_icon.pack(side=tk.LEFT)
        self.settings_content = tk.Frame(self.settings_frame)
        self.settings_header.bind("<Button-1>", lambda e: self.toggle_settings())
        self.settings_label.bind("<Button-1>", lambda e: self.toggle_settings())
        self.toggle_icon.bind("<Button-1>", lambda e: self.toggle_settings())
        row1 = tk.Frame(self.settings_content)
        row1.pack(fill=tk.X, pady=3)
        group_btn_frame = tk.Frame(row1)
        group_btn_frame.pack(side=tk.LEFT)
        tk.Button(group_btn_frame, text="添加分组", command=self.add_group_dialog).pack(side=tk.LEFT, padx=2)
        tk.Button(group_btn_frame, text="编辑分组", command=self.open_edit_group_window).pack(side=tk.LEFT, padx=2)
        tk.Button(group_btn_frame, text="删除分组", command=self.open_delete_group_window).pack(side=tk.LEFT, padx=2)
        tk.Label(row1, text="  ").pack(side=tk.LEFT)
        tk.Label(row1, text="落款签名").pack(side=tk.LEFT)
        self.signature_entry = tk.Entry(row1, width=15)
        self.signature_entry.pack(side=tk.LEFT, padx=5)
        self.signature_entry.insert(0, DEFAULT_SIGNATURE)
        tk.Checkbutton(
            row1,
            text="生成txt",
            variable=self.generate_txt_var,
            font=("微软雅黑", 9)
        ).pack(side=tk.LEFT, padx=5)
        tk.Checkbutton(
            row1,
            text="生成docx",
            variable=self.generate_docx_var,
            font=("微软雅黑", 9)
        ).pack(side=tk.LEFT, padx=5)
        # 第二行：保存路径
        row2 = tk.Frame(self.settings_content)
        row2.pack(fill=tk.X, pady=3)
        tk.Label(row2, text="保存路径").pack(side=tk.LEFT, padx=8)
        self.path_var = tk.StringVar()
        self.path_entry = tk.Entry(row2, textvariable=self.path_var, width=50)
        self.path_entry.pack(side=tk.LEFT, padx=2)
        tk.Button(row2, text="浏览", command=self.choose_custom_path).pack(side=tk.LEFT, padx=2)
        tk.Button(row2, text="清除", command=self.clear_custom_path).pack(side=tk.LEFT, padx=2)
        tk.Label(row2, text="  ").pack(side=tk.LEFT, expand=True)
        self.settings_expanded = False
        # ==================== 主操作区 ====================
        top_main_frame = tk.Frame(root)
        top_main_frame.pack(fill=tk.X, padx=10, pady=5)
        row1 = tk.Frame(top_main_frame)
        row1.pack(fill=tk.X, pady=2)
        hint_frame = tk.Frame(row1)
        hint_frame.pack(side=tk.LEFT)
        tk.Label(hint_frame, text="文件名备注").pack(side=tk.LEFT)
        self.hint_entry = tk.Entry(hint_frame, width=20)
        self.hint_entry.pack(side=tk.LEFT, padx=5)
        export_frame = tk.Frame(row1)
        export_frame.pack(side=tk.LEFT, padx=20)
        tk.Button(export_frame, text="导出为Word", command=self.export_content_to_word, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Button(export_frame, text="从Word导入", command=self.import_content_from_word, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Label(row1, text="  ").pack(side=tk.LEFT, expand=True)
        self.folder_per_member_var = tk.BooleanVar(value=False)
        tk.Checkbutton(
            row1,
            text="每人独立文件夹",
            variable=self.folder_per_member_var,
            font=("微软雅黑", 9)
        ).pack(side=tk.LEFT, padx=5)
        self.attach_content_var = tk.BooleanVar(value=False)
        attach_frame = tk.Frame(row1)
        attach_frame.pack(side=tk.LEFT, padx=5)
        tk.Checkbutton(
            attach_frame,
            text="拷贝文件到附件",
            variable=self.attach_content_var,
            font=("微软雅黑", 9)
        ).pack(side=tk.LEFT)
        self.attach_count_label = tk.Label(attach_frame, text="(0)", fg="#888", font=("微软雅黑", 8))
        self.attach_count_label.pack(side=tk.LEFT, padx=2)
        tk.Button(attach_frame, text="管理附件", command=self.manage_attachments, font=("微软雅黑", 9)).pack(side=tk.LEFT, padx=2)
        tk.Button(
            row1,
            text="生成信件",
            command=self.start_generation,
            font=("微软雅黑", 10, "bold"),
            bg="#0078D7",
            fg="white",
            relief="flat",
            padx=14,
            pady=2
        ).pack(side=tk.LEFT, padx=(15, 0))
        # ==================== 分组区（带滚动条） ====================
        self.group_container = tk.Frame(root)
        self.group_container.pack(fill=tk.X, padx=10)
        self.group_canvas = tk.Canvas(self.group_container, borderwidth=0, highlightthickness=0, height=180)
        self.group_scrollbar = tk.Scrollbar(self.group_container, orient="vertical", command=self.group_canvas.yview)
        self.group_frame = tk.Frame(self.group_canvas)
        for c in range(5):
            self.group_frame.columnconfigure(c, weight=1)
        self.group_frame.bind("<Configure>", lambda e: self.group_canvas.configure(scrollregion=self.group_canvas.bbox("all")))
        self._group_window_id = self.group_canvas.create_window((0, 0), window=self.group_frame, anchor="nw")
        def _resize_group(event):
            self.group_canvas.itemconfig(self._group_window_id, width=event.width)
        self.group_canvas.bind("<Configure>", _resize_group)
        self.group_canvas.configure(yscrollcommand=self.group_scrollbar.set)
        self.group_canvas.pack(side="left", fill="both", expand=True)
        self.group_scrollbar.pack(side="right", fill="y")
        def _gmw(event):
            self.group_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        self.group_canvas.bind("<MouseWheel>", _gmw)
        self.group_frame.bind("<MouseWheel>", _gmw)
        self.group_canvas.bind("<Button-4>", lambda e: self.group_canvas.yview_scroll(-1, "units"))
        self.group_frame.bind("<Button-4>", lambda e: self.group_canvas.yview_scroll(-1, "units"))
        self.group_canvas.bind("<Button-5>", lambda e: self.group_canvas.yview_scroll(1, "units"))
        self.group_frame.bind("<Button-5>", lambda e: self.group_canvas.yview_scroll(1, "units"))
        # ==================== 信件内容区 ====================
        self.content_header_frame = tk.Frame(root, relief="groove", bd=1)
        self.content_header_frame.pack(fill=tk.X, padx=10, pady=(10, 0))
        self.content_title_frame = tk.Frame(self.content_header_frame)
        self.content_title_frame.pack(expand=True)
        self.content_label = tk.Label(
            self.content_title_frame,
            text="信件内容",
            font=("微软雅黑", 10, "bold"),
            cursor="hand2"
        )
        self.content_label.pack(side=tk.LEFT)
        self.content_toggle_icon = tk.Label(
            self.content_title_frame,
            text="▼",
            width=3,
            cursor="hand2",
            font=("Segoe UI Symbol", 10)
        )
        self.content_toggle_icon.pack(side=tk.LEFT, padx=(8, 0))
        self.content_label.bind("<Button-1>", lambda e: self.toggle_content_section())
        self.content_toggle_icon.bind("<Button-1>", lambda e: self.toggle_content_section())
        self.content_container = tk.Frame(root)
        self.content_container.pack(fill=tk.BOTH, padx=10, pady=(0,10), expand=True)
        self.line_numbers = LineNumbers(self.content_container, width=30, bg="#FFEAEA", highlightthickness=0)
        self.line_numbers.pack(side=tk.LEFT, fill=tk.Y)
        v_scroll = tk.Scrollbar(self.content_container, orient=tk.VERTICAL)
        v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.content_text = tk.Text(
            self.content_container,
            height=25,
            font=self.text_font,
            wrap='word',
            yscrollcommand=v_scroll.set,
            undo=True,
            padx=8,
            spacing1=2,
            spacing3=2
        )
        self.content_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scroll.config(command=self.content_text.yview)
        self.line_numbers.attach(self.content_text)
        self.content_text.insert(tk.END, DEFAULT_CONTENT)
        self.content_text.bind("<FocusIn>", self.clear_default_content)
        self.content_text.bind("<Control-MouseWheel>", self.on_ctrl_mousewheel)
        self.content_text.bind("<Control-Button-4>", self.on_ctrl_mousewheel)
        self.content_text.bind("<Control-Button-5>", self.on_ctrl_mousewheel)
        self.text_menu = tk.Menu(self.content_text, tearoff=0)
        self.text_menu.add_command(label="设置字体颜色", command=self.set_selected_fg_color)
        self.text_menu.add_command(label="设置背景颜色", command=self.set_selected_bg_color)
        self.text_menu.add_separator()
        self.text_menu.add_command(label="使用上一次颜色", command=self.use_last_color)
        self.text_menu.add_separator()
        self.text_menu.add_command(label="设为黑体", command=lambda: self.set_font_family("黑体"))
        self.text_menu.add_command(label="恢复宋体", command=lambda: self.set_font_family("宋体"))
        self.content_text.bind("<Button-3>", self.show_text_menu)
        # ==================== 加载配置 ====================
        self.groups = self.load_config()
        self.update_group_display()
        self.update_path_display()
        self.content_expanded = True

    # ==================== 折叠控制 ====================
    def toggle_settings(self):
        if self.settings_expanded:
            self.settings_content.pack_forget()
            self.toggle_icon.config(text="▲")
            self.settings_expanded = False
        else:
            self.settings_content.pack(fill=tk.X, padx=8, pady=(0,8))
            self.toggle_icon.config(text="▼")
            self.settings_expanded = True
    def toggle_content_section(self):
        if self.content_expanded:
            self.content_container.pack_forget()
            self.content_toggle_icon.config(text="▲")
            self.content_expanded = False
        else:
            self.content_container.pack(fill=tk.BOTH, padx=10, pady=(0,10), expand=True)
            self.content_toggle_icon.config(text="▼")
            self.content_expanded = True
# ==================== 配置读写 ====================
    def load_config(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.groups = data.get("groups", {})
                    signature = data.get("signature", DEFAULT_SIGNATURE)
#                    hint = data.get("title_hint", DEFAULT_HINT)
                    self.default_save_dir = data.get("default_save_dir", BASE_DIR)
                    self.custom_save_path = data.get("custom_save_path", None)
                    self.signature_entry.delete(0, tk.END)
                    self.signature_entry.insert(0, signature)
                    return self.groups
            except Exception as e:
                print("读取配置失败:", e)
        return {}
    def save_config(self):
        data = {
            "groups": self.groups,
            "signature": self.signature_entry.get().strip(),
            "default_save_dir": self.default_save_dir,
            "custom_save_path": self.custom_save_path,
        }
        try:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print("保存配置失败：", e)

    # ==================== 路径操作 ====================
    def choose_custom_path(self):
        path = filedialog.askdirectory(title="选择默认保存目录", initialdir=self.default_save_dir)
        if path:
            self.custom_save_path = path
            self.path_var.set(path)
            self.save_config()
    def clear_custom_path(self):
        self.custom_save_path = None
        self.path_var.set("")
        self.save_config()
    def update_path_display(self):
        if self.custom_save_path and os.path.isdir(self.custom_save_path):
            self.path_var.set(self.custom_save_path)
        else:
            self.path_var.set(self.default_save_dir)

    # ==================== 文本操作 ====================
    def clear_default_content(self, event=None):
        if self.content_text.get("1.0", "end-1c").strip() == DEFAULT_CONTENT.strip():
            self.content_text.delete("1.0", tk.END)
            self.content_text.unbind("<FocusIn>")
    def on_ctrl_mousewheel(self, event):
        try:
            if hasattr(event, 'delta'):
                delta = event.delta
                if delta > 0:
                    self.text_font_size += 1
                else:
                    self.text_font_size = max(6, self.text_font_size - 1)
            else:
                if event.num == 4:
                    self.text_font_size += 1
                elif event.num == 5:
                    self.text_font_size = max(6, self.text_font_size - 1)
            self.text_font.configure(size=self.text_font_size)
            self.line_numbers.redraw()
            return "break"
        except Exception as e:
            print("调整字体大小出错:", e)
    def show_text_menu(self, event):
        try:
            self.content_text.index("sel.first")
            self.text_menu.post(event.x_root, event.y_root)
        except tk.TclError:
            pass
    def set_selected_fg_color(self):
        try:
            start = self.content_text.index("sel.first")
            end = self.content_text.index("sel.last")
        except tk.TclError:
            return
        color = colorchooser.askcolor(title="选择字体颜色")[1]
        if color:
            self.last_fg_color = color
            tag = f"fg_{color}"
            self.content_text.tag_add(tag, start, end)
            self.content_text.tag_config(tag, foreground=color)
    def set_selected_bg_color(self):
        try:
            start = self.content_text.index("sel.first")
            end = self.content_text.index("sel.last")
        except tk.TclError:
            return
        color = colorchooser.askcolor(title="选择背景颜色")[1]
        if color:
            self.last_bg_color = color
            tag = f"bg_{color}"
            self.content_text.tag_add(tag, start, end)
            self.content_text.tag_config(tag, background=color)
    def use_last_color(self):
        try:
            start = self.content_text.index("sel.first")
            end = self.content_text.index("sel.last")
        except tk.TclError:
            return
        if self.last_fg_color:
            tag = f"fg_{self.last_fg_color}"
            self.content_text.tag_add(tag, start, end)
            self.content_text.tag_config(tag, foreground=self.last_fg_color)
        if self.last_bg_color:
            tag = f"bg_{self.last_bg_color}"
            self.content_text.tag_add(tag, start, end)
            self.content_text.tag_config(tag, background=self.last_bg_color)
    def set_font_family(self, family):
        try:
            start = self.content_text.index("sel.first")
            end = self.content_text.index("sel.last")
        except tk.TclError:
            return
        tag = f"font_{family}"
        self.content_text.tag_add(tag, start, end)
        self.content_text.tag_config(tag, font=(family, self.text_font_size))

    # ==================== 生成信件（关键修改） ====================
# ==================== 生成信件（关键修改） ====================

    def start_generation(self):
        selected = [(g, m) for (g, m), var in self.member_vars.items() if var.get()]
        if not selected:
            messagebox.showerror("错误", "请勾选部门/成员")
            return

        # === 生成格式（从设置读取） ===
        generate_txt = self.generate_txt_var.get()
        generate_docx = self.generate_docx_var.get()
        if not generate_txt and not generate_docx:
            messagebox.showerror("错误", "请至少勾选一种输出格式（生成txt/生成docx）")
            return

        # === 附带内容（已通过“管理附件”配置） ===

        # 选择保存目录（自动创建）
        if self.custom_save_path:
            os.makedirs(self.custom_save_path, exist_ok=True)
            folder = self.custom_save_path
        else:
            folder = filedialog.askdirectory(title="选择保存目录", initialdir=self.default_save_dir)
            if not folder:
                return
            self.default_save_dir = folder

        # 创建进度窗口
        progress_win = self.create_toplevel("正在生成信件...", 400, 120)
        progress_win.attributes("-topmost", True)
        lbl = tk.Label(progress_win, text=f"0 / {len(selected)}")
        lbl.pack(pady=8)
        pb = ttk.Progressbar(progress_win, length=350, mode='determinate', maximum=len(selected))
        pb.pack(pady=5)

        cancel_flag = {"cancel": False}

        def on_cancel():
            if messagebox.askyesno("取消", "确定取消生成？", parent=progress_win):
                cancel_flag["cancel"] = True

        tk.Button(progress_win, text="取消", command=on_cancel).pack(pady=3)

        threading.Thread(
            target=self.generate_letters,
            args=(selected, folder, pb, lbl, progress_win, cancel_flag, generate_txt, generate_docx),
            daemon=True
        ).start()


    def generate_letters(self, selected, folder, progressbar, label, progress_win, cancel_flag, generate_txt, generate_docx):
        signature = self.signature_entry.get().strip()
        title_hint = self.hint_entry.get().strip()
        total = len(selected)

        for idx_item, (group, member) in enumerate(selected, start=1):
            if cancel_flag.get("cancel"):
                break

            try:
                replacement_uuid_for_doc = str(uuid.uuid4())
                doc = Document()

                # ====================== 正文内容 ======================
                try:
                    content_len = int(self.content_text.index('end-1c').split('.')[0])
                except:
                    content_len = int(self.content_text.index('end').split('.')[0])

                for i in range(1, content_len + 1):
                    line_text = self.content_text.get(f"{i}.0", f"{i}.end")
                    if not line_text.strip():
                        doc.add_paragraph()
                        continue

                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.first_line_indent = Cm(0.74)

                    parts = line_text.split("AAA")
                    for part_idx, part in enumerate(parts):
                        if part_idx > 0:
                            uuid_run = p.add_run(replacement_uuid_for_doc)
                            uuid_run.font.size = Pt(12)
                            uuid_run.font.name = '宋体'
                            rPr = uuid_run._r.get_or_add_rPr()
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:eastAsia'), '宋体')
                            rPr.append(rFonts)
                            add_vanish_to_run(uuid_run)

                        if part:
                            line_offset = sum(len(parts[k]) for k in range(part_idx)) + part_idx * 3
                            pos = 0
                            while pos < len(part):
                                col = line_offset + pos
                                fmt = self._get_char_format(i, col)
                                end = pos + 1
                                while end < len(part) and self._get_char_format(i, line_offset + end) == fmt:
                                    end += 1
                                run = p.add_run(part[pos:end])
                                font_name, fg_hex, bg_hex = fmt
                                self.apply_font_to_run(run, font_name, 12)
                                if fg_hex:
                                    try:
                                        r, g, b = self.hex_to_rgb(fg_hex)
                                        run.font.color.rgb = RGBColor(r, g, b)
                                    except:
                                        pass
                                if bg_hex:
                                    try:
                                        self.set_run_background(run, bg_hex)
                                    except:
                                        pass
                                pos = end

                # ====================== sectPr ======================
                body = doc.element.body
                for child in list(body):
                    if child.tag == qn('w:sectPr'):
                        body.remove(child)
                body.append(OxmlElement('w:sectPr'))

                # ====================== 开头两行（强制宋体） ======================
                p_dept = doc.add_paragraph()
                p_dept.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p_dept.paragraph_format.first_line_indent = Cm(0)
                p_dept.paragraph_format.space_after = Pt(0)
                run_dept = p_dept.add_run(f"{member}：")
                self._apply_songti_font(run_dept)

                p_hello = doc.add_paragraph()
                p_hello.paragraph_format.first_line_indent = Cm(0.74)
                p_hello.paragraph_format.space_after = Pt(0)
                run_hello = p_hello.add_run("你们好！")
                self._apply_songti_font(run_hello)

                body.insert(0, p_dept._p)
                body.insert(1, p_hello._p)

                # ====================== 落款 ======================
                p_sig = doc.add_paragraph()
                p_sig.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                run_sig = p_sig.add_run(signature)
                self._apply_songti_font(run_sig)
                p_sig.paragraph_format.space_after = Pt(0)

                p_date = doc.add_paragraph()
                p_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                run_date = p_date.add_run(datetime.datetime.today().strftime("%Y年%m月%d日"))
                self._apply_songti_font(run_date)

                sectPr_list = body.findall(qn('w:sectPr'))
                if sectPr_list:
                    insert_pos = list(body).index(sectPr_list[0])
                    body.insert(insert_pos, p_date._p)
                    body.insert(insert_pos, p_sig._p)
                else:
                    body.append(p_sig._p)
                    body.append(p_date._p)

                # ====================== 保存 ======================
                date_str = datetime.datetime.now().strftime("%m%d")
                hint_part = f"【{title_hint}】" if title_hint else ""
                safe_member = "".join(c for c in member if c not in r'\/:*?"<>|')
                base_filename = f"{signature}-{safe_member}_{date_str}{hint_part}"

                target_dir = os.path.join(folder, base_filename) if self.folder_per_member_var.get() else folder
                os.makedirs(target_dir, exist_ok=True)

                # ====================== 生成 txt ======================
                if generate_txt:
                    txt_path = os.path.join(target_dir, base_filename + ".txt")
                    with open(txt_path, "w", encoding="utf-8") as txt_file:
                        txt_file.write(f"{member}：\n")
                        txt_file.write("        你们好！\n")
                        txt_file.write(self.content_text.get("1.0", "end-1c"))
                        txt_file.write(f"\n\n{signature}\n{datetime.datetime.today().strftime('%Y年%m月%d日')}\n")

                # ====================== 生成 docx ======================
                if generate_docx:
                    doc.save(os.path.join(target_dir, base_filename + ".docx"))

                # ====================== 处理附件 ======================
                if self.attach_content_var.get():
                    today = datetime.datetime.now().strftime("%Y%m%d")
                    attachment_dir = os.path.join(target_dir, f"附件_{today}")
                    os.makedirs(attachment_dir, exist_ok=True)

                    for ap in self.attach_content_paths:
                        if os.path.exists(ap):
                            try:
                                if os.path.isdir(ap):
                                    folder_name = os.path.basename(ap)
                                    dest_path = os.path.join(attachment_dir, folder_name)
                                    if os.path.exists(dest_path):
                                        shutil.rmtree(dest_path)
                                    shutil.copytree(ap, dest_path)
                                else:
                                    file_name = os.path.basename(ap)
                                    dest_path = os.path.join(attachment_dir, file_name)
                                    shutil.copy2(ap, dest_path)
                            except Exception as copy_e:
                                print(f"附带内容复制失败 {group}/{member}: {copy_e}")

            except Exception as e:
                print(f"生成失败 {group}/{member}: {e}")
                import traceback
                traceback.print_exc()

            # 更新进度
            self.root.after(0, lambda v=idx_item: (
                progressbar.config(value=v),
                label.config(text=f"{v}/{total}")
            ))

        # 完成
        self.root.after(0, progress_win.destroy)
        try:
            if sys.platform.startswith('win'):
                os.startfile(folder)
        except:
            pass

        self.attach_content_paths.clear()
        self.save_config()


    # 小工具方法（建议加到类里，避免重复代码）
    def _apply_songti_font(self, run):
        """统一设置宋体字体"""
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.bold = False
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), '宋体')
        rFonts.set(qn('w:hAnsi'), '宋体')
        rPr.append(rFonts)

    def _get_char_format(self, line_num, col):
        char_index = f"{line_num}.{col}"
        tags = self.content_text.tag_names(char_index)
        font_name = "宋体"
        for tag in tags:
            if tag.startswith("font_"):
                font_name = tag.split("_", 1)[1]
                break
        fg = ""
        bg = ""
        for tag in tags:
            try:
                f = self.content_text.tag_cget(tag, "foreground")
                if f:
                    fg = f
            except:
                pass
            try:
                b = self.content_text.tag_cget(tag, "background")
                if b:
                    bg = b
            except:
                pass
        return (font_name, fg, bg)

    # ==================== 生成信件（关键修改） ====================

    @staticmethod
    def hex_to_rgb(hex_color):
        hex_color = hex_color.lstrip("#")
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    @staticmethod
    def set_run_background(run, hex_color):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color.lstrip('#').upper())
        run._r.get_or_add_rPr().append(shd)
    def apply_font_to_run(self, run, name, size_pt):
        try:
            run.font.name = name
            run.font.size = Pt(size_pt)
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), name)
            rPr.append(rFonts)
        except:
            try:
                run.font.size = Pt(size_pt)
            except: pass

    # ==================== 其余函数保持不变（省略） ====================
    # （update_group_display, toggle_group_members, show_members_popup, add_group_dialog, ...）
    # （为了篇幅，此处省略，但必须保留原代码）


    # ==================== 分组管理等其余函数保持不变 ====================
    def update_group_display(self):
        saved_states = {g: self.group_vars[g].get() for g in self.group_vars if g in self.groups}
        for widget in self.group_frame.winfo_children():
            widget.destroy()
        self.group_vars = {}
        self.group_frame.update_idletasks()
        cw = self.group_frame.winfo_width()
        if cw < 50:
            cw = 800
        col_w = (cw - 40) // 5
        font = tkfont.Font(font=("微软雅黑", 9))
        chk_w = 22
        for i, group in enumerate(self.groups):
            old = saved_states.get(group, False)
            var_group = tk.BooleanVar(value=old)
            tw = font.measure(group)
            if tw > col_w - chk_w:
                d = group
                while font.measure(d + "…") > col_w - chk_w and len(d) > 1:
                    d = d[:-1]
                display = d + "…"
            else:
                display = group
            chk = tk.Checkbutton(self.group_frame, text=display, variable=var_group)
            chk.grid(row=i//5, column=i%5, sticky="w", padx=5, pady=1)
            chk.bind("<Button-1>", lambda e, g=group: self.toggle_group_members(g))
            chk.bind("<Button-3>", lambda e, g=group, w=chk: self.show_members_popup(g, w))
            if display != group:
                ToolTip(chk, group)
            self.group_vars[group] = var_group
    def toggle_group_members(self, group):
        now = time.time()
        if (hasattr(self, '_last_click') and now - self._last_click['time'] < 0.4
                and self._last_click['group'] == group):
            self.edit_group_members(group)
            self._last_click = None
            return "break"
        self._last_click = {'time': now, 'group': group}
        var_group = self.group_vars[group]
        new_state = not var_group.get()
        var_group.set(new_state)
        for member in self.groups[group]:
            self.member_vars[(group, member)] = self.member_vars.get((group, member), tk.BooleanVar())
            self.member_vars[(group, member)].set(new_state)
        return "break"
    def edit_group_members(self, group_name):
        members = self.groups.get(group_name, [])
        initial = "\n".join(members)
        dlg = self.create_toplevel(f"编辑 {group_name}", 400, 350)
        label_text = f"编辑 {group_name} 的成员："
        tk.Label(dlg, text=label_text, wraplength=380, anchor="w", justify="left").pack(anchor="w", padx=10, pady=(10,0))
        tk.Label(dlg, text="每行一个成员：", anchor="w", fg="#666").pack(anchor="w", padx=10)
        text_frame = tk.Frame(dlg)
        text_frame.pack(fill=tk.BOTH, padx=10, pady=(5,10), expand=True)
        t = tk.Text(text_frame, height=12, width=45)
        t.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scroll = tk.Scrollbar(text_frame, orient=tk.VERTICAL, command=t.yview)
        v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        t.config(yscrollcommand=v_scroll.set)
        t.insert("1.0", initial)
        ok_res = {"ok": False}
        def ok():
            raw = t.get("1.0", "end-1c").strip().split("\n")
            members_list = [m.strip() for m in raw if m.strip()]
            members_list = list(dict.fromkeys(members_list))
            self.groups[group_name] = members_list
            ok_res["ok"] = True
            dlg.destroy()
        def cancel():
            dlg.destroy()
        btnf = tk.Frame(dlg)
        btnf.pack(pady=5)
        tk.Button(btnf, text="确定", command=ok).pack(side=tk.LEFT, padx=6)
        tk.Button(btnf, text="取消", command=cancel).pack(side=tk.LEFT, padx=6)
        dlg.wait_window()
        if ok_res["ok"]:
            self.save_config()
            self.update_group_display()
    def show_members_popup(self, group, widget):
        members = self.groups.get(group, [])
        x = widget.winfo_rootx()
        y = widget.winfo_rooty() + widget.winfo_height()
        screen_height = self.root.winfo_screenheight()
        max_height = int(screen_height * 0.7)
        popup_width = 300
        popup = tk.Toplevel(self.root)
        popup.title(f"{group} 成员选择")
        popup.transient(self.root)
        popup.attributes("-topmost", True)
        container = tk.Frame(popup)
        container.pack(fill=tk.BOTH, expand=True)
        canvas = tk.Canvas(container, borderwidth=0, highlightthickness=0)
        scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill=tk.Y)
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind("<MouseWheel>", _on_mousewheel)
        scrollable_frame.bind("<MouseWheel>", _on_mousewheel)
        canvas.bind("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        scrollable_frame.bind("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        canvas.bind("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))
        scrollable_frame.bind("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))
        for member in members:
            var = self.member_vars.get((group, member))
            if not var:
                var = tk.BooleanVar()
                self.member_vars[(group, member)] = var
            chk = tk.Checkbutton(scrollable_frame, text=member, variable=var, anchor="w")
            chk.pack(fill='x', padx=5, pady=2)
        tk.Button(popup, text="关闭", command=popup.destroy).pack(pady=5)
        # 根据实际内容高度调整窗口高度
        popup.update_idletasks()
        content_height = scrollable_frame.winfo_reqheight() + 50
        popup_height = min(content_height, max_height)
        popup.geometry(f"{popup_width}x{popup_height}+{x}+{y}")
    def manage_attachments(self):
        win = self.create_toplevel("管理附件", 520, 350)
        tk.Label(win, text="已选择的附件的文件和文件夹：").pack(anchor="w", padx=10, pady=(8,0))
        frame = tk.Frame(win)
        frame.pack(fill=tk.BOTH, padx=10, pady=5, expand=True)
        listbox = tk.Listbox(frame)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar = tk.Scrollbar(frame, orient="vertical", command=listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        listbox.configure(yscrollcommand=scrollbar.set)
        def refresh():
            listbox.delete(0, tk.END)
            for p in self.attach_content_paths:
                listbox.insert(tk.END, os.path.basename(p) if os.path.isfile(p) else f"[文件夹] {os.path.basename(p)}")
            self.attach_count_label.config(text=f"({len(self.attach_content_paths)})")
        refresh()
        btn_frame = tk.Frame(win)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        def add_files():
            win.attributes("-topmost", False)
            paths = filedialog.askopenfilenames(title="选择要附带的文件", filetypes=[("所有文件", "*.*")], parent=win)
            win.attributes("-topmost", True)
            for p in paths:
                if p not in self.attach_content_paths:
                    self.attach_content_paths.append(p)
            refresh()
        def add_folder():
            win.attributes("-topmost", False)
            path = filedialog.askdirectory(title="选择要附带的文件夹", parent=win)
            win.attributes("-topmost", True)
            if path and path not in self.attach_content_paths:
                self.attach_content_paths.append(path)
                refresh()
        def remove_selected():
            sel = listbox.curselection()
            for i in reversed(sel):
                self.attach_content_paths.pop(i)
            refresh()
        def clear_all():
            self.attach_content_paths.clear()
            refresh()
        tk.Button(btn_frame, text="添加文件", command=add_files).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="添加文件夹", command=add_folder).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="移除选中", command=remove_selected).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="清空", command=clear_all).pack(side=tk.LEFT, padx=3)
        tk.Button(btn_frame, text="关闭", command=win.destroy).pack(side=tk.RIGHT, padx=3)
    def add_group_dialog(self):
        win = self.create_toplevel("添加分组", 400, 320)
        tk.Label(win, text="分组名称：").pack(anchor="w", padx=10, pady=(10,0))
        entry_group = tk.Entry(win)
        entry_group.pack(fill=tk.X, padx=10, pady=(0,5))
        tk.Label(win, text="成员（每行一个）：").pack(anchor="w", padx=10)
        t = tk.Text(win, height=8, width=40)
        t.pack(fill=tk.BOTH, padx=10, pady=(0,10), expand=True)
        result = {"ok": False, "group": None, "members": None}
        def on_ok():
            group_name = entry_group.get().strip()
            raw = t.get("1.0", "end-1c").strip().split("\n")
            members = [m.strip() for m in raw if m.strip()]
            members = list(dict.fromkeys(members))
            if not group_name:
                messagebox.showerror("错误", "分组名称不能为空", parent=win)
                return
            result["ok"] = True
            result["group"] = group_name
            result["members"] = members
            win.destroy()
        def on_cancel():
            win.destroy()
        btn_frame = tk.Frame(win)
        btn_frame.pack(pady=5)
        tk.Button(btn_frame, text="确定", command=on_ok).pack(side=tk.LEFT, padx=10)
        tk.Button(btn_frame, text="取消", command=on_cancel).pack(side=tk.LEFT, padx=10)
        win.wait_window()
        if result["ok"]:
            self.groups[result["group"]] = result["members"]
            self.save_config()
            self.update_group_display()
    def open_edit_group_window(self):
        win = self.create_toplevel("编辑分组", 400, 350)
        info_label = tk.Label(win, text="双击分组名或选中后点击编辑来修改成员")
        info_label.pack(pady=5)
        listbox_frame = tk.Frame(win)
        listbox_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        listbox = tk.Listbox(listbox_frame)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        lb_scroll = tk.Scrollbar(listbox_frame, orient="vertical", command=listbox.yview)
        lb_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        listbox.configure(yscrollcommand=lb_scroll.set)
        def _listbox_refresh():
            listbox.delete(0, tk.END)
            name_map.clear()
            for g in self.groups:
                display = g if len(g) <= MAX_GROUP_WIDTH else g[:MAX_GROUP_WIDTH] + "…"
                name_map[display] = g
                listbox.insert(tk.END, display)
        name_map = {}
        _listbox_refresh()
        def edit_selected(event=None):
            sel = listbox.curselection()
            if not sel:
                messagebox.showinfo("提示", "请选择一个分组", parent=win)
                return
            display = listbox.get(sel[0])
            group_name = name_map.get(display, display)
            initial = "\n".join(self.groups.get(group_name, []))
            dlg = self.create_toplevel(f"编辑 {group_name}", 450, 450)
            tk.Label(dlg, text=f"编辑 {group_name} 的成员：", wraplength=420, anchor="w", justify="left").pack(anchor="w", padx=10, pady=(10,0))
            tk.Label(dlg, text="每行一个成员：", anchor="w", fg="#666").pack(anchor="w", padx=10)
            text_frame = tk.Frame(dlg)
            text_frame.pack(fill=tk.BOTH, padx=10, pady=(5,10), expand=True)
            t = tk.Text(text_frame, height=15, width=50)
            t.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            v_scroll = tk.Scrollbar(text_frame, orient=tk.VERTICAL, command=t.yview)
            v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
            t.config(yscrollcommand=v_scroll.set)
            t.insert("1.0", initial)
            ok_res = {"ok": False}
            def ok():
                raw = t.get("1.0", "end-1c").strip().split("\n")
                members = [m.strip() for m in raw if m.strip()]
                members = list(dict.fromkeys(members))
                self.groups[group_name] = members
                ok_res["ok"] = True
                dlg.destroy()
            def cancel():
                dlg.destroy()
            btnf = tk.Frame(dlg)
            btnf.pack(pady=5)
            tk.Button(btnf, text="确定", command=ok).pack(side=tk.LEFT, padx=6)
            tk.Button(btnf, text="取消", command=cancel).pack(side=tk.LEFT, padx=6)
            dlg.wait_window()
            if ok_res["ok"]:
                self.save_config()
                self.update_group_display()
                _listbox_refresh()
        listbox.bind("<Double-Button-1>", edit_selected)
        btn_frame = tk.Frame(win)
        btn_frame.pack(pady=5)
        tk.Button(btn_frame, text="编辑", command=edit_selected).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="关闭", command=win.destroy).pack(side=tk.LEFT, padx=5)
    def open_delete_group_window(self):
        win = self.create_toplevel("删除分组", 400, 300)
        tk.Label(win, text="勾选要删除的分组，然后点击确定").pack(pady=5)
        container = tk.Frame(win)
        container.pack(fill=tk.BOTH, expand=True, padx=10)
        canvas = tk.Canvas(container, borderwidth=0, highlightthickness=0)
        scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill=tk.Y)
        def _on_mw(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind("<MouseWheel>", _on_mw)
        scrollable_frame.bind("<MouseWheel>", _on_mw)
        canvas.bind("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        scrollable_frame.bind("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        canvas.bind("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))
        scrollable_frame.bind("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))
        check_vars = {}
        for g in list(self.groups.keys()):
            v = tk.BooleanVar(value=False)
            display = g if len(g) <= MAX_GROUP_WIDTH else g[:MAX_GROUP_WIDTH] + "…"
            chk = tk.Checkbutton(scrollable_frame, text=display, variable=v)
            chk.pack(anchor="w")
            check_vars[g] = v
            if len(g) > MAX_GROUP_WIDTH:
                ToolTip(chk, g)
        def confirm_delete():
            to_delete = [g for g, v in check_vars.items() if v.get()]
            if not to_delete:
                messagebox.showinfo("提示", "没有勾选任何分组", parent=win)
                return
            confirm_win = self.create_toplevel("确认删除", 320, 140)
            tk.Label(confirm_win, text=f"确定删除选中的 {len(to_delete)} 个分组？").pack(pady=10, padx=10)
            ok_flag = {"ok": False}
            def do_ok():
                ok_flag["ok"] = True
                confirm_win.destroy()
            def do_cancel():
                confirm_win.destroy()
            btnf = tk.Frame(confirm_win)
            btnf.pack(pady=8)
            tk.Button(btnf, text="确定", command=do_ok).pack(side=tk.LEFT, padx=8)
            tk.Button(btnf, text="取消", command=do_cancel).pack(side=tk.LEFT, padx=8)
            confirm_win.wait_window()
            if not ok_flag["ok"]:
                return
            for g in to_delete:
                if g in self.groups:
                    self.groups.pop(g)
                for member in list(self.member_vars.keys()):
                    if member[0] == g:
                        self.member_vars.pop(member, None)
            self.save_config()
            self.update_group_display()
            win.destroy()
        btn_frame = tk.Frame(win)
        btn_frame.pack(pady=5)
        tk.Button(btn_frame, text="确定删除", command=confirm_delete).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="取消", command=win.destroy).pack(side=tk.LEFT, padx=5)
    def center_root_window(self):
        self.root.update_idletasks()
        sw = self.root.winfo_screenwidth()
        sh = self.root.winfo_screenheight()
        w, h = 950, 650
        self.root.geometry(f"{w}x{h}+{(sw-w)//2}+{(sh-h)//2}")
    def center_window(self, win, w, h):
        sw = win.winfo_screenwidth()
        sh = win.winfo_screenheight()
        x = (sw - w) // 2
        y = (sh - h) // 2
        win.geometry(f"{w}x{h}+{max(0, x)}+{max(0, y)}")
    def create_toplevel(self, title, width=None, height=None, resizable=(True, True)):
        win = tk.Toplevel(self.root)
        win.title(title)
        win.transient(self.root)
        # 先设置位置再显示，避免闪烁
        if width and height:
            self.center_window(win, width, height)
        else:
            win.geometry(f"+{win.winfo_screenwidth()//4}+{win.winfo_screenheight()//4}")
        win.wait_visibility()
        win.grab_set()
        win.attributes("-topmost", True)
        if self.icon_path and os.path.exists(self.icon_path):
            try:
                win.iconbitmap(self.icon_path)
            except:
                pass
        win.resizable(*resizable)
        return win
    def export_content_to_word(self):
        file_path = filedialog.asksaveasfilename(
            title="导出统一信为 Word",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")],
            initialdir=self.default_save_dir
        )
        if not file_path: return
        doc = Document()
        content = self.content_text.get("1.0", tk.END)
        lines = content.splitlines()
        for i, line in enumerate(lines, 1):
            if not line.strip():
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(6)
            pos = 0
            while pos < len(line):
                fmt = self._get_char_format(i, pos)
                end = pos + 1
                while end < len(line) and self._get_char_format(i, end) == fmt:
                    end += 1
                run = p.add_run(line[pos:end])
                font_name, fg_hex, bg_hex = fmt
                self.apply_font_to_run(run, font_name, 12)
                if fg_hex:
                    try:
                        r, g, b = self.hex_to_rgb(fg_hex)
                        run.font.color.rgb = RGBColor(r, g, b)
                    except: pass
                if bg_hex:
                    try:
                        self.set_run_background(run, bg_hex)
                    except: pass
                pos = end
        try:
            doc.save(file_path)
            messagebox.showinfo("成功", f"已导出为 Word：\n{file_path}\n（AAA 未替换，格式完整保留）")
            self.default_save_dir = os.path.dirname(file_path)
            self.save_config()
        except Exception as e:
            messagebox.showerror("错误", f"导出失败：{e}")
    def import_content_from_word(self):
        file_path = filedialog.askopenfilename(
            title="从 Word 导入统一信",
            filetypes=[("Word 文档", "*.docx")],
            initialdir=self.default_save_dir
        )
        if not file_path: return
        try:
            doc = Document(file_path)
            self.content_text.delete("1.0", tk.END)
            def ensure_tag_config(tag, **options):
                if tag not in self.content_text.tag_names():
                    self.content_text.tag_config(tag, **options)
            for para_idx, para in enumerate(doc.paragraphs):
                if not para.text.strip() and not para.runs:
                    self.content_text.insert(tk.END, "\n")
                    continue
                for run_idx, run in enumerate(para.runs):
                    text = run.text
                    is_hidden = False
                    try:
                        if run.font.color and run.font.color.rgb == RGBColor(255, 255, 255):
                            is_hidden = True
                    except: pass
                    if is_hidden and UUID_REGEX.fullmatch(text.strip()):
                        text = "AAA"
                    font_name = "宋体"
                    is_bold = False
                    try:
                        if getattr(run, 'bold', False) or getattr(run.font, 'bold', False):
                            is_bold = True
                        ea = (getattr(run.font, 'eastAsia', '') or '').lower()
                        fn = (getattr(run.font, 'name', '') or '').lower()
                        if any(x in ea for x in ['黑体', 'hei', 'simhei']):
                            font_name = "黑体"
                        elif any(x in ea for x in ['宋体', 'song', 'simsun']):
                            font_name = "宋体"
                        elif any(x in fn for x in ['黑体', 'hei', 'simhei']):
                            font_name = "黑体"
                        elif any(x in fn for x in ['宋体', 'song', 'simsun']):
                            font_name = "宋体"
                    except: pass
                    if is_bold:
                        font_name = "黑体"
                    fg_color = None
                    try:
                        if run.font.color and run.font.color.rgb:
                            r, g, b = run.font.color.rgb
                            if (r, g, b) != (255, 255, 255):
                                fg_color = f"#{r:02x}{g:02x}{b:02x}"
                    except: pass
                    bg_color = None
                    try:
                        rPr = run._element.get_or_add_rPr()
                        shd = rPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill.strip() and fill.lower() not in ['auto', 'none', '']:
                                fill = fill.strip().upper()
                                if len(fill) == 3:
                                    fill = ''.join(c*2 for c in fill)
                                if len(fill) == 6:
                                    bg_color = f"#{fill}"
                    except: pass
                    if not bg_color:
                        try:
                            if run.font.highlight_color:
                                hc = run.font.highlight_color
                                map_dict = {
                                    'YELLOW': '#FFFF00', 'BRIGHT_GREEN': '#00FF00', 'TURQUOISE': '#00FFFF',
                                    'PINK': '#FFC0CB', 'BLUE': '#0000FF', 'RED': '#FF0000', 'DARK_BLUE': '#00008B',
                                    'DARK_RED': '#8B0000', 'TEAL': '#008080', 'GREEN': '#008000', 'VIOLET': '#EE82EE',
                                    'DARK_YELLOW': '#808000', 'GRAY_50': '#808080', 'GRAY_25': '#C0C0C0', 'BLACK': '#000000',
                                }
                                key = hc.name.upper().replace(' ', '_')
                                bg_color = map_dict.get(key)
                        except: pass
                    start_idx = self.content_text.index(tk.INSERT)
                    self.content_text.insert(tk.INSERT, text)
                    end_idx = self.content_text.index(tk.INSERT)
                    if font_name == "黑体":
                        tag = "font_黑体"
                        ensure_tag_config(tag, font=("黑体", self.text_font_size))
                        self.content_text.tag_add(tag, start_idx, end_idx)
                    if fg_color:
                        tag = f"fg_{fg_color}"
                        ensure_tag_config(tag, foreground=fg_color)
                        self.content_text.tag_add(tag, start_idx, end_idx)
                    if bg_color:
                        tag = f"bg_{bg_color}"
                        ensure_tag_config(tag, background=bg_color)
                        self.content_text.tag_add(tag, start_idx, end_idx)
                self.content_text.insert(tk.INSERT, "\n")
            messagebox.showinfo("成功", "已从 Word 导入\n（软件内“设置背景颜色”100% 还原）")
            self.default_save_dir = os.path.dirname(file_path)
            self.line_numbers.redraw()
        except Exception as e:
            messagebox.showerror("错误", f"导入失败：{e}")
            import traceback
            print(traceback.format_exc())
if __name__ == "__main__":
    root = tk.Tk()
    app = LetterGenerator(root)
    root.mainloop()
    
##添加生成txt格式的信件##
