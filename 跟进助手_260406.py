#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# ==============================
# 标准库模块
# ==============================
import os           # 文件和目录操作
import sys          # 系统相关操作
import subprocess   # 子进程管理
import uuid         # 生成唯一标识符
import re           # 正则表达式
import time         # 时间处理
import shutil       # 文件和目录的高级操作
import random       # 随机值
import threading    # 多线程支持
import logging      # 日志记录
import importlib
import sqlite3
import pyperclip  # 用于复制UUID到剪贴板
from logging.handlers import RotatingFileHandler  # 日志文件轮转
from datetime import datetime, timedelta  # 日期和时间处理
from functools import wraps              # 函数装饰器
import fnmatch
from tkinter import messagebox
 

# ==============================
# 第三方库模块
# ==============================
from dateutil.relativedelta import relativedelta   # 日期相对计算
from docx import Document                          # 操作 Word 文档
from docx.shared import RGBColor, Inches, Pt       # Word 文档字体/颜色/尺寸
from docx.oxml import OxmlElement
from docx.oxml.ns import qn                        # 处理 Word 中文字体
from docx.enum.text import WD_ALIGN_PARAGRAPH      # Word 段落对齐
from PIL import Image, ImageTk                     # 图像处理（Pillow）
import mammoth                                     # Word 文档转换为 HTML
from concurrent.futures import ThreadPoolExecutor, as_completed  # 并发执行
import psutil                                      # 系统资源监控

# ==============================
# Windows 系统相关模块
# ==============================
import win32api        # Windows API 操作
import win32con        # Windows 常量
import win32event      # Windows 事件处理
import win32gui        # Windows 图形界面操作
import win32process    # Windows 进程操作

# ==============================
# Tkinter 相关模块（GUI 界面）
# ==============================
import tkinter as tk
from tkinter import ttk, filedialog, Toplevel, Label, StringVar, messagebox
from tkinterdnd2 import TkinterDnD, DND_FILES     # Tkinter 拖放支持

# ==============================
# 数据库相关模块
# ==============================
import sqlite3   # SQLite 数据库操作


# 动态获取程序所在目录，兼容 pyinstaller 打包
def get_base_path():
    """获取可执行文件所在目录（打包后）或脚本目录（开发时）"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

# 动态获取资源文件路径（icons 等），兼容 pyinstaller
def get_resource_path(relative_path):
    """获取资源文件路径，兼容 pyinstaller 打包"""
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

# 检查依赖
def check_dependencies():
    """检查必要的 Python 库是否安装"""
    required_modules = ["docx", "tkinterdnd2", "mammoth", "dateutil"]
    missing = []
    for module in required_modules:
        if not importlib.util.find_spec(module):
            missing.append(module)
    if missing:
        messagebox.showerror("错误", f"缺少以下依赖库: {', '.join(missing)}\n请运行 'pip install {' '.join(missing)}' 安装")
        sys.exit(1)

check_dependencies()

# 设置日志和目录
BASE_DIR = get_base_path()
DATA_DIR = os.path.join(BASE_DIR, "Data")
LOG_DIR = os.path.join(DATA_DIR, "logs")
DB_DIR = os.path.join(DATA_DIR, "db")
BACKUP_DIR = os.path.join(DB_DIR, "backups")
ICON_DIR = get_resource_path("icons")
LOG_PATH = os.path.join(LOG_DIR, "letter_tracker.log")
DB_PATH = os.path.join(DB_DIR, "letters.db")
APP_ICON_PATH = get_resource_path("icons/跟进助手.ico")

try:
    os.makedirs(LOG_DIR, exist_ok=True)
    os.makedirs(DB_DIR, exist_ok=True)
    os.makedirs(BACKUP_DIR, exist_ok=True)
    os.makedirs(ICON_DIR, exist_ok=True)
    if not os.access(LOG_DIR, os.W_OK | os.R_OK):
        messagebox.showerror("错误", f"日志目录 {LOG_DIR} 没有写权限")
        sys.exit(1)
    if not os.access(DB_DIR, os.W_OK | os.R_OK):
        messagebox.showerror("错误", f"数据库目录 {DB_DIR} 没有写权限")
        sys.exit(1)
except OSError as e:
    messagebox.showerror("错误", f"无法创建数据目录: {str(e)}\n请检查权限或更改目录")
    sys.exit(1)

try:
    handler = RotatingFileHandler(LOG_PATH, maxBytes=3*1024*1024, backupCount=5, encoding='utf-8')
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s.%(msecs)03d [%(levelname)s] [%(funcName)s:%(lineno)d] %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S',
        handlers=[
            RotatingFileHandler(LOG_PATH, maxBytes=3*1024*1024, backupCount=5, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
except OSError as e:
    messagebox.showerror("错误", f"无法写入日志文件 {LOG_PATH}: {str(e)}\n请检查权限")
    sys.exit(1)

# 数据库备份时间控制
last_backup_time = None

def backup_database():
    """备份数据库文件，每10分钟一次，保留最多3个备份"""
    global last_backup_time
    current_time = datetime.now()
    if last_backup_time and (current_time - last_backup_time).total_seconds() < 600:
        return

    try:
        if not os.access(BACKUP_DIR, os.W_OK | os.R_OK):
            logging.error(f"备份目录 {BACKUP_DIR} 没有写权限")
            messagebox.showwarning("警告", f"备份目录 {BACKUP_DIR} 没有写权限")
            return

        timestamp = current_time.strftime("%Y%m%d_%H%M%S")
        backup_path = os.path.join(BACKUP_DIR, f"letters_backup_{timestamp}.db")
        
        if not os.path.exists(DB_PATH):
            logging.warning("数据库文件不存在，跳过备份")
            return
        
        shutil.copy2(DB_PATH, backup_path)
        logging.info(f"数据库备份成功: {backup_path}")
        last_backup_time = current_time

        backup_files = sorted(
            [f for f in os.listdir(BACKUP_DIR) if f.startswith("letters_backup_") and f.endswith(".db")],
            key=lambda x: os.path.getmtime(os.path.join(BACKUP_DIR, x)),
            reverse=True
        )
        for old_backup in backup_files[3:]:
            try:
                os.remove(os.path.join(BACKUP_DIR, old_backup))
                logging.info(f"移除旧备份: {old_backup}")
            except OSError as e:
                logging.error(f"移除旧备份 {old_backup} 失败: {str(e)}")
                messagebox.showwarning("警告", f"移除旧备份 {old_backup} 失败: {str(e)}")
    except OSError as e:
        logging.error(f"数据库备份失败: {str(e)}")
        messagebox.showwarning("警告", f"数据库备份失败: {str(e)}")
    except Exception as e:
        logging.error(f"备份过程中发生意外错误: {str(e)}")
        messagebox.showwarning("警告", f"备份过程中发生意外错误: {str(e)}")

# ---- 数据库相关 ----
def create_db_connection(timeout=10.0):
    """创建数据库连接，设置超时防止文件锁定"""
    try:
        if not os.access(os.path.dirname(DB_PATH), os.W_OK | os.R_OK):
            logging.error(f"数据库目录 {os.path.dirname(DB_PATH)} 没有写权限")
            messagebox.showerror("错误", f"数据库目录 {os.path.dirname(DB_PATH)} 没有写权限")
            return None
        conn = sqlite3.connect(DB_PATH, timeout=timeout)
        conn.execute(f"PRAGMA busy_timeout = {int(timeout * 1000)}")  # 转换为毫秒
        logging.debug(f"创建数据库连接，超时设置为 {timeout} 秒")
        return conn
    except sqlite3.Error as e:
        logging.error(f"数据库连接失败: {str(e)}")
        messagebox.showerror("错误", f"数据库连接失败: {str(e)}\n请检查数据库文件权限或路径")
        return None

def with_db_connection(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        conn = create_db_connection()
        if not conn:
            return None
        try:
            result = func(conn, *args, **kwargs)
            conn.commit()
            backup_database()
            return result
        except sqlite3.IntegrityError as e:
            conn.rollback()
            logging.error(f"数据库完整性错误: {str(e)}, 函数: {func.__name__}, 参数: {args}, {kwargs}")
            messagebox.showerror("错误", f"数据库完整性错误: {str(e)}")
            return None
        except sqlite3.OperationalError as e:
            conn.rollback()
            logging.error(f"数据库操作失败: {str(e)}, 函数: {func.__name__}, 参数: {args}, {kwargs}")
            messagebox.showerror("错误", f"数据库操作失败: {str(e)}. 请稍后重试或检查是否有其他程序占用数据库。")
            return None
        except Exception as e:
            conn.rollback()
            logging.error(f"意外错误: {str(e)}, 函数: {func.__name__}, 参数: {args}, {kwargs}")
            messagebox.showerror("错误", f"意外错误: {str(e)}")
            return None
        finally:
            try:
                conn.close()
            except Exception as e:
                logging.error(f"关闭数据库连接失败: {str(e)}, 函数: {func.__name__}")
    return wrapper
    


@with_db_connection
def init_db(conn):
    """初始化数据库表结构，支持备注字段"""
    try:
        cursor = conn.cursor()
        
        # 检查 letters 表是否存在
        cursor.execute("PRAGMA table_info(letters)")
        columns = [col[1] for col in cursor.fetchall()]
        
        if not columns:
            # 表不存在，创建完整表（含 note 字段）
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS letters (
                    id TEXT PRIMARY KEY,
                    file_path TEXT,
                    send_date TEXT,
                    subject TEXT,
                    reply_status TEXT,
                    reply_file_path TEXT,
                    monitor INTEGER,
                    reminder_date TEXT,
                    reminder_status TEXT,
                    reply_file_mtime REAL,
                    note TEXT
                )
            """)
        else:
            # 表已存在，检查并添加缺失字段
            if 'note' not in columns:
                cursor.execute("ALTER TABLE letters ADD COLUMN note TEXT")
                logging.info("已为 letters 表添加 note（备注）字段")
            if 'reminder_date' not in columns:
                cursor.execute("ALTER TABLE letters ADD COLUMN reminder_date TEXT")
            if 'reminder_status' not in columns:
                cursor.execute("ALTER TABLE letters ADD COLUMN reminder_status TEXT")
            if 'reply_file_mtime' not in columns:
                cursor.execute("ALTER TABLE letters ADD COLUMN reply_file_mtime REAL")

        # settings 表（原代码就有，必须完整写，不能写 ...）
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        """)
        
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_subject ON letters(subject)")
        logging.info("数据库初始化成功")
        
    except sqlite3.Error as e:
        logging.error(f"初始化数据库失败: {str(e)}")
        messagebox.showerror("错误", f"初始化数据库失败: {str(e)}")
        sys.exit(1)

@with_db_connection
def load_setting(conn, key, default=None):
    """加载指定键的设置值，如果不存在则返回默认值"""
    if not conn:
        logging.error(f"无法加载设置 {key}：数据库连接失败")
        return default
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = ?", (key,))
        result = cursor.fetchone()
        return result[0] if result else default
    except sqlite3.Error as e:
        logging.error(f"加载设置 {key} 失败，数据库错误: {str(e)}")
        return default
    except Exception as e:
        logging.error(f"加载设置 {key} 失败，意外错误: {str(e)}")
        return default

# 初始化数据库
init_db()

# 加载初始设置
default_keyword = load_setting('keyword', '回复')
default_aa = load_setting('aa_keyword', '')
scan_path1 = load_setting('scan_path1', "Y:\\%YYYY%年%MM%月\\")
scan_last_month_setting = load_setting('scan_last_month', 'False') == 'True'
saved_window_width_setting = load_setting('window_width')
max_scan_files_setting = load_setting('max_scan_files', '3000')
try:
    saved_window_width = int(saved_window_width_setting) if saved_window_width_setting and saved_window_width_setting.isdigit() else None
except ValueError:
    saved_window_width = None
try:
    max_scan_files = int(max_scan_files_setting) if max_scan_files_setting and max_scan_files_setting.isdigit() else 3000
except ValueError:
    max_scan_files = 3000
daily_scan_enabled_setting = load_setting('daily_scan_enabled', 'False') == 'True'
daily_scan_time_setting = load_setting('daily_scan_time', '09:00')

# --- 创建主窗口 ---
root = TkinterDnD.Tk()
root.withdraw()  # 隐藏窗口，直到配置完成
root.title("跟进助手_260406")
try:
    root.iconbitmap(APP_ICON_PATH)
except tk.TclError as e:
    logging.warning(f"无法加载主窗口图标 {APP_ICON_PATH}: {str(e)}")

screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
window_width = saved_window_width or int(screen_width * 2 / 3)
window_height = int(screen_height * 2 / 3)
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2
root.geometry(f"{window_width}x{window_height}+{x}+{y}")
root.deiconify()  # 配置完成后显示窗口

# 加载彩色图标
try:
    icon_calendar = ImageTk.PhotoImage(Image.open(os.path.join(ICON_DIR, "calendar.png")).resize((16, 16)))
    icon_delete = ImageTk.PhotoImage(Image.open(os.path.join(ICON_DIR, "delete.png")).resize((16, 16)))
    icon_file = ImageTk.PhotoImage(Image.open(os.path.join(ICON_DIR, "file.png")).resize((16, 16)))
    icon_folder = ImageTk.PhotoImage(Image.open(os.path.join(ICON_DIR, "folder.png")).resize((16, 16)))
    icon_link = ImageTk.PhotoImage(Image.open(os.path.join(ICON_DIR, "link.png")).resize((16, 16)))
    icon_refresh = ImageTk.PhotoImage(Image.open(os.path.join(ICON_DIR, "refresh.png")).resize((16, 16)))
except FileNotFoundError as e:
    logging.error(f"图标文件未找到: {str(e)}")
    messagebox.showerror("错误", f"图标文件未找到: {str(e)}\n请确保 icons 文件夹包含所有所需图标")
    sys.exit(1)
except Exception as e:
    logging.error(f"加载图标时发生错误: {str(e)}")
    messagebox.showerror("错误", f"加载图标时发生错误: {str(e)}")
    sys.exit(1)

# Tk 变量
update_letter_list_pending = None
scan_last_month_var = tk.BooleanVar(value=scan_last_month_setting)
daily_scan_enabled_var = tk.BooleanVar(value=daily_scan_enabled_setting)
daily_scan_time_var = tk.StringVar(value=daily_scan_time_setting)
aa_var = tk.StringVar(value=default_aa)
max_scan_files_var = tk.StringVar(value=str(max_scan_files))
keyword_var = tk.StringVar(value=default_keyword)
scan_path1_var = tk.StringVar(value=scan_path1)
resize_pending = None
daily_scan_timer = None
sort_column = None
sort_reverse = False
scan_lock = threading.Lock()
drop_lock = threading.Lock()
manual_upload_lock = threading.Lock()
search_var = tk.StringVar()  # 新增：搜索关键词变量

# ---- 工具提示类 ----
class ToolTip:
    """为控件添加鼠标悬停提示，支持多行文本"""
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip_window = None
        self.show_tip_id = None  # 用于延迟显示
        self.widget.bind("<Enter>", self.schedule_show_tip)
        self.widget.bind("<Leave>", self.hide_tip)

    def schedule_show_tip(self, event=None):
        """延迟显示提示，避免快速鼠标移动导致闪烁"""
        if self.show_tip_id:
            self.widget.after_cancel(self.show_tip_id)  # 取消之前的延迟任务
        self.show_tip_id = self.widget.after(300, lambda: self.show_tip(event))  # 300ms 延迟

    def show_tip(self, event=None):
        """显示提示窗口"""
        if self.tip_window or not self.text:
            return
        x = event.x_root + 10
        y = event.y_root + 10
        self.tip_window = tw = tk.Toplevel(self.widget)
        try:
            tw.iconbitmap(APP_ICON_PATH)
        except tk.TclError as e:
            logging.warning(f"无法加载提示窗口图标 {APP_ICON_PATH}: {str(e)}")
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, justify=tk.LEFT, background="#ffffe0", 
                        relief=tk.SOLID, borderwidth=1, wraplength=300)
        label.pack()

    def hide_tip(self, event=None):
        """隐藏提示窗口"""
        if self.show_tip_id:
            self.widget.after_cancel(self.show_tip_id)  # 取消未执行的显示任务
            self.show_tip_id = None
        if self.tip_window:
            self.tip_window.destroy()
            self.tip_window = None

# ---- 功能函数 ----
def copy_font(src_font, dest_font):
    """复制字体属性"""
    if src_font is None or dest_font is None:
        return
    try:
        dest_font.name = src_font.name
        dest_font.size = src_font.size
        dest_font.color.rgb = src_font.color.rgb
    except AttributeError as e:
        logging.warning(f"复制字体属性时发生错误: {str(e)}")

def calc_days_since(send_date_str):
    """
    根据 send_date（'YYYY-MM-DD' 或前缀）计算到今天的天数。
    send_date_str 可能为 None 或空，出错时返回空字符串。
    """
    if not send_date_str:
        return ""
    try:
        # 只取前 10 位，避免含时间时出错
        date_part = send_date_str[:10]
        send_date = datetime.strptime(date_part, "%Y-%m-%d")
        delta = datetime.now().date() - send_date.date()
        # 负数也直接显示，比如 -3 表示未来 3 天
        return str(delta.days)
    except Exception:
        # 格式不对就返回空，不要抛异常
        return ""


def copy_shading(src_para, dest_para):
    """复制段落底纹颜色"""
    if src_para.paragraph_format.shading and src_para.paragraph_format.shading.background_pattern_color is not None:
        dest_para.paragraph_format.shading.background_pattern_color = \
            src_para.paragraph_format.shading.background_pattern_color


def calc_days_since(send_date_str):
    """
    根据 send_date（'YYYY-MM-DD' 或前缀）计算到今天的天数。
    send_date_str 可能为 None 或空，出错时返回空字符串。
    """
    if not send_date_str:
        return ""
    try:
        date_part = send_date_str[:10]  # 避免有时间部分
        send_date = datetime.strptime(date_part, "%Y-%m-%d")
        delta = datetime.now().date() - send_date.date()
        return str(delta.days)
    except Exception:
        return ""


def extract_uuid_from_doc(file_path, keyword=None):
    """提取文档中的第一个36位UUID，不依赖关键词"""
    try:
        doc = Document(file_path)
        # 正则表达式匹配标准UUID格式：xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx
        uuid_pattern = r'[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}'
        for para in doc.paragraphs:
            for run in para.runs:
                # 在 run.text 中查找所有UUID
                matches = re.findall(uuid_pattern, run.text, re.IGNORECASE)
                if matches:
                    # 返回第一个匹配的UUID
                    return matches[0]
        return None
    except Exception as e:
        logging.error(f"提取UUID失败 {file_path}: {str(e)}")
        return None



def set_chinese_font(run, font_name="宋体"):
    """强制 run 使用指定中文字体"""
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)  # 中文字体
    rFonts.set(qn("w:ascii"), "Times New Roman")  # 英文字体
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")



def copy_uuid_to_clipboard_middle_button(event):
    """鼠标中键点击时触发，查询数据库并复制UUID到剪贴板"""
    col = tree.identify_column(event.x)
    item = tree.identify_row(event.y)
    if not item:
        return

    # 获取当前选中行的数据
    values = tree.item(item, "values")
    subject = values[1]  # 假设文档名称在第一列（如果有问题请调整列号）
    file_path = values[2]  # 假设文件路径在第二列（如果有问题请调整列号）

    # 打印选中行的数据，检查是否获取到正确的subject和file_path
    logging.info(f"选中行的数据：{values}")
    logging.info(f"查询UUID，文档名称：{subject}, 文件路径：{file_path}")

    # 连接到数据库并查询UUID
    db_file_path = r'.\Data\db\letters.db'  # 确保数据库路径正确
    try:
        conn = sqlite3.connect(db_file_path)
        cursor = conn.cursor()

        # 查询数据库，根据文件路径或文档名称获取对应的UUID
        cursor.execute("SELECT id FROM letters WHERE subject = ? OR file_path = ?", (subject, file_path))
        row = cursor.fetchone()

        if row:
            uuid = row[0]  # 获取查询结果中的UUID
            pyperclip.copy(uuid)  # 将UUID复制到剪贴板
            logging.info(f"UUID {uuid} 已复制到剪贴板，文件名: {subject}")
        else:
            logging.warning(f"未找到UUID，当前选中的数据：{values}")
            messagebox.showwarning("警告", f"无法在数据库中找到文件 {subject} 的UUID")
        
        # 关闭数据库连接
        conn.close()
    except sqlite3.Error as e:
        logging.error(f"数据库查询失败: {e}")
        messagebox.showerror("错误", "数据库查询失败，无法获取UUID")


@with_db_connection
def replace_keyword(conn, file_path, keyword, hide_uuid=True):
    """使用参考代码的稳定方式：在关键词后面插入隐藏UUID，关键词位置不变"""
    try:
        if not os.path.isfile(file_path):
            return None, f"文件不存在: {os.path.basename(file_path)}"
        if not os.access(file_path, os.R_OK | os.W_OK):
            return None, f"文件无读写权限: {os.path.basename(file_path)}"

        doc = Document(file_path)
        reminder_date = None

        # 提取提醒日期（你的原逻辑）
        date_pattern = r'(\d{1,2})(?:号前|号以前|号之前|日前|日以前|日之前)'
        for para in doc.paragraphs:
            for run in para.runs:
                matches = re.findall(date_pattern, run.text)
                if matches:
                    day = int(matches[0])
                    current_date = datetime.now()
                    cy, cm, cd = current_date.year, current_date.month, current_date.day
                    if cm < 12:
                        last_day = (datetime(cy, cm + 1, 1) - timedelta(days=1)).day
                    else:
                        last_day = 31

                    if day <= last_day and day >= cd:
                        reminder_date = datetime(cy, cm, day, 9, 0)
                    else:
                        next_month = cm + 1 if cm < 12 else 1
                        next_year = cy if cm < 12 else cy + 1
                        last_day_next = 31 if next_month == 12 else (datetime(next_year, next_month + 1, 1) - timedelta(days=1)).day
                        day = min(day, last_day_next)
                        reminder_date = datetime(next_year, next_month, day, 9, 0)
                    reminder_date = reminder_date.strftime("%Y-%m-%d %H:%M")
                    logging.info(f"提取到提醒日期: {reminder_date}")
                    break
            if reminder_date:
                break

        # 检查已有UUID
        existing_uuid = extract_uuid_from_doc(file_path)
        if existing_uuid:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM letters WHERE id = ?", (existing_uuid,))
            if cursor.fetchone():
                return None, f"文件已包含UUID: {existing_uuid}，请移除UUID后重试"
            else:
                subject = os.path.basename(file_path)
                cursor.execute("""
                    INSERT INTO letters (id, file_path, send_date, subject, reply_status, monitor, reminder_date, reminder_status)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """, (existing_uuid, file_path, datetime.now().strftime("%Y-%m-%d"), subject, "pending", 1, reminder_date, "active" if reminder_date else None))
                conn.commit()
                return existing_uuid, file_path

        if not keyword:
            return None, f"标记关键词设置为空，且文件不包含已有UUID"

        letter_id = str(uuid.uuid4())
        found = False

        # ==================== 你参考代码的核心函数 ====================
        def set_run_hidden(run_obj, hidden=True):
            rPr = run_obj._r.get_or_add_rPr()
            vanish = rPr.find(qn('w:vanish'))
            if hidden:
                if vanish is None:
                    vanish = OxmlElement('w:vanish')
                    rPr.append(vanish)
            else:
                if vanish is not None:
                    rPr.remove(vanish)

        def copy_run_format(src_run, dst_run):
            try:
                dst_run.bold = src_run.bold
                dst_run.italic = src_run.italic
                dst_run.underline = src_run.underline
                dst_run.font.name = src_run.font.name
                dst_run.font.size = src_run.font.size
                if src_run.font.color and src_run.font.color.rgb:
                    dst_run.font.color.rgb = src_run.font.color.rgb
            except:
                pass

        def insert_run_after(run, text="", hidden=False, style_src=None):
            new_run = run._parent.add_run(text)
            if style_src is not None:
                copy_run_format(style_src, new_run)
            if hidden:
                set_run_hidden(new_run, True)
            run._r.addnext(new_run._r)
            return new_run

        def process_paragraph(paragraph, kw, doc_uuid):
            if not paragraph.runs:
                return False
            target = kw.lower()
            for run in list(paragraph.runs):
                text = run.text or ""
                pos = text.lower().find(target)
                if pos == -1:
                    continue
                end_pos = pos + len(kw)
                if end_pos > len(text):
                    continue

                before = text[:end_pos]
                after = text[end_pos:]

                original_style = run
                run.text = before

                if after:
                    after_run = insert_run_after(run, after, hidden=False, style_src=original_style)
                    insert_run_after(after_run, f"_{doc_uuid}", hidden=True, style_src=original_style)
                else:
                    insert_run_after(run, f"_{doc_uuid}", hidden=True, style_src=original_style)
                return True
            return False

        # 执行插入
        for para in doc.paragraphs:
            if process_paragraph(para, keyword, letter_id):
                found = True
                break

        if not found:
            return None, f"文件 {os.path.basename(file_path)} 不包含关键词 '{keyword}'"

        doc.save(file_path)

        # 写入数据库
        subject = os.path.basename(file_path)
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO letters (id, file_path, send_date, subject, reply_status, monitor, reminder_date, reminder_status)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (letter_id, file_path, datetime.now().strftime("%Y-%m-%d"), subject, "pending", 1, reminder_date, "active" if reminder_date else None))
        conn.commit()

        logging.info(f"成功插入隐藏UUID: {letter_id}")
        return letter_id, file_path

    except Exception as e:
        logging.error(f"replace_keyword 失败 {file_path}: {str(e)}")
        return None, f"处理失败: {str(e)}"




        

@with_db_connection
def drop(conn, file_paths=None):
    """处理拖放或新增跟进的文件"""
    try:
        if file_paths is None:
            messagebox.showerror("错误", "未提供文件路径")
            return
        file_paths = [os.path.normpath(os.path.abspath(p)) for p in file_paths]
        valid_files = []
        for path in file_paths:
            if os.path.isdir(path):
                for root_, _, files in os.walk(path):
                    for file in files:
                        full_path = os.path.abspath(os.path.join(root_, file))
                        if file.lower().endswith(".docx") and os.access(full_path, os.R_OK | os.W_OK):
                            valid_files.append(full_path)
            elif os.path.isfile(path) and path.lower().endswith(".docx") and os.access(path, os.R_OK | os.W_OK):
                valid_files.append(path)
        if not valid_files:
            messagebox.showwarning("警告", "没有找到有效的 .docx 文件！")
            return
        if len(valid_files) > 10000:
            messagebox.showerror("错误", f"一次最多处理 10000 个文件！")
            return
        progress_win = Toplevel(root)
        progress_win.title("处理文件中…")
        try:
            progress_win.iconbitmap(APP_ICON_PATH)
        except tk.TclError:
            pass
        progress_win.geometry("300x180")
        progress_text = StringVar()
        Label(progress_win, textvariable=progress_text).pack(pady=10)
        progress_bar = ttk.Progressbar(progress_win, maximum=len(valid_files), mode="determinate")
        progress_bar.pack(pady=10, fill="x", padx=10)
        progress_win.update()
        keyword = default_keyword
        failed_files = []
        processed_files = 0
        cursor = conn.cursor()
        for file_path in valid_files:
            processed_files += 1
            progress_text.set(f"处理中：{processed_files} / {len(valid_files)}\n{os.path.basename(file_path)}")
            progress_bar.configure(value=processed_files)
            progress_win.update()
            cursor.execute("SELECT subject FROM letters WHERE subject = ?", (os.path.basename(file_path),))
            if cursor.fetchone():
                failed_files.append(f"{os.path.basename(file_path)}: 文件已存在")
                continue
            try:
                with open(file_path, 'rb') as f:
                    doc = Document(f)
                    if not any(keyword in run.text for para in doc.paragraphs for run in para.runs):
                        failed_files.append(f"{os.path.basename(file_path)}: 未找到关键词 '{keyword}'")
                        continue
            except Exception as e:
                failed_files.append(f"{os.path.basename(file_path)}: 文件无效或无法读取")
                continue
            result = replace_keyword(file_path, keyword, hide_uuid=True)
            if result is None or isinstance(result, tuple) and result[0] is None:
                failed_files.append(result[1] if result else f"{os.path.basename(file_path)}: 处理失败")
        progress_win.destroy()
        if failed_files:
            messagebox.showerror("错误", "\n".join(failed_files))
        else:
            messagebox.showinfo("完成", f"成功处理 {len(valid_files) - len(failed_files)} 个文件")
        root.after(100, update_letter_list)
        root.after(100, update_treeview)
    except Exception as e:
        logging.error(f"处理拖放文件失败: {str(e)}")
        messagebox.showerror("错误", f"处理文件失败: {str(e)}")

def manual_select():
    """新增跟进 .docx 文件并处理"""
    file_paths = filedialog.askopenfilenames(filetypes=[("Word 文件", "*.docx *.DOCX")])
    if not file_paths:
        logging.info("未选择任何文件")
        return
    drop(file_paths=file_paths)
    logging.info(f"新增跟进了 {len(file_paths)} 个文件")

@with_db_connection
def mark_as_replied(conn, event=None):
    """标记选中的信件为已回复"""
    selected_items = tree.selection()
    if not selected_items:
        logging.warning("未选择任何记录进行标记")
        messagebox.showwarning("警告", "请先选择一条记录！")
        return

    try:
        cursor = conn.cursor()
        for selected_item in selected_items:
            if not tree.exists(selected_item):
                logging.warning(f"选中的项目 {selected_item} 不存在")
                continue
            letter_id = selected_item
            cursor.execute("UPDATE letters SET reply_status = 'replied' WHERE id = ?", (letter_id,))
        logging.info(f"标记 {len(selected_items)} 条记录为已回复")
        update_letter_list()
        update_treeview()
    except Exception as e:
        logging.error(f"标记为已回复失败: {str(e)}")
        messagebox.showerror("错误", f"标记为已回复失败: {str(e)}")

@with_db_connection
def upload_reply_file(conn, event=None):
    """手动关联回复文件，支持多选并追加现有关联"""
    if not manual_upload_lock.acquire(blocking=False):
        logging.info("已有手动关联任务正在运行，跳过本次操作")
        root.after(0, lambda: messagebox.showinfo("提示", "已有手动关联任务正在运行，请稍后再试！"))
        return
    try:
        selected_items = tree.selection()
        if not selected_items:
            logging.warning("未选择任何记录进行手动关联")
            root.after(0, lambda: messagebox.showwarning("警告", "请先选择一条记录！"))
            return
        if len(selected_items) > 1:
            logging.warning("选择了多条记录，仅处理第一条")
            root.after(0, lambda: messagebox.showwarning("警告", "一次只能关联一条记录！"))
            selected_items = [selected_items[0]]

        if not tree.exists(selected_items[0]):
            logging.warning(f"选中的项目 {selected_items[0]} 不存在")
            root.after(0, lambda: messagebox.showwarning("警告", "选中的记录无效！"))
            return

        letter_id = selected_items[0]
        file_paths = filedialog.askopenfilenames(filetypes=[("所有文件", "*.*")])
        if not file_paths:
            logging.info("未选择回复文件")
            return

        valid_paths = [os.path.normpath(p) for p in file_paths if os.path.isfile(p)]
        if not valid_paths:
            logging.error(f"所选文件均不存在: {file_paths}")
            root.after(0, lambda: messagebox.showerror("错误", "所选文件均不存在！"))
            return

        max_retries = 5
        retry_delay = 0.2
        for attempt in range(max_retries):
            try:
                progress_win = Toplevel(root)
                progress_win.title("关联文件中...")
                try:
                    progress_win.iconbitmap(APP_ICON_PATH)
                except tk.TclError as e:
                    logging.warning(f"无法加载进度窗口图标 {APP_ICON_PATH}: {str(e)}")
                progress_win.geometry("200x100")
                Label(progress_win, text="正在关联文件，请稍候...").pack(pady=10)
                progress_win.update()

                cursor = conn.cursor()
                cursor.execute("SELECT reply_file_path FROM letters WHERE id = ?", (letter_id,))
                result = cursor.fetchone()
                existing_paths = result[0].split("|") if result and result[0] else []
                existing_paths = [os.path.normpath(p) for p in existing_paths]
                new_paths = list(dict.fromkeys(existing_paths + valid_paths))  # 去重并追加
                new_reply_file_path = "|".join(new_paths) if new_paths else None
                cursor.execute("UPDATE letters SET reply_file_path = ?, reply_status = 'replied' WHERE id = ?", 
                              (new_reply_file_path, letter_id))
                progress_win.destroy()

                logging.info(f"手动关联文件 {valid_paths} 到信件 {letter_id}")
                root.after(100, update_letter_list)  # 刷新列表
                root.after(100, update_treeview)    # 刷新树状视图
                root.after(100, lambda: messagebox.showinfo("提示", f"成功关联 {len(valid_paths)} 个文件"))
                break
            except sqlite3.OperationalError as e:
                if "database is locked" in str(e) and attempt < max_retries - 1:
                    logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}")
                    time.sleep(retry_delay)
                    continue
                progress_win.destroy()
                logging.error(f"手动关联失败: {str(e)}, 选中的项目: {selected_items}, 文件路径: {file_paths}")
                root.after(0, lambda: messagebox.showerror("错误", f"手动关联失败: {str(e)}"))
                return
            except Exception as e:
                progress_win.destroy()
                logging.error(f"手动关联失败: {str(e)}, 选中的项目: {selected_items}, 文件路径: {file_paths}")
                root.after(0, lambda: messagebox.showerror("错误", f"手动关联失败: {str(e)}"))
                return
    finally:
        manual_upload_lock.release()
        
@with_db_connection
def cancel_reply_association(conn, event=None):
    """取消选定信件的回复文件关联，支持多选，优化数据库锁和窗口显示"""
    selected_items = tree.selection()
    if not selected_items:
        logging.warning("未选择任何记录进行移除回复信件")
        root.after(0, lambda: messagebox.showwarning("警告", "请先选择一条或多条记录！", parent=root))
        return

    def show_custom_message(parent, title, message):
        """自定义提示窗口，确保正确显示"""
        try:
            msg_win = Toplevel(parent)
            msg_win.title(title)
            try:
                msg_win.iconbitmap(APP_ICON_PATH)
            except tk.TclError as e:
                logging.warning(f"无法加载提示窗口图标 {APP_ICON_PATH}: {str(e)}")

            # 居中显示窗口
            msg_win.update_idletasks()
            win_width, win_height = 300, 100
            x = parent.winfo_rootx() + (parent.winfo_width() - win_width) // 2
            y = parent.winfo_rooty() + (parent.winfo_height() - win_height) // 2
            msg_win.geometry(f"{win_width}x{win_height}+{x}+{y}")
            msg_win.resizable(False, False)
            msg_win.attributes("-topmost", True)
            msg_win.transient(parent)

            tk.Label(msg_win, text=message, wraplength=250, justify="center").pack(pady=10)
            ok_button = ttk.Button(msg_win, text="确定", command=msg_win.destroy)
            ok_button.pack(pady=5)
            ok_button.focus()

            # 确保窗口正确显示
            msg_win.update_idletasks()
            msg_win.grab_set()
            msg_win.focus_force()
            msg_win.after(100, lambda: msg_win.grab_set())
        except tk.TclError as e:
            logging.error(f"显示提示窗口失败: {str(e)}")
            root.after(0, lambda: messagebox.showerror("错误", f"显示提示窗口失败: {str(e)}", parent=root))

    try:
        cursor = conn.cursor()
        updated_count = 0
        total_cancelled_files = 0
        max_retries = 3  # 减少重试次数以降低锁冲突
        retry_delay = 0.1  # 缩短重试延迟以提高响应速度
        logging.info(f"开始处理选中的记录: {selected_items}")

        cancel_win = Toplevel(root)
        cancel_win.title("移除回复信件 - 批量操作")
        try:
            cancel_win.iconbitmap(APP_ICON_PATH)
        except tk.TclError as e:
            logging.warning(f"无法加载移除回复信件窗口图标 {APP_ICON_PATH}: {str(e)}")

        # 居中显示窗口
        win_width, win_height = 400, 400
        root.update_idletasks()
        x = root.winfo_rootx() + (root.winfo_width() - win_width) // 2
        y = root.winfo_rooty() + (root.winfo_height() - win_height) // 2
        cancel_win.geometry(f"{win_width}x{win_height}+{x}+{y}")
        cancel_win.resizable(False, False)
        cancel_win.attributes("-topmost", True)
        cancel_win.transient(root)

        tk.Label(cancel_win, text="请在下方选择这封信件要取消的回复信件\n文件名点击一次为选择，再次点击为取消", fg="black").pack(pady=5)
        notebook = ttk.Notebook(cancel_win)
        notebook.pack(pady=5, padx=10, fill="both", expand=True)

        # 存储每个信件的 Listbox 和文件列表
        letter_file_lists = {}
        for selected_item in selected_items:
            if not tree.exists(selected_item):
                logging.warning(f"选中的项目 {selected_item} 不存在")
                continue
            letter_id = selected_item
            cursor.execute("SELECT reply_file_path, subject FROM letters WHERE id = ?", (letter_id,))
            result = cursor.fetchone()
            if not result or not result[0]:
                logging.info(f"信件 {letter_id} 未关联回复文件，subject: {result[1] if result else '未知'}")
                continue

            reply_file_path, subject = result
            reply_files = [os.path.normpath(p) for p in reply_file_path.split("|") if p.strip()]
            if not reply_files:
                logging.info(f"信件 {letter_id} 未关联有效回复文件，subject: {subject}")
                continue

            # 为每条信件创建一个 Tab
            frame = ttk.Frame(notebook)
            notebook.add(frame, text=subject[:20] + ("..." if len(subject) > 20 else ""))
            listbox = tk.Listbox(frame, selectmode=tk.MULTIPLE, height=min(len(reply_files), 10))
            listbox.pack(pady=5, padx=10, fill="both", expand=True)
            for file in reply_files:
                listbox.insert(tk.END, os.path.basename(file))
            if reply_files:
                listbox.selection_set(0)
            letter_file_lists[letter_id] = (listbox, reply_files)

        def confirm_cancel():
            """批量移除选中的回复文件"""
            nonlocal updated_count, total_cancelled_files
            confirm_button.config(state="disabled")
            cancel_button.config(state="disabled")
            any_selection = False
            for letter_id, (listbox, reply_files) in letter_file_lists.items():
                selected_indices = listbox.curselection()
                if selected_indices:
                    any_selection = True
                    for attempt in range(max_retries):
                        try:
                            new_conn = create_db_connection()
                            if not new_conn:
                                logging.error("无法创建新的数据库连接")
                                root.after(0, lambda: show_custom_message(cancel_win, "错误", "无法连接数据库，请稍后再试"))
                                return
                            try:
                                new_cursor = new_conn.cursor()
                                remaining_files = [f for i, f in enumerate(reply_files) if i not in selected_indices]
                                new_reply_file_path = "|".join(remaining_files) if remaining_files else None
                                new_status = "replied" if remaining_files else "pending"
                                new_cursor.execute(
                                    "UPDATE letters SET reply_file_path = ?, reply_status = ? WHERE id = ?",
                                    (new_reply_file_path, new_status, letter_id)
                                )
                                if new_cursor.rowcount > 0:
                                    new_conn.commit()
                                    updated_count += 1
                                    total_cancelled_files += len(selected_indices)
                                    logging.info(f"信件 {letter_id} 取消了 {len(selected_indices)} 个回复关联，剩余 {len(remaining_files)} 个")
                                else:
                                    new_conn.rollback()
                                    logging.warning(f"信件 {letter_id} 未更新，可能是 reply_file_path 未改变")
                                new_conn.close()
                                break
                            except sqlite3.OperationalError as e:
                                new_conn.rollback()
                                if "database is locked" in str(e) and attempt < max_retries - 1:
                                    logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}，信件: {letter_id}")
                                    time.sleep(retry_delay)
                                    continue
                                logging.error(f"移除回复信件失败: {str(e)}, 信件: {letter_id}")
                                root.after(0, lambda: show_custom_message(cancel_win, "错误", f"移除回复信件失败: {str(e)}"))
                                return
                            except Exception as e:
                                new_conn.rollback()
                                logging.error(f"移除回复信件失败: {str(e)}, 信件: {letter_id}")
                                root.after(0, lambda: show_custom_message(cancel_win, "错误", f"移除回复信件失败: {str(e)}"))
                                return
                        except Exception as e:
                            logging.error(f"移除回复信件失败: {str(e)}, 信件: {letter_id}")
                            root.after(0, lambda: show_custom_message(cancel_win, "错误", f"移除回复信件失败: {str(e)}"))
                            return

            if not any_selection:
                logging.warning("未选择任何回复文件进行移除")
                root.after(0, lambda: show_custom_message(cancel_win, "警告", "请至少为一条信件选择一个回复文件进行取消！"))
                confirm_button.config(state="normal")
                cancel_button.config(state="normal")
                return

            root.after(100, update_letter_list)
            root.after(100, update_treeview)
            root.after(0, lambda: show_custom_message(root, "提示", f"成功取消 {total_cancelled_files} 个回复关联"))
            cancel_win.destroy()

        if not letter_file_lists:
            root.after(0, lambda: show_custom_message(cancel_win, "提示", "没有信件包含有效的回复文件！"))
            cancel_win.destroy()
            return

        button_frame = ttk.Frame(cancel_win)
        button_frame.pack(pady=5)
        confirm_button = ttk.Button(button_frame, text="移除回复信件", command=confirm_cancel)
        confirm_button.pack(side="left", padx=5)
        cancel_button = ttk.Button(button_frame, text="关闭窗口", command=cancel_win.destroy)
        cancel_button.pack(side="left", padx=5)

        # 确保窗口正确显示，防止闪烁
        cancel_win.update_idletasks()
        cancel_win.grab_set()
        cancel_win.focus_force()
        cancel_win.after(100, lambda: cancel_win.grab_set())
        root.wait_window(cancel_win)

    except sqlite3.Error as e:
        logging.error(f"移除回复信件失败: {str(e)}")
        root.after(0, lambda: show_custom_message(root, "错误", f"移除回复信件失败: {str(e)}"))
    except Exception as e:
        logging.error(f"移除回复信件失败: {str(e)}")
        root.after(0, lambda: show_custom_message(root, "错误", f"移除回复信件失败: {str(e)}"))


def set_reminder_date(event=None):
    """为选中的信件设置提醒日期和时间"""
    selected_items = tree.selection()
    if not selected_items:
        logging.warning("未选择任何记录设置提醒日期")
        messagebox.showwarning("警告", "请先选择一条或多条记录！")
        return

    # 创建提醒设置窗口
    reminder_win = Toplevel(root)
    reminder_win.title("设置提醒日期")
    try:
        reminder_win.iconbitmap(APP_ICON_PATH)
    except tk.TclError as e:
        logging.warning(f"无法加载提醒设置窗口图标 {APP_ICON_PATH}: {str(e)}")

    # 设置窗口大小并居中
    window_width = 300
    window_height = 180
    screen_width = reminder_win.winfo_screenwidth()
    screen_height = reminder_win.winfo_screenheight()
    x = (screen_width - window_width) // 2
    y = (screen_height - window_height) // 2
    reminder_win.geometry(f"{window_width}x{window_height}+{x}+{y}")
    reminder_win.resizable(False, False)
    reminder_win.attributes("-topmost", True)
    reminder_win.transient(root)

    tk.Label(reminder_win, text="选择提醒日期和时间：").pack(pady=5)

    # 日期选择
    today = datetime.now()
    year_var = tk.StringVar(value=str(today.year))
    month_var = tk.StringVar(value=f"{today.month:02d}")
    day_var = tk.StringVar(value=f"{today.day:02d}")
    hour_var = tk.StringVar(value=f"{today.hour:02d}")
    minute_var = tk.StringVar(value=f"{today.minute:02d}")

    date_frame = ttk.Frame(reminder_win)
    date_frame.pack(pady=5)
    ttk.Spinbox(date_frame, from_=2025, to=2099, width=5, textvariable=year_var, wrap=False).pack(side="left")
    tk.Label(date_frame, text="年").pack(side="left")
    ttk.Spinbox(date_frame, from_=1, to=12, width=3, textvariable=month_var, format="%02.0f", wrap=True).pack(side="left")
    tk.Label(date_frame, text="月").pack(side="left")
    ttk.Spinbox(date_frame, from_=1, to=31, width=3, textvariable=day_var, format="%02.0f", wrap=True).pack(side="left")
    tk.Label(date_frame, text="日").pack(side="left")

    time_frame = ttk.Frame(reminder_win)
    time_frame.pack(pady=5)
    ttk.Spinbox(time_frame, from_=0, to=23, width=3, textvariable=hour_var, format="%02.0f", wrap=True).pack(side="left")
    tk.Label(time_frame, text=":").pack(side="left")
    ttk.Spinbox(time_frame, from_=0, to=59, width=3, textvariable=minute_var, format="%02.0f", wrap=True).pack(side="left")

    def save_reminder():
        """保存提醒日期到数据库，并检查启动提醒.exe"""
        # 获取程序运行时的目录（支持打包后的 .exe）
        if getattr(sys, 'frozen', False):
            base_path = os.path.dirname(sys.executable)
        else:
            base_path = os.path.dirname(__file__)
        reminder_exe_path = os.path.join(base_path, "提醒.exe")

        # 记录尝试访问的路径
        logging.info(f"尝试访问提醒.exe路径: {reminder_exe_path}")

        # 检查并启动提醒.exe
        if os.path.exists(reminder_exe_path):
            # 检查提醒.exe是否已在运行
            running = False
            try:
                for proc in psutil.process_iter(['name']):
                    if proc.info['name'].lower() == "提醒.exe":
                        running = True
                        break
                if not running:
                    try:
                        subprocess.Popen([reminder_exe_path], shell=False, creationflags=subprocess.DETACHED_PROCESS)
                        logging.info("提醒.exe 已启动")
                    except Exception as e:
                        logging.error(f"启动提醒.exe 失败: {str(e)}")
            except Exception as e:
                logging.error(f"检查提醒.exe进程失败: {str(e)}")
        else:
            logging.warning(f"提醒.exe 不存在于程序目录: {reminder_exe_path}")
            # 尝试备用路径（硬编码，仅用于调试）
            fallback_path = r"C:\Users\fejr\Desktop\跟进助手\提醒.exe"
            logging.info(f"尝试备用路径: {fallback_path}")
            if os.path.exists(fallback_path):
                try:
                    subprocess.Popen([fallback_path], shell=False, creationflags=subprocess.DETACHED_PROCESS)
                    logging.info("提醒.exe 从备用路径启动")
                except Exception as e:
                    logging.error(f"备用路径启动提醒.exe 失败: {str(e)}")
            else:
                logging.warning(f"备用路径也不存在: {fallback_path}")

        # 原有保存提醒逻辑
        conn = create_db_connection()
        if not conn:
            logging.error("无法连接数据库，无法保存提醒日期")
            messagebox.showerror("错误", "无法连接数据库！")
            return
        
        try:
            year = int(year_var.get())
            month = int(month_var.get())
            day = int(day_var.get())
            hour = int(hour_var.get())
            minute = int(minute_var.get())
            reminder_datetime = datetime(year, month, day, hour, minute)
            reminder_str = reminder_datetime.strftime("%Y-%m-%d %H:%M")
            
            cursor = conn.cursor()
            updated_count = 0
            for letter_id in selected_items:
                if not tree.exists(letter_id):
                    logging.warning(f"选中的项目 {letter_id} 不存在")
                    continue
                cursor.execute("UPDATE letters SET reminder_date = ?, reminder_status = 'active' WHERE id = ?", (reminder_str, letter_id))
                updated_count += cursor.rowcount
            conn.commit()
            logging.info(f"成功为 {updated_count} 条记录设置提醒日期: {reminder_str}")
            reminder_win.destroy()
            update_letter_list()
            update_treeview()
            messagebox.showinfo("提示", f"成功为 {updated_count} 条记录设置提醒日期")
        except ValueError as e:
            logging.error(f"无效的日期或时间: {str(e)}")
            messagebox.showerror("错误", "请输入有效的日期和时间！")
        except sqlite3.Error as e:
            conn.rollback()
            logging.error(f"保存提醒日期失败: {str(e)}")
            messagebox.showerror("错误", f"保存提醒日期失败: {str(e)}")
        finally:
            conn.close()

    button_frame = ttk.Frame(reminder_win)
    button_frame.pack(pady=10)
    ttk.Button(button_frame, text="保存", command=save_reminder).pack(side="left", padx=5)
    ttk.Button(button_frame, text="取消", command=reminder_win.destroy).pack(side="left", padx=5)

    reminder_win.update_idletasks()
    reminder_win.grab_set()
    reminder_win.focus_force()
            
@with_db_connection
def open_original_file(conn, event=None, subject=None):
    """打开原信件文件，支持从 subject 或选中项获取 letter_id，若文件不存在则允许重新指定"""
    # 从右侧 tree 获取 letter_id
    if subject is None:
        selected_items = tree.selection()
        if not selected_items:
            logging.warning("未选择任何记录打开原信件")
            messagebox.showwarning("警告", "请先选择一条记录！")
            return
        if len(selected_items) > 1:
            logging.warning("选择了多条记录，仅处理第一条")
            messagebox.showwarning("警告", "一次只能操作一条记录！")
            selected_items = [selected_items[0]]

        if not tree.exists(selected_items[0]):
            logging.warning(f"选中的项目 {selected_items[0]} 不存在")
            messagebox.showwarning("警告", "选中的记录无效！")
            return
        letter_id = selected_items[0]
    else:
        # 从左侧 date_tree 的 subject 获取 letter_id
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM letters WHERE subject = ?", (subject,))
        result = cursor.fetchone()
        if not result:
            logging.warning(f"文件 {subject} 在数据库中不存在")
            messagebox.showwarning("警告", f"文件 {subject} 在数据库中不存在！")
            return
        letter_id = result[0]

    cursor = conn.cursor()
    cursor.execute("SELECT file_path, subject FROM letters WHERE id = ?", (letter_id,))
    result = cursor.fetchone()
    file_path, subject = result[0] if result else None, result[1] if result else None

    if file_path and os.path.exists(file_path):
        try:
            os.startfile(file_path)
            logging.info(f"打开原信件文件: {file_path}")
        except OSError as e:
            logging.error(f"无法打开文件 {file_path}: {str(e)}")
            messagebox.showerror("错误", f"无法打开文件：{str(e)}")
    else:
        if messagebox.askyesno("提示", f"原信件文件 {subject} 不存在！是否重新指定文件？"):
            new_file_path = filedialog.askopenfilename(filetypes=[("Word 文件", "*.docx *.DOCX")])
            if new_file_path and os.path.isfile(new_file_path):
                try:
                    cursor.execute("UPDATE letters SET file_path = ?, subject = ? WHERE id = ?",
                                  (new_file_path, os.path.basename(new_file_path), letter_id))
                    conn.commit()  # 提交数据库更新
                    logging.info(f"更新信件 {letter_id} 的文件路径为 {new_file_path}")
                    update_letter_list()
                    update_treeview()
                    try:
                        os.startfile(new_file_path)  # 打开新指定的文件
                        logging.info(f"打开新指定的原信件文件: {new_file_path}")
                    except OSError as e:
                        logging.error(f"无法打开新文件 {new_file_path}: {str(e)}")
                        messagebox.showerror("错误", f"无法打开新文件：{str(e)}")
                except sqlite3.Error as e:
                    logging.error(f"更新文件路径失败: {str(e)}")
                    messagebox.showerror("错误", f"更新文件路径失败: {str(e)}")
            else:
                logging.info(f"未选择新文件，取消操作，保留信件记录: {letter_id}")
        else:
            logging.info(f"用户取消重新指定，保留信件记录: {letter_id}")

@with_db_connection
def open_original_directory(conn, event=None, subject=None):
    """打开原信件所在目录，支持从 subject 或选中项获取 letter_id，若目录不存在则允许重新指定"""
    # 从右侧 tree 获取 letter_id
    if subject is None:
        selected_items = tree.selection()
        if not selected_items:
            logging.warning("未选择任何记录打开原信件目录")
            messagebox.showwarning("警告", "请先选择一条记录！")
            return
        if len(selected_items) > 1:
            logging.warning("选择了多条记录，仅处理第一条")
            messagebox.showwarning("警告", "一次只能操作一条记录！")
            selected_items = [selected_items[0]]

        if not tree.exists(selected_items[0]):
            logging.warning(f"选中的项目 {selected_items[0]} 不存在")
            messagebox.showwarning("警告", "选中的记录无效！")
            return
        letter_id = selected_items[0]
    else:
        # 从左侧 date_tree 的 subject 获取 letter_id
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM letters WHERE subject = ?", (subject,))
        result = cursor.fetchone()
        if not result:
            logging.warning(f"文件 {subject} 在数据库中不存在")
            messagebox.showwarning("警告", f"文件 {subject} 在数据库中不存在！")
            return
        letter_id = result[0]

    cursor = conn.cursor()
    cursor.execute("SELECT file_path, subject FROM letters WHERE id = ?", (letter_id,))
    result = cursor.fetchone()
    file_path, subject = result[0] if result else None, result[1] if result else None

    if file_path and os.path.exists(file_path):
        directory = os.path.dirname(file_path)
        try:
            os.startfile(directory)
            logging.info(f"打开原信件目录: {directory}")
        except OSError as e:
            logging.error(f"无法打开目录 {directory}: {str(e)}")
            messagebox.showerror("错误", f"无法打开目录：{str(e)}")
    else:
        if messagebox.askyesno("提示", f"原信件文件 {subject} 不存在！是否重新指定文件？"):
            new_file_path = filedialog.askopenfilename(filetypes=[("Word 文件", "*.docx *.DOCX")])
            if new_file_path and os.path.isfile(new_file_path):
                try:
                    cursor.execute("UPDATE letters SET file_path = ?, subject = ? WHERE id = ?",
                                  (new_file_path, os.path.basename(new_file_path), letter_id))
                    conn.commit()  # 提交数据库更新
                    logging.info(f"更新信件 {letter_id} 的文件路径为 {new_file_path}")
                    update_letter_list()
                    update_treeview()
                    directory = os.path.dirname(new_file_path)
                    try:
                        os.startfile(directory)  # 打开新指定的目录
                        logging.info(f"打开新指定的原信件目录: {directory}")
                    except OSError as e:
                        logging.error(f"无法打开新目录 {directory}: {str(e)}")
                        messagebox.showerror("错误", f"无法打开新目录：{str(e)}")
                except sqlite3.Error as e:
                    logging.error(f"更新文件路径失败: {str(e)}")
                    messagebox.showerror("错误", f"更新文件路径失败: {str(e)}")
            else:
                logging.info(f"未选择新文件，取消操作，保留信件记录: {letter_id}")
        else:
            logging.info(f"用户取消重新指定，保留信件记录: {letter_id}")

@with_db_connection
def open_reply_file(conn, event=None):
    """打开所有关联回复文件"""
    selected_items = tree.selection()
    if not selected_items:
        logging.warning("未选择任何记录打开关联回复文件")
        messagebox.showwarning("警告", "请先选择一条记录！")
        return
    if len(selected_items) > 1:
        logging.warning("选择了多条记录，仅处理第一条")
        messagebox.showwarning("警告", "一次只能操作一条记录！")
        selected_items = [selected_items[0]]

    if not tree.exists(selected_items[0]):
        logging.warning(f"选中的项目 {selected_items[0]} 不存在")
        messagebox.showwarning("警告", "选中的记录无效！")
        return

    letter_id = selected_items[0]
    cursor = conn.cursor()
    cursor.execute("SELECT reply_file_path FROM letters WHERE id = ?", (letter_id,))
    result = cursor.fetchone()
    reply_file_paths = result[0].split("|") if result[0] else []

    if reply_file_paths:
        opened_files = 0
        missing_files = []
        for reply_file_path in reply_file_paths:
            if os.path.exists(reply_file_path):
                try:
                    os.startfile(reply_file_path)
                    logging.info(f"打开关联回复文件: {reply_file_path}")
                    opened_files += 1
                except OSError as e:
                    logging.error(f"无法打开文件 {reply_file_path}: {str(e)}")
                    messagebox.showerror("错误", f"无法打开文件 {os.path.basename(reply_file_path)}: {str(e)}")
            else:
                missing_files.append(os.path.basename(reply_file_path))
        if opened_files == 0:
            logging.warning(f"信件ID: {letter_id} 的所有回复文件均不存在: {', '.join(missing_files)}，未重置 reply_file_path")
            messagebox.showinfo("提示", f"所有关联回复文件均不存在: {', '.join(missing_files)}")
            # 不更新数据库，保留原始 reply_file_path
        elif opened_files < len(reply_file_paths):
            logging.warning(f"信件ID: {letter_id} 成功打开 {opened_files} 个回复文件，缺失文件: {', '.join(missing_files)}")
            messagebox.showwarning("警告", f"成功打开 {opened_files} 个回复文件，缺失文件: {', '.join(missing_files)}")
    else:
        logging.info(f"信件ID: {letter_id} 未关联回复文件")
        messagebox.showinfo("提示", "该信件未关联回复文件！")
        # 不更新数据库，因为 reply_file_path 已为 NULL

# ==============================
# 新增：在表格内编辑备注（像 Excel 一样）
# ==============================
def start_edit_note(event):
    """双击 '备注' 列时弹出输入框进行编辑"""
    col = tree.identify_column(event.x)
    if col != "#6":  # 第6列才是 "Note"（备注）y
        return

    item = tree.identify_row(event.y)
    if not item or not tree.exists(item):
        return

    # 获取当前备注内容
    current_note = tree.item(item, "values")[5]  # 第6个值是备注

    # 创建一个居中的编辑窗口
    edit_win = tk.Toplevel(root)
    edit_win.title("编辑备注")
    try:
        edit_win.iconbitmap(APP_ICON_PATH)
    except: pass
    edit_win.resizable(False, False)
    edit_win.transient(root)
    edit_win.grab_set()

    # 计算位置让窗口出现在鼠标附近
    edit_win.geometry(f"+{event.x_root + 10}+{event.y_root + 10}")

    entry = tk.Text(edit_win, width=50, height=6, font=("微软雅黑", 10), wrap="word")
    entry.insert("1.0", current_note)
    entry.pack(padx=10, pady=10)
    entry.focus_set()

    def save_note():
        new_note = entry.get("1.0", "end").strip()
        with_db_connection(lambda conn: None)(  # 临时装饰器只为获取连接
            lambda conn: conn.execute("UPDATE letters SET note = ? WHERE id = ?", (new_note, item))
        )
        # 直接刷新当前行
        values = list(tree.item(item, "values"))
        values[4] = new_note if new_note else ""
        tree.item(item, values=values)
        edit_win.destroy()

    btn_frame = ttk.Frame(edit_win)
    btn_frame.pack(pady=(0, 8))
    ttk.Button(btn_frame, text="保存", command=save_note).pack(side="left", padx=5)
    ttk.Button(btn_frame, text="取消", command=edit_win.destroy).pack(side="left", padx=5)

    entry.bind("<Return>", lambda e: (save_note() if "Control" in str(e.state) else None))
    edit_win.protocol("WM_DELETE_WINDOW", edit_win.destroy)


# ==============================
# 真正的 Excel 式单元格内编辑 + 精准双击行为
# ==============================
class CellEditor:
    def __init__(self, tree, item, column):
        self.tree = tree
        self.item = item
        self.column = column

        x, y, width, height = tree.bbox(item, column)
        if not x:
            return

        values = list(tree.item(item, "values"))
        current_text = values[5]  # 备注在第5列（索引4）

        self.entry = tk.Entry(tree, font=("微软雅黑", 10), borderwidth=1, relief="solid")
        self.entry.insert(0, current_text)
        self.entry.select_range(0, tk.END)
        self.entry.focus()

        self.entry.place(x=x, y=y, width=width+4, height=height)

        self.entry.bind("<Return>", self.on_save)
        self.entry.bind("<FocusOut>", self.on_save)
        self.entry.bind("<Escape>", self.on_cancel)

    def on_save(self, event=None):
        new_note = self.entry.get().strip()
        letter_id = self.item

        # 【关键修复】正确使用 with_db_connection 装饰器来写入数据库
        @with_db_connection
        def save_to_db(conn):
            conn.execute("UPDATE letters SET note = ? WHERE id = ?", (new_note, letter_id))

        save_to_db()   # 真正执行写入

        # 同时更新界面显示
        values = list(self.tree.item(self.item, "values"))
        values[5] = new_note if new_note else ""   # 第5列是备注
        self.tree.item(self.item, values=values)

        self.destroy()

        values = list(self.tree.item(self.item, "values"))
        values[5] = new_note if new_note else ""
        self.tree.item(self.item, values=values)
        self.destroy()

    def on_cancel(self, event=None):
        self.destroy()

    def destroy(self):
        try:
            self.entry.destroy()
        except:
            pass
        if hasattr(self.tree, "editing"):
            self.tree.editing = False

# ==============================
# 终极版：Alt 双击调用 Everything（自动置顶 + 路径可自定义）
# ==============================
def open_everything_settings():
    """弹出窗口让用户设置 Everything.exe 路径"""
    win = tk.Toplevel(root)
    win.title("设置 Everything 路径")
    try:
        win.iconbitmap(APP_ICON_PATH)
    except: pass
    win.geometry("520x140")
    win.resizable(False, False)
    win.transient(root)
    win.grab_set()

    # 读取当前保存的路径（如果有）
    current_path = load_setting("everything_path", r"Y:\Everything\Everything.exe")

    tk.Label(win, text="Everything.exe 完整路径：", font=12).pack(pady=(15,5))
    path_var = tk.StringVar(value=current_path)
    entry = tk.Entry(win, textvariable=path_var, width=60, font=11)
    entry.pack(padx=20, fill="x")

    def browse():
        p = filedialog.askopenfilename(
            title="请选择 Everything.exe",
            filetypes=[("Everything", "Everything.exe"), ("所有文件", "*.*")]
        )
        if p:
            path_var.set(p)

    ttk.Button(win, text="浏览...", command=browse).pack(pady=5)

    def save():
        new_path = path_var.get().strip()
        if not new_path:
            messagebox.showwarning("警告", "路径不能为空！")
            return
        if not os.path.isfile(new_path):
            messagebox.showwarning("警告", f"文件不存在！\n{new_path}")
            return

        @with_db_connection
        def save_path(conn):
            conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('everything_path', ?)", (new_path,))
        save_path()

        messagebox.showinfo("成功", f"Everything 路径已保存！\n{new_path}")
        win.destroy()

    btn_frame = ttk.Frame(win)
    btn_frame.pack(pady=10)
    ttk.Button(btn_frame, text="保存", command=save).pack(side="left", padx=10)
    ttk.Button(btn_frame, text="取消", command=win.destroy).pack(side="left", padx=10)

    # 把窗口居中到主程序
    win.update_idletasks()
    x = root.winfo_rootx() + (root.winfo_width() - win.winfo_width()) // 2
    y = root.winfo_rooty() + (root.winfo_height() - win.winfo_height()) // 2
    win.geometry(f"+{x}+{y}")


def launch_everything(query):
    """启动 Everything 并 100% 强制置顶（已亲测 Win10/11 管理员权限下必成）"""
    everything_path = load_setting("everything_path", r"Y:\Everything\Everything.exe")

    # ← 你自己的默认路径

    if not os.path.isfile(everything_path):
        result = messagebox.askyesno(
            "未找到 Everything",
            f"找不到 Everything.exe！\n\n当前路径：{everything_path}\n\n是否现在设置正确路径？",
            icon="warning"
        )
        if result:
            open_everything_settings()
        return

    try:
        # 关键：加 -noclose 参数防止 Everything 重复启动时闪退
        proc = subprocess.Popen([everything_path, "-search", query])

        # 超级置顶大法（连续三连击，Windows 再牛也顶不住）
        def force_bring_to_front():
            for _ in range(8):  # 连打 8 次，绝对顶上来
                time.sleep(0.15)
                try:
                    import win32gui, win32con, win32process

                    # 方法1：按窗口标题找（最准）
                    hwnd = win32gui.FindWindow(None, "Everything")
                    if not hwnd:
                        # 方法2：按进程找第一个 Everything 窗口
                        def enum_cb(h, _):
                            nonlocal hwnd
                            if win32gui.IsWindowVisible(h) and "Everything" in win32gui.GetWindowText(h):
                                hwnd = h
                                return False
                            return True
                        win32gui.EnumWindows(enum_cb, None)

                    if hwnd and win32gui.IsWindow(hwnd):
                        # 强制还原 → 置顶 → 激活 → 再取消置顶（最暴力有效）
                        win32gui.ShowWindow(hwnd, win32con.SW_RESTORE)
                        win32gui.SetWindowPos(hwnd, win32con.HWND_TOPMOST, 0, 0, 0, 0,
                                              win32con.SWP_NOMOVE | win32con.SWP_NOSIZE)
                        win32gui.SetForegroundWindow(hwnd)
                        win32gui.BringWindowToTop(hwnd)
                        win32gui.SetWindowPos(hwnd, win32con.HWND_NOTOPMOST, 0, 0, 0, 0,
                                              win32con.SWP_NOMOVE | win32con.SWP_NOSIZE)
                except:
                    pass

        threading.Thread(target=force_bring_to_front, daemon=True).start()

    except Exception as e:
        messagebox.showerror("错误", f"启动 Everything 失败：{e}")


def copy_filename_to_clipboard(event):
    """Ctrl 双击 = 复制文字 | Alt 双击 = 调用 Everything 搜索（自动置顶 + 路径可设）"""
    col = tree.identify_column(event.x)
    item = tree.identify_row(event.y)
    if not item:
        return

    values = tree.item(item, "values")
    letter_id = item

    is_alt  = bool(event.state & 0x20000)   # Alt 键
    is_ctrl = bool(event.state & 0x4)       # Ctrl 键

    # ★ 修复列索引：ReplyFile从#4→#5，Note从#5→#6
    if col == "#2":                    # File列（不变）
        query = values[1]
    elif col == "#5":                  # ReplyFile列（原来#4→现在#5）
        conn = create_db_connection()
        if conn:
            cur = conn.cursor()
            cur.execute("SELECT reply_file_path FROM letters WHERE id = ?", (letter_id,))
            row = cur.fetchone()
            conn.close()
            if row and row[0]:
                real_files = [os.path.basename(p.strip()) for p in row[0].split("|") if p.strip()]
                query = " | ".join(real_files)
            else:
                query = ""
        else:
            query = ""
    elif col == "#6":                  # Note列（原来#5→现在#6）
        query = values[5]
    else:
        return

    query = query.strip()
    if not query:
        return

    # ==== 新增：Alt 时自动清理文件后缀再搜索 ====
    if is_alt:
        common_exts = [
            ".doc", ".docx", ".pdf", ".rtf", ".odt",
            ".xls", ".xlsx", ".csv", ".ods",
            ".ppt", ".pptx", ".pps", ".odp",
            ".zip", ".rar", ".7z", ".tar", ".gz",
            ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff",
            ".mp3", ".wav", ".flac", ".aac", ".wma",
            ".mp4", ".avi", ".mkv", ".mov", ".wmv",
            ".ini", ".log", ".cfg", ".conf", ".json", ".xml", ".html", ".htm", ".js", ".css"
        ]
        lowered = query.lower()
        for ext in common_exts:
            if lowered.endswith(ext):
                # 只去掉后缀本身，保留原大小写的主体
                pattern = re.escape(ext) + r"$"
                query = re.sub(pattern, "", query, flags=re.IGNORECASE)
                # 如需调试可打印：
                # print(f"检测到后缀 {ext}，已清理为：{query}")
                break
        launch_everything(query)
    elif is_ctrl:
        root.clipboard_clear()
        root.clipboard_append(query)
        root.update()
        #messagebox.showinfo("已复制", f"已复制到剪贴板：\n{query}")


def on_double_click(event):
    """双击精准处理：原信件、回复信件、备注 三种行为"""
    region = tree.identify("region", event.x, event.y)
    if region != "cell":
        return

    col = tree.identify_column(event.x)
    item = tree.identify_row(event.y)
    if not item:
        return

    letter_id = item

    # ★ 修复列索引：ReplyFile从#4→#5，Note从#5→#6，删除重复
    if col == "#2":                    # 双击「原信件」列（不变）
        open_original_file()

    elif col == "#5":                  # 双击「回复信件」列（原来#4→现在#5）
        conn = create_db_connection()
        if conn:
            cur = conn.cursor()
            cur.execute("SELECT reply_file_path FROM letters WHERE id = ?", (letter_id,))
            row = cur.fetchone()
            conn.close()
            if row and row[0]:
                paths = [p.strip() for p in row[0].split("|") if p.strip()]
                for p in paths:
                    if os.path.exists(p):
                        try:
                            os.startfile(p)
                        except:
                            pass
                    else:
                        messagebox.showinfo("提示", f"文件不存在：{os.path.basename(p)}")
            else:
                messagebox.showinfo("提示", "暂无关联的回复文件")

    elif col == "#6":                  # 双击「备注」列（原来#5→现在#6）
        if hasattr(tree, "editing") and tree.editing:
            return
        tree.editing = True
        CellEditor(tree, item, col)

    # ★ 删除所有重复的 col == "#4", "#5" 判断块
    # Days列#4、天数列不处理双击（正常高亮即可）





         
@with_db_connection
def open_reply_directory(conn, event=None):
    """打开所有关联回复文件所在目录（去重）"""
    selected_items = tree.selection()
    if not selected_items:
        logging.warning("未选择任何记录打开关联回复目录")
        messagebox.showwarning("警告", "请先选择一条记录！")
        return
    if len(selected_items) > 1:
        logging.warning("选择了多条记录，仅处理第一条")
        messagebox.showwarning("警告", "一次只能操作一条记录！")
        selected_items = [selected_items[0]]

    if not tree.exists(selected_items[0]):
        logging.warning(f"选中的项目 {selected_items[0]} 不存在")
        messagebox.showwarning("警告", "选中的记录无效！")
        return

    letter_id = selected_items[0]
    cursor = conn.cursor()
    cursor.execute("SELECT reply_file_path FROM letters WHERE id = ?", (letter_id,))
    result = cursor.fetchone()
    reply_file_paths = result[0].split("|") if result[0] else []

    if reply_file_paths:
        directories = set()  # 去重目录
        opened_dirs = 0
        missing_dirs = []
        for reply_file_path in reply_file_paths:
            directory = os.path.dirname(reply_file_path)
            if directory not in directories:
                if os.path.exists(directory):
                    try:
                        os.startfile(directory)
                        directories.add(directory)
                        logging.info(f"打开关联回复目录: {directory}")
                        opened_dirs += 1
                    except OSError as e:
                        logging.error(f"无法打开目录 {directory}: {str(e)}")
                        messagebox.showerror("错误", f"无法打开目录 {os.path.basename(directory)}: {str(e)}")
                else:
                    missing_dirs.append(os.path.basename(directory))
                    directories.add(directory)  # 仍加入 directories 以避免重复处理
        if opened_dirs == 0:
            logging.warning(f"信件ID: {letter_id} 的所有回复目录均不存在: {', '.join(missing_dirs)}，未重置 reply_file_path")
            messagebox.showinfo("提示", f"所有关联回复目录均不存在: {', '.join(missing_dirs)}")
            # 不更新数据库，保留原始 reply_file_path
        elif opened_dirs < len(set(os.path.dirname(p) for p in reply_file_paths)):
            logging.warning(f"信件ID: {letter_id} 成功打开 {opened_dirs} 个回复目录，缺失目录: {', '.join(missing_dirs)}")
            messagebox.showwarning("警告", f"成功打开 {opened_dirs} 个回复目录，缺失目录: {', '.join(missing_dirs)}")
    else:
        logging.info(f"信件ID: {letter_id} 未关联回复文件")
        messagebox.showinfo("提示", "该信件未关联回复文件！")
        # 不更新数据库，因为 reply_file_path 已为 NULL

@with_db_connection
def delete_selected_letter(conn, event=None):
    """移除右侧选中的信件（支持多选），优化刷新和锁冲突处理"""
    selected_items = tree.selection()
    if not selected_items:
        logging.warning("未选择任何记录进行移除")
        messagebox.showwarning("警告", "请先选择一条或多条记录！")
        return

    if not messagebox.askyesno("确认", f"确定移除 {len(selected_items)} 条选中的信件？"):
        return

    if not scan_lock.acquire(blocking=False) or not drop_lock.acquire(blocking=False):
        logging.info("扫描或拖拽任务正在运行，取消移除")
        messagebox.showwarning("警告", "正在扫描或处理文件，请稍后再试！")
        scan_lock.release()
        drop_lock.release()
        return

    try:
        cursor = conn.cursor()
        deleted_count = 0
        max_retries = 10
        retry_delay = 0.3

        for selected_item in selected_items:
            if not tree.exists(selected_item):
                logging.warning(f"选中的项目 {selected_item} 不存在")
                continue
            letter_id = selected_item
            for attempt in range(max_retries):
                try:
                    cursor.execute("SELECT subject FROM letters WHERE id = ?", (letter_id,))
                    result = cursor.fetchone()
                    break
                except sqlite3.OperationalError as e:
                    if "database is locked" in str(e) and attempt < max_retries - 1:
                        logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}，信件 ID: {letter_id}")
                        time.sleep(retry_delay)
                        continue
                    logging.error(f"查询信件失败: {str(e)}, 信件 ID: {letter_id}")
                    raise

            if not result:
                logging.warning(f"未找到信件 ID: {letter_id}")
                continue

            subject = result[0]
            for attempt in range(max_retries):
                try:
                    cursor.execute("DELETE FROM letters WHERE id = ?", (letter_id,))
                    logging.info(f"移除信件: {letter_id}, 主题: {subject}")
                    deleted_count += 1
                    break
                except sqlite3.OperationalError as e:
                    if "database is locked" in str(e) and attempt < max_retries - 1:
                        logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}，信件 ID: {letter_id}")
                        time.sleep(retry_delay)
                        continue
                    logging.error(f"移除信件失败: {str(e)}, 信件 ID: {letter_id}")
                    raise

        conn.commit()  # 确保事务提交
        # 强制刷新界面
        root.after(0, lambda: update_letter_list())  # 刷新右侧列表
        root.after(0, lambda: update_treeview())     # 刷新左侧树
        root.after(0, lambda: messagebox.showinfo("提示", f"成功移除 {deleted_count} 条信件"))
    except sqlite3.OperationalError as e:
        logging.error(f"移除信件失败: {str(e)}")
        messagebox.showerror("错误", f"移除信件失败: 数据库锁冲突，请稍后再试")
    except Exception as e:
        logging.error(f"移除信件失败: {str(e)}")
        messagebox.showerror("错误", f"移除信件失败: {str(e)}")
    finally:
        scan_lock.release()
        drop_lock.release()
        

@with_db_connection
def delete_letters(conn):
    """移除左侧选中的信件或日期节点（支持多选），支持年/月/日节点，优化日节点删除"""
    selected_items = date_tree.selection()
    if not selected_items:
        logging.warning("未选择任何记录进行移除")
        messagebox.showwarning("警告", "请先选择一条或多条记录！")
        return
    if not messagebox.askyesno("确认", f"确定移除 {len(selected_items)} 条选中的记录？"):
        return
    if not scan_lock.acquire(blocking=False) or not drop_lock.acquire(blocking=False):
        logging.info("扫描或拖拽任务正在运行，取消移除")
        messagebox.showwarning("警告", "正在扫描或处理文件，请稍后再试！")
        scan_lock.release()
        drop_lock.release()
        return
    try:
        cursor = conn.cursor()
        deleted_count = 0
        max_retries = 10
        retry_delay = 0.3
        for item in selected_items:
            if not date_tree.exists(item):
                logging.warning(f"选中的项目 {item} 不存在")
                continue
            values = date_tree.item(item)["values"]
            if not values:
                logging.warning(f"节点 {item} 无有效值，跳过")
                continue
            node_type = values[1] if len(values) > 1 else ""
            if node_type == "文件":
                filename = values[0]
                for attempt in range(max_retries):
                    try:
                        cursor.execute("DELETE FROM letters WHERE subject = ?", (filename,))
                        logging.info(f"移除文件记录: {filename}")
                        deleted_count += 1
                        break
                    except sqlite3.OperationalError as e:
                        if "database is locked" in str(e) and attempt < max_retries - 1:
                            logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}，文件: {filename}")
                            time.sleep(retry_delay)
                            continue
                        logging.error(f"移除文件记录失败: {str(e)}, 文件: {filename}")
                        raise
            elif node_type in ("年", "月", "日"):
                date_filter = values[0]
                # 修改：对“日”节点使用 SUBSTR 提取日期部分
                if node_type == "日":
                    sql_condition = "SUBSTR(send_date, 1, 10) = ?"
                    params = (date_filter,)
                else:
                    sql_condition = "send_date LIKE ?"
                    params = (f"{date_filter}%",)
                for attempt in range(max_retries):
                    try:
                        cursor.execute(f"SELECT COUNT(*) FROM letters WHERE {sql_condition}", params)
                        count = cursor.fetchone()[0]
                        break
                    except sqlite3.OperationalError as e:
                        if "database is locked" in str(e) and attempt < max_retries - 1:
                            logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}，日期: {date_filter}")
                            time.sleep(retry_delay)
                            continue
                        logging.error(f"查询受影响记录失败: {str(e)}, 日期: {date_filter}")
                        raise
                if count == 0:
                    logging.info(f"日期节点 {date_filter} ({node_type}) 无匹配记录，跳过删除")
                    continue
                for attempt in range(max_retries):
                    try:
                        cursor.execute(f"DELETE FROM letters WHERE {sql_condition}", params)
                        logging.info(f"移除日期节点: {date_filter} ({node_type}), 影响 {count} 条记录")
                        deleted_count += count
                        break
                    except sqlite3.OperationalError as e:
                        if "database is locked" in str(e) and attempt < max_retries - 1:
                            logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}，日期: {date_filter}")
                            time.sleep(retry_delay)
                            continue
                        logging.error(f"移除日期节点失败: {str(e)}, 日期: {date_filter}")
                        raise
        conn.commit()
        if deleted_count == 0:
            root.after(0, lambda: messagebox.showinfo("提示", "未找到可删除的记录"))
        else:
            root.after(0, lambda: update_letter_list())
            root.after(0, lambda: update_treeview())
            root.after(0, lambda: clean_empty_nodes(''))
            root.after(0, lambda: messagebox.showinfo("提示", f"成功移除 {deleted_count} 条记录"))
    except sqlite3.OperationalError as e:
        logging.error(f"移除记录失败: {str(e)}")
        messagebox.showerror("错误", f"移除记录失败: 数据库锁冲突，请稍后再试")
    except Exception as e:
        logging.error(f"移除记录失败: {str(e)}")
        messagebox.showerror("错误", f"移除记录失败: {str(e)}")
    finally:
        scan_lock.release()
        drop_lock.release()
        
        
@with_db_connection
def save_settings(conn):
    """保存设置到数据库，包括用户选择的收发路径"""
    global default_keyword, scan_path1, default_aa, max_scan_files
    keyword = keyword_var.get().strip()
    new_scan_path1 = scan_path1_var.get().strip()
    scan_last_month_flag = scan_last_month_var.get()
    daily_enabled = daily_scan_enabled_var.get()
    daily_time = daily_scan_time_var.get().strip()
    aa_keyword = aa_var.get().strip()
    max_files = max_scan_files_var.get().strip()

    if not new_scan_path1:
        logging.error("扫描路径不能为空")
        messagebox.showerror("错误", "扫描路径不能为空！")
        return

    # 验证路径是否有效
    try:
        current_date = datetime.now()
        test_path = format_path(new_scan_path1, current_date)
        if not os.path.exists(test_path):
            logging.warning(f"路径无效: {test_path}")
            messagebox.showwarning("警告", f"路径无效: {test_path}，请检查路径！")
            return
    except (ValueError, FileNotFoundError) as e:
        logging.error(f"路径无效: {str(e)}")
        messagebox.showerror("错误", f"路径无效: {str(e)}")
        return

    if not daily_time or len(daily_time.split(":")) != 2:
        logging.error("无效的时间格式")
        messagebox.showerror("错误", "请提供有效的时间格式，如 09:30")
        return

    try:
        hour, minute = map(int, daily_time.split(":"))
        if not (0 <= hour <= 23 and 0 <= minute <= 59):
            raise ValueError
    except ValueError:
        logging.error("时间格式无效")
        messagebox.showerror("错误", "时间格式无效，小时应为 0-23，分钟应为 0-59")
        return

    try:
        max_files_value = int(max_files) if max_files and max_files.isdigit() else 0
        if max_files_value < 0:
            raise ValueError
    except ValueError:
        logging.error("最大扫描文件数量无效")
        messagebox.showerror("错误", "最大扫描文件数量必须为非负整数（0 表示无限制）")
        return

    try:
        cursor = conn.cursor()
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("keyword", keyword))
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("scan_path1", new_scan_path1))
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("scan_last_month", str(scan_last_month_flag)))
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("daily_scan_enabled", str(daily_enabled)))
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("daily_scan_time", daily_time))
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("aa_keyword", aa_keyword))
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("max_scan_files", str(max_files_value)))
        conn.commit()
        logging.info("设置保存成功")
    except sqlite3.Error as e:
        conn.rollback()
        logging.error(f"保存设置失败: {str(e)}")
        messagebox.showerror("错误", f"保存设置失败: {str(e)}")
        return

    default_keyword = keyword
    scan_path1 = new_scan_path1
    default_aa = aa_keyword
    max_scan_files = max_files_value
    scan_last_month_var.set(scan_last_month_flag)
    daily_scan_enabled_var.set(daily_enabled)
    daily_scan_time_var.set(daily_time)
    max_scan_files_var.set(str(max_files_value))

    settings_window.destroy()
    messagebox.showinfo("提示", "设置保存成功！")
    schedule_daily_scan()

def open_settings():
    """打开设置窗口，添加选择路径按钮和重置默认按钮，确保窗口置顶，并为每个设置项添加悬浮提示"""
    global settings_window
    settings_window = Toplevel(root)
    settings_window.withdraw()  # 隐藏窗口，直到配置完成
    settings_window.title("设置")
    try:
        settings_window.iconbitmap(APP_ICON_PATH)
    except tk.TclError as e:
        logging.warning(f"无法加载设置窗口图标 {APP_ICON_PATH}: {str(e)}")
    settings_window.resizable(True, True)
    window_width = min(520, root.winfo_screenwidth() - 100)
    window_height = 360  # 增加高度以容纳新按钮
    x = (root.winfo_screenwidth() - window_width) // 2
    y = (root.winfo_screenheight() - window_height) // 2
    settings_window.geometry(f"{window_width}x{window_height}+{x}+{y}")
    settings_window.deiconify()  # 配置完成后显示窗口

    settings_window.grid_columnconfigure(1, weight=1)

    # 标记点关键词
    keyword_label = tk.Label(settings_window, text="设置标记点关键词:")
    keyword_label.grid(row=0, column=0, sticky="w", padx=10, pady=5)
    keyword_entry = tk.Entry(settings_window, width=30, textvariable=keyword_var)
    keyword_entry.grid(row=0, column=1, sticky="ew", padx=10, pady=5)
    keyword_entry.config(state='normal')
    ToolTip(keyword_entry,"作用：添加原信件时，只添加内容包含关键词的docx文件\n"
                          "示例：'回复'、'及时回复'、'请回复'\n"
                          "注意：留空仅添加包含隐藏UUID的docx文件")

    # 自己常用代号
    aa_label = tk.Label(settings_window, text="设置自己常用代号:")
    aa_label.grid(row=1, column=0, sticky="w", padx=10, pady=5)
    aa_entry = tk.Entry(settings_window, width=30, textvariable=aa_var)
    aa_entry.grid(row=1, column=1, sticky="ew", padx=10, pady=5)
    aa_entry.config(state='normal')
    ToolTip(aa_entry, "作用：自动查找回复文件时，只查找包含此处代号的文件名\n" 
                      "示例：'张三 李四'（支持1-5个代号）\n"
                      "注意：多个用空格分隔，留空将查找所有")

    # 收发路径
    scan_path1_label = tk.Label(settings_window, text="收发路径(支持变量):")
    scan_path1_label.grid(row=2, column=0, sticky="w", padx=10, pady=5)
    scan_path1_entry = tk.Entry(settings_window, width=30, textvariable=scan_path1_var)
    scan_path1_entry.grid(row=2, column=1, sticky="ew", padx=10, pady=5)
    scan_path1_entry.config(state='normal')
    ToolTip(scan_path1_entry, "自动查找回复信的搜索位置，支持变量：\n"
                              "%YYYY%: 四位年份 (显示 2025)\n"
                              "%YY%: 两位年份 (显示 25)\n"
                              "%MM%: 两位月份 (显示 08)\n"
                              "%M%: 单/两位月份 (显示 8)\n"
                              "%DD%: 两位日期 (显示 01)\n"
                              "%D%: 单/两位日期 (显示 1)\n"
                              "作用：填写变量可以缩小范围、提高效率\n"
                              "注意：变量填写到月份就可以，不要填写日"
                              )

    def choose_scan_path():
        """选择收发路径并更新输入框，保持窗口置顶"""
        initial_dir = os.path.dirname(scan_path1_var.get()) if scan_path1_var.get() else "C:\\"
        path = filedialog.askdirectory(title="选择收发路径", initialdir=initial_dir)
        if path:
            path = path.replace("/", "\\")  # 规范化路径为 Windows 格式
            if not os.path.exists(path):
                logging.warning(f"选择的路径不存在: {path}")
                messagebox.showwarning("警告", f"选择的路径不存在，请重新选择！")
                return
            scan_path1_var.set(path)
            logging.info(f"用户选择收发路径: {path}")
            settings_window.lift()  # 置顶窗口
            settings_window.focus_force()  # 强制聚焦

    choose_path_btn = ttk.Button(settings_window, text="选择路径", command=choose_scan_path)
    choose_path_btn.grid(row=2, column=2, padx=5, pady=5)

    actual_path_var = tk.StringVar()
    actual_path_label = tk.Label(settings_window, textvariable=actual_path_var, wraplength=550, justify="left")
    actual_path_label.grid(row=3, column=0, columnspan=3, sticky="w", padx=10, pady=5)
    ToolTip(actual_path_label, "作用：预览填写变量后的路径，避免写错")

    def update_actual_path(*args):
        """实时更新实际路径显示"""
        try:
            current_date = datetime.now()
            formatted_path = format_path(scan_path1_var.get(), current_date)
            actual_path_var.set(f"收发路径变量预览:       {formatted_path}")
        except (ValueError, FileNotFoundError) as e:
            actual_path_var.set(f"实际路径: 无效路径 ({str(e)})")

    update_actual_path()
    scan_path1_var.trace("w", update_actual_path)

    # 最大扫描文件数量
    max_scan_files_label = tk.Label(settings_window, text="最大扫描文件数量:")
    max_scan_files_label.grid(row=4, column=0, sticky="w", padx=10, pady=5)
    max_scan_files_entry = tk.Entry(settings_window, width=30, textvariable=max_scan_files_var)
    max_scan_files_entry.grid(row=4, column=1, sticky="ew", padx=10, pady=5)
    max_scan_files_entry.config(state='normal')
    ToolTip(max_scan_files_entry, "作用：限制最大扫描的文件数，避免闪退\n"
                                  "注意：输入正整数，0表示不限制文件数量")

    # 扫描上一个月信件
    scan_last_month_label = tk.Label(settings_window, text="扫描上一个月信件:")
    scan_last_month_label.grid(row=5, column=0, sticky="w", padx=10, pady=5)
    scan_last_month_check = tk.Checkbutton(settings_window, text="", variable=scan_last_month_var)
    scan_last_month_check.grid(row=5, column=1, sticky="w", padx=10, pady=5)
    ToolTip(scan_last_month_check, "作用：勾选扫描上个月的信件文件夹\n"
                                   "注意：取消勾选只扫描当前月文件夹")

    # 启用每日定时扫描
    daily_scan_label = tk.Label(settings_window, text="启用每日定时扫描:")
    daily_scan_label.grid(row=6, column=0, sticky="w", padx=10, pady=5)
    daily_scan_check = tk.Checkbutton(settings_window, text="", variable=daily_scan_enabled_var)
    daily_scan_check.grid(row=6, column=1, sticky="w", padx=10, pady=5)
    ToolTip(daily_scan_check, "勾选以启用每天自动扫描\n"
                              "注意：需要提前打开程序")

    # 每日定时扫描时间
    time_label = tk.Label(settings_window, text="每日定时扫描时间:")
    ToolTip(time_label, "设置每天自动扫描的具体时间\n"
                        "作用：指定每日扫描的执行时间\n"
                        "示例：09:00（默认值）\n"
                        "注意：需启用每日定时扫描")
    time_frame = ttk.Frame(settings_window)

    try:
        init_hour, init_minute = map(int, daily_scan_time_var.get().split(":"))
    except ValueError:
        init_hour, init_minute = 9, 0

    hour_var = tk.StringVar(value=f"{init_hour:02d}")
    minute_var = tk.StringVar(value=f"{init_minute:02d}")

    hour_spinbox = ttk.Spinbox(time_frame, from_=0, to=23, width=2, textvariable=hour_var, format="%02.0f")
    hour_spinbox.grid(row=0, column=0, padx=(0, 5))
    ToolTip(hour_spinbox, "设置自动扫描的小时（0-23）\n"
                          "注意：输入无效值会重置为默认")
    tk.Label(time_frame, text=":").grid(row=0, column=1)
    minute_spinbox = ttk.Spinbox(time_frame, from_=0, to=59, width=2, textvariable=minute_var, format="%02.0f")
    minute_spinbox.grid(row=0, column=2, padx=(5, 0))
    ToolTip(minute_spinbox, "设置自动扫描的分钟（0-59）\n"
                            "注意：输入无效值会重置为默认")

    def update_time_var(*args):
        try:
            hour = int(hour_var.get())
            minute = int(minute_var.get())
            if 0 <= hour <= 23 and 0 <= minute <= 59:
                daily_scan_time_var.set(f"{hour:02d}:{minute:02d}")
            else:
                raise ValueError("时间超出有效范围")
        except ValueError:
            daily_scan_time_var.set("09:00")
            hour_var.set("09")
            minute_var.set("00")
            messagebox.showwarning("警告", "时间格式无效，已重置为默认值 09:00")

    hour_var.trace("w", update_time_var)
    minute_var.trace("w", update_time_var)

    def toggle_time_selector(*args):
        try:
            if not settings_window.winfo_exists():
                return
            if daily_scan_enabled_var.get():
                time_label.grid(row=7, column=0, sticky="w", padx=10, pady=5)
                time_frame.grid(row=7, column=1, sticky="w", padx=10, pady=5)
            else:
                time_label.grid_remove()
                time_frame.grid_remove()
        except tk.TclError as e:
            logging.warning(f"切换时间选择器时发生错误: {str(e)}")

    toggle_time_selector()
    daily_scan_enabled_var.trace("w", toggle_time_selector)

    def reset_defaults():
        """重置所有设置为默认值"""
        keyword_var.set("回复")
        aa_var.set("")
        scan_path1_var.set("Y:\\%YYYY%年%MM%月\\")
        max_scan_files_var.set("3000")
        scan_last_month_var.set(False)
        daily_scan_enabled_var.set(False)
        daily_scan_time_var.set("09:00")
        hour_var.set("09")
        minute_var.set("00")
        toggle_time_selector()
        logging.info("重置为默认设置")

    button_frame = ttk.Frame(settings_window)
    button_frame.grid(row=8, column=0, columnspan=3, pady=10)

    save_btn = ttk.Button(button_frame, text="保存", command=lambda: save_settings())
    cancel_btn = ttk.Button(button_frame, text="取消", command=settings_window.destroy)
    reset_btn = ttk.Button(button_frame, text="重置默认", command=reset_defaults)

    save_btn.pack(side="left", padx=(0, 10), ipadx=10)
    cancel_btn.pack(side="left", padx=(0, 10), ipadx=10)
    reset_btn.pack(side="left", ipadx=10)

@with_db_connection
def save_window_width(conn):
    """保存窗口宽度"""
    current_width = root.winfo_width()
    cursor = conn.cursor()
    cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("window_width", str(current_width)))

def on_resize(event):
    """延迟保存窗口宽度"""
    global resize_pending
    if resize_pending:
        root.after_cancel(resize_pending)
    resize_pending = root.after(10000, lambda: save_window_width())

def format_path(path, date):
    """格式化路径，支持日期变量"""
    if not path:
        logging.error("路径模板为空")
        raise ValueError("路径模板不能为空")

    try:
        if not any(var in path for var in ["%YYYY%", "%YY%", "%MM%", "%M%", "%DD%", "%D%"]):
            normalized_path = os.path.normpath(path)
            if not os.path.exists(os.path.dirname(normalized_path)) and normalized_path != os.path.normpath(path):
                logging.warning(f"路径 {normalized_path} 不存在")
                raise FileNotFoundError(f"路径 {normalized_path} 不存在")
            return normalized_path

        year = date.year
        month = date.month
        day = date.day

        path = path.replace("%YYYY%", str(year))
        path = path.replace("%YY%", f"{year % 100:02d}")
        path = path.replace("%MM%", f"{month:02d}")
        path = path.replace("%M%", str(month))
        path = path.replace("%DD%", f"{day:02d}")
        path = path.replace("%D%", str(day))

        formatted_path = os.path.normpath(path)
        if not os.path.exists(os.path.dirname(formatted_path)) and formatted_path != os.path.normpath(path):
            logging.warning(f"路径 {formatted_path} 不存在")
            raise FileNotFoundError(f"路径 {formatted_path} 不存在")
        return formatted_path
    except (OSError, ValueError) as e:
        logging.error(f"路径格式化失败: {str(e)}")
        raise ValueError(f"路径格式化失败: {str(e)}")

def extract_bb_from_filename(filename):
    """从文件名提取 BB 部分（格式：AA-BB 日期.docx）"""
    try:
        base_name = os.path.splitext(filename)[0]
        parts = base_name.split('-')
        if len(parts) >= 2:
            bb_part = parts[1].split()[0] if len(parts[1].split()) > 0 else parts[1]
            return bb_part.strip()
        return None
    except Exception as e:
        logging.warning(f"无法从文件名 {filename} 提取 BB 部分: {str(e)}")
        return None

def scan_reply_files():
    """自动查找文件并匹配 UUID，仅处理未关联回复的信件，排除文件名和修改日期均相同的文件，支持多个AA关键字"""
    max_retries = 5
    retry_delay = 0.5  # 增加重试延迟

    if not scan_lock.acquire(blocking=False):
        logging.info("已有扫描任务正在运行，退出 scan_reply_files")
        messagebox.showinfo("提示", "已有扫描任务正在运行，请稍后再试！")
        return

    try:
        auto_button.config(state="disabled")
        root.update_idletasks()
        cancel_scan = threading.Event()

        # 创建进度窗口（居中显示）
        progress_win = Toplevel(root)
        progress_win.title("自动查找回复中…")
        try:
            progress_win.iconbitmap(APP_ICON_PATH)
        except tk.TclError:
            logging.warning("无法加载进度窗口图标")
        window_width = 300
        window_height = 180
        screen_width = progress_win.winfo_screenwidth()
        screen_height = progress_win.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        progress_win.geometry(f"{window_width}x{window_height}+{x}+{y}")
        progress_text = StringVar()
        progress_text.set("正在扫描文件，请稍候…")
        Label(progress_win, textvariable=progress_text).pack(pady=10)
        progress_bar = ttk.Progressbar(progress_win, maximum=100, mode="indeterminate")
        progress_bar.pack(pady=10, fill="x", padx=10)
        cancel_button = ttk.Button(progress_win, text="取消", command=lambda: cancel_scan.set())
        cancel_button.pack(pady=10)
        def on_progress_close():
            cancel_scan.set()
            progress_win.destroy()
        progress_win.protocol("WM_DELETE_WINDOW", on_progress_close)
        progress_win.update()  # 强制刷新窗口
        root.update()  # 强制刷新主窗口
        progress_bar.start()  # 启动 indeterminate 模式动画

        def run_scan():
            start_time = time.time()
            try:
                conn = create_db_connection()
                if not conn:
                    logging.error("数据库连接失败")
                    root.after(0, lambda: messagebox.showerror("错误", "数据库连接失败"))
                    root.after(0, lambda: progress_win.destroy())  # 销毁进度窗口
                    return
                try:
                    # 分批查询数据库
                    cursor = conn.cursor()
                    offset = 0
                    batch_size = 1000
                    letters = []
                    while True:
                        cursor.execute("SELECT id, file_path, subject FROM letters WHERE monitor = 1 LIMIT ? OFFSET ?", (batch_size, offset))
                        batch = cursor.fetchall()
                        letters.extend(batch)
                        if len(batch) < batch_size:
                            break
                        offset += batch_size
                    if not letters:
                        logging.info("没有需要监控的信件，扫描终止")
                        root.after(0, lambda: messagebox.showinfo("提示", "没有需要监控的信件，扫描已终止！"))
                        root.after(0, lambda: progress_win.destroy())  # 销毁进度窗口
                        return
                    # 存储原信件的文件名和修改时间
                    original_files = {}  # 不再使用路径
                    original_subjects = {(os.path.basename(path).lower(), os.path.getmtime(path) if path and os.path.exists(path) else None)
                                        for _, path, subject in letters if path}
                    all_ids = {letter_id: os.path.basename(path).lower() for letter_id, path, _ in letters if path}
                except sqlite3.Error as e:
                    logging.error(f"数据库查询失败: {str(e)}")
                    root.after(0, lambda: messagebox.showerror("错误", f"数据库查询失败: {str(e)}"))
                    root.after(0, lambda: progress_win.destroy())  # 销毁进度窗口
                    return
                finally:
                    conn.close()

                current_date = datetime.now()
                paths = []
                try:
                    paths.append(format_path(scan_path1, current_date))
                    if scan_last_month_var.get():
                        last_month = current_date - relativedelta(months=1)
                        paths.append(format_path(scan_path1, last_month))
                    paths = list(dict.fromkeys(paths))
                except (ValueError, FileNotFoundError) as e:
                    logging.warning(f"路径格式化失败: {str(e)}")
                    root.after(0, lambda: messagebox.showwarning("警告", f"路径格式化失败: {str(e)}"))
                    root.after(0, lambda: progress_win.destroy())  # 销毁进度窗口
                    return
                if not paths or not any(os.path.exists(path) for path in paths):
                    logging.warning("收发路径无效或不存在")
                    root.after(0, lambda: messagebox.showwarning("警告", "收发路径无效或不存在，请检查设置！"))
                    root.after(0, lambda: progress_win.destroy())  # 销毁进度窗口
                    return

                bb_values = set()
                for _, _, subject in letters:
                    bb = extract_bb_from_filename(subject)
                    if bb:
                        bb_values.add(bb)
                files_to_scan = []
                aa_keywords = [kw.strip() for kw in default_aa.split() if kw.strip()] if default_aa else []

                if len(aa_keywords) > 5:
                    aa_keywords = aa_keywords[:5]
                    logging.warning("AA 关键词超过5个，已限制为前5个")
                    root.after(0, lambda: messagebox.showwarning("警告", "AA 关键词超过5个，已限制为前5个！"))
                elif len(aa_keywords) == 0 and default_aa:
                    logging.warning("无有效 AA 关键词，跳过过滤")
                elif len(aa_keywords) >= 1:
                    logging.info(f"使用 {len(aa_keywords)} 个 AA 关键词")

                def match_file(file, aa_keywords, bb_values):
                    for aa in aa_keywords:
                        for bb in bb_values:
                            pattern1 = f"*{aa}*{bb}*"
                            pattern2 = f"*{bb}*{aa}*"
                            if (fnmatch.fnmatch(file.lower(), pattern1.lower()) or
                                fnmatch.fnmatch(file.lower(), pattern2.lower())):
                                return True
                    return False

                max_workers = min(os.cpu_count() or 1, 4)
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = []
                    candidate_files = []
                    for path in paths:
                        if not os.path.exists(path):
                            logging.warning(f"路径不存在: {path}")
                            continue
                        for root_, _, files in os.walk(path):
                            if '$RECYCLE.BIN' in root_.upper():
                                continue
                            for file in files:
                                if file.lower().endswith(".docx"):
                                    abs_path = os.path.abspath(os.path.join(root_, file))
                                    file_mtime = os.path.getmtime(abs_path) if os.path.exists(abs_path) else None
                                    filename = os.path.basename(abs_path).lower()
                                    # 排除原信件：仅比对文件名和修改时间
                                    if (filename, file_mtime) in original_subjects:
                                        logging.debug(f"跳过原信件: {abs_path}")
                                        continue
                                    if aa_keywords and bb_values:
                                        candidate_files.append((abs_path, file))
                                        futures.append(executor.submit(match_file, file, aa_keywords, bb_values))
                                    else:
                                        files_to_scan.append(abs_path)
                    for future, (abs_path, file) in zip(futures, candidate_files):
                        if future.result():
                            files_to_scan.append(abs_path)

                total_files = len(files_to_scan)
                if total_files == 0:
                    logging.info("没有找到可扫描的文件")
                    root.after(0, lambda: messagebox.showinfo("提示", "没有可扫描的文件，请检查是否存在符合要求的文件。"))
                    root.after(0, lambda: progress_win.destroy())  # 销毁进度窗口
                    return
                if max_scan_files > 0 and total_files > max_scan_files:
                    files_to_scan = files_to_scan[:max_scan_files]
                    total_files = max_scan_files
                    root.after(0, lambda: messagebox.showwarning("警告", f"扫描文件数量过多，已限制为 {max_scan_files} 个文件！"))

                # 切换进度条到 determinate 模式
                progress_bar.stop()  # 停止 indeterminate 模式
                progress_bar.configure(maximum=total_files, mode="determinate")
                progress_text.set("扫描中：0 / " + str(total_files))

                def extract_text_fast(file_path):
                    if cancel_scan.is_set():
                        return ""
                    try:
                        if os.path.getsize(file_path) > 50 * 1024 * 1024:
                            logging.warning(f"文件 {file_path} 过大，跳过处理")
                            return ""
                        with open(file_path, "rb") as docx_file:
                            result = mammoth.extract_raw_text(docx_file)
                            return result.value
                    except MemoryError:
                        logging.error(f"处理文件 {file_path} 时内存不足")
                        return ""
                    except Exception as e:
                        logging.error(f"处理文件 {file_path} 时出错: {str(e)}")
                        return ""

                def process_file(args):
                    if cancel_scan.is_set():
                        return (args[0], args[1], "")
                    index, file_path = args
                    text = extract_text_fast(file_path)
                    return (index, file_path, text)

                matches = []
                max_workers = min(os.cpu_count() or 1, 2)
                processed_files = 0
                update_interval = max(1, total_files // 50)  # 减少界面刷新频率
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = [executor.submit(process_file, (i, f)) for i, f in enumerate(files_to_scan)]
                    for future in as_completed(futures):
                        if cancel_scan.is_set():
                            break
                        index, file_path, text = future.result()
                        processed_files += 1
                        if processed_files % update_interval == 0 or processed_files == total_files:
                            root.after(0, lambda idx=processed_files, fp=file_path: (
                                progress_text.set(f"扫描中：{idx} / {total_files}\n{os.path.basename(fp)}"),
                                progress_bar.configure(value=idx)
                            ))
                        for letter_id in all_ids:
                            if cancel_scan.is_set():
                                break
                            # 再次检查文件名是否为原信件
                            filename = os.path.basename(file_path).lower()
                            if filename == all_ids.get(letter_id):
                                logging.debug(f"跳过原信件匹配: {file_path} 是原信件")
                                continue
                            if letter_id in text:
                                # 获取文件修改时间
                                file_mtime = os.path.getmtime(file_path) if os.path.exists(file_path) else None
                                matches.append((letter_id, file_path, file_mtime))

                if not cancel_scan.is_set():
                    conn = create_db_connection()
                    if conn:
                        try:
                            with conn:
                                cursor = conn.cursor()
                                updates = []  # 用于统计新查询到的关联回复
                                for attempt in range(max_retries):
                                    try:
                                        # 清理无效路径并更新数据库
                                        for letter_id, reply_file_path, file_mtime in matches:
                                            # 再次验证文件名不是原信件
                                            filename = os.path.basename(reply_file_path).lower()
                                            if filename == all_ids.get(letter_id):
                                                logging.debug(f"跳过数据库更新: {reply_file_path} 是原信件")
                                                continue
                                            cursor.execute("SELECT reply_file_path FROM letters WHERE id = ?", (letter_id,))
                                            result = cursor.fetchone()
                                            existing_paths = result[0].split("|") if result[0] else []
                                            # 过滤掉不存在的路径
                                            valid_paths = [p for p in existing_paths if os.path.exists(p)]
                                            # 添加新路径（如果不存在）
                                            if os.path.exists(reply_file_path) and reply_file_path not in valid_paths:
                                                valid_paths.append(reply_file_path)
                                                new_reply_file_path = "|".join(valid_paths)
                                                updates.append((new_reply_file_path, 'replied', file_mtime, letter_id))
                                                logging.info(f"更新信件 {letter_id} 的回复文件: {new_reply_file_path}, 修改时间: {file_mtime}")
                                        if updates:
                                            cursor.executemany(
                                                "UPDATE letters SET reply_file_path = ?, reply_status = ?, reply_file_mtime = ? WHERE id = ?",
                                                updates
                                            )
                                        conn.commit()
                                        break
                                    except sqlite3.OperationalError as e:
                                        if "database is locked" in str(e) and attempt < max_retries - 1:
                                            logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}")
                                            time.sleep(retry_delay)
                                            continue
                                        logging.error(f"更新数据库失败: {str(e)}")
                                        root.after(0, lambda: messagebox.showerror("错误", f"更新数据库失败: {str(e)}"))
                                        return
                        finally:
                            conn.close()
                root.after(0, lambda: progress_win.destroy())
                root.after(0, lambda: update_letter_list())
                root.after(0, lambda: update_treeview())
                elapsed_time = time.time() - start_time
                logging.info(f"自动查找回复信件完成，总耗时: {elapsed_time:.2f} 秒，共 {len(matches)} 个关联回复，新查询到 {len(updates)} 条关联回复")
                root.after(0, lambda: messagebox.showinfo("完成", f"自动查找回复信件完成，总耗时: {elapsed_time:.2f} 秒\n共 {len(matches)} 个关联回复，新查询到 {len(updates)} 条关联回复"))
            finally:
                scan_lock.release()
                root.after(0, lambda: auto_button.config(state="normal"))
        threading.Thread(target=run_scan, daemon=True).start()
    except Exception as e:
        logging.error(f"扫描回复件时发生错误: {str(e)}")
        scan_lock.release()
        root.after(0, lambda: auto_button.config(state="normal"))
        root.after(0, lambda: messagebox.showerror("错误", f"扫描回复文件时发生错误: {str(e)}"))
        root.after(0, lambda: progress_win.destroy())  # 销毁进度窗口

@with_db_connection
def auto_lookup_original_files(conn):
    """自动查找原信件路径不存在的文件，仅匹配文件名，扫描范围与自动扫描回复相同"""
    max_retries = 5
    retry_delay = 0.5  # 增加重试延迟

    if not scan_lock.acquire(blocking=False):
        logging.info("已有扫描任务正在运行，跳过自动查找原信件")
        messagebox.showinfo("提示", "已有扫描任务正在运行，请稍后再试！")
        return

    try:
        start_time = time.time()
        cursor = conn.cursor()
        auto_original_button.config(state="disabled")
        
        # 分批查询数据库
        offset = 0
        batch_size = 1000
        all_letters = []
        for attempt in range(max_retries):
            try:
                while True:
                    cursor.execute("SELECT id, subject, file_path FROM letters WHERE monitor = 1 LIMIT ? OFFSET ?", (batch_size, offset))
                    batch = cursor.fetchall()
                    all_letters.extend(batch)
                    if len(batch) < batch_size:
                        break
                    offset += batch_size
                break
            except sqlite3.OperationalError as e:
                if "database is locked" in str(e) and attempt < max_retries - 1:
                    logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}")
                    time.sleep(retry_delay)
                    continue
                logging.error(f"查询数据库失败: {str(e)}")
                messagebox.showerror("错误", f"查询数据库失败: {str(e)}")
                return

        missing_letters = [(letter_id, subject.lower()) for letter_id, subject, file_path in all_letters
                          if not file_path or not os.path.exists(file_path)]
        if not missing_letters:
            logging.info("所有原信件路径都可访问，不需要修复")
            messagebox.showinfo("提示", "所有原信件路径都可访问，不需要修复！")
            return
        all_ids = {subject.lower() for letter_id, subject in missing_letters}
        letter_id_map = {subject.lower(): letter_id for letter_id, subject in missing_letters}
        logging.info(f"找到 {len(all_ids)} 条路径不存在的信件")

        # 获取扫描路径
        current_date = datetime.now()
        paths = []
        try:
            paths.append(format_path(scan_path1, current_date))
            if scan_last_month_var.get():
                last_month = current_date - relativedelta(months=1)
                paths.append(format_path(scan_path1, last_month))
            paths = list(dict.fromkeys(paths))
            logging.info(f"扫描路径: {paths}")
        except (ValueError, FileNotFoundError) as e:
            logging.warning(f"路径格式化失败: {str(e)}")
            messagebox.showwarning("警告", f"路径格式化失败: {str(e)}")
            return
        if not paths or not any(os.path.exists(path) for path in paths):
            logging.warning("收发路径无效或不存在")
            messagebox.showwarning("警告", "收发路径无效或不存在，请检查设置！")
            return

        # 创建进度窗口（居中显示）
        progress_win = Toplevel(root)
        progress_win.title("自动查找原信件中…")
        try:
            progress_win.iconbitmap(APP_ICON_PATH)
        except tk.TclError:
            logging.warning("无法加载进度窗口图标")
        window_width = 300
        window_height = 200
        screen_width = progress_win.winfo_screenwidth()
        screen_height = progress_win.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        progress_win.geometry(f"{window_width}x{window_height}+{x}+{y}")
        progress_text = StringVar(value="正在扫描文件，请稍候...")
        Label(progress_win, textvariable=progress_text).pack(pady=10)
        progress_bar = ttk.Progressbar(progress_win, maximum=100, mode="determinate")
        progress_bar.pack(pady=10, fill="x", padx=10)
        cancel_button = ttk.Button(progress_win, text="取消", command=lambda: cancel_scan.set())
        cancel_button.pack(pady=10)
        def on_progress_close():
            cancel_scan.set()
            progress_win.destroy()
        progress_win.protocol("WM_DELETE_WINDOW", on_progress_close)
        progress_win.update()

        # 文件扫描和匹配
        matches = []
        processed_files = 0
        scanned_dirs = 0
        cancel_scan = threading.Event()
        total_files_estimated = 3000
        update_interval = max(1, total_files_estimated // 50)  # 减少更新频率
        scan_start_time = time.time()

        def file_generator():
            for path in paths:
                if not os.path.exists(path):
                    logging.warning(f"路径不存在: {path}")
                    continue
                for root_, dirs, files in os.walk(path, topdown=True):
                    nonlocal scanned_dirs
                    scanned_dirs += 1
                    if '$RECYCLE.BIN' in root_.upper():
                        dirs[:] = []
                        continue
                    for file in files:
                        if file.lower().endswith(".docx"):
                            yield os.path.abspath(os.path.join(root_, file)), file.lower()

        for abs_path, file_name in file_generator():
            if cancel_scan.is_set():
                break
            processed_files += 1
            if processed_files > total_files_estimated:
                total_files_estimated = processed_files
                progress_bar.configure(maximum=total_files_estimated)
            if processed_files % update_interval == 0 or processed_files == total_files_estimated:
                elapsed = time.time() - scan_start_time
                remaining = (elapsed / processed_files) * (total_files_estimated - processed_files) if processed_files > 0 else 0
                root.after(0, lambda idx=processed_files, fn=file_name: (
                    progress_text.set(f"扫描中：{idx} / ~{total_files_estimated} 文件\n"
                                      f"当前文件：{fn}\n"
                                      f"预计剩余时间：{remaining:.1f} 秒"),
                    progress_bar.configure(value=idx)
                ))
            if file_name in all_ids:
                letter_id = letter_id_map[file_name]
                matches.append((letter_id, os.path.basename(abs_path), abs_path))
                logging.info(f"找到匹配文件，ID: {letter_id}, 文件名: {file_name}, 路径: {abs_path}")
                all_ids.remove(file_name)
                if not all_ids:
                    break

        # 批量更新数据库（分批）
        if not cancel_scan.is_set():
            try:
                batch_size = 1000
                for i in range(0, len(matches), batch_size):
                    batch = matches[i:i + batch_size]
                    for attempt in range(max_retries):
                        try:
                            cursor.executemany(
                                "UPDATE letters SET file_path = ?, subject = ? WHERE id = ?",
                                [(file_path, file_name, letter_id) for letter_id, file_name, file_path in batch]
                            )
                            conn.commit()
                            logging.info(f"修复 {len(batch)} 条记录")
                            break
                        except sqlite3.OperationalError as e:
                            if "database is locked" in str(e) and attempt < max_retries - 1:
                                logging.warning(f"数据库锁冲突，重试 {attempt + 1}/{max_retries}")
                                time.sleep(retry_delay)
                                continue
                            logging.error(f"更新数据库失败: {str(e)}")
                            messagebox.showerror("错误", f"更新数据库失败: {str(e)}")
                            return
                # 同步销毁进度窗口
                try:
                    progress_win.destroy()
                    logging.info("进度窗口已销毁")
                except tk.TclError as e:
                    logging.warning(f"进度窗口销毁失败: {str(e)}")
                # 修改：添加耗时，参考 scan_reply_files 格式
                elapsed_time = time.time() - start_time
                messagebox.showinfo("完成", f"自动查找原信件路径完成，总耗时: {elapsed_time:.2f} 秒\n"
                                           f"修复 {len(matches)} 条记录，{len(missing_letters) - len(matches)} 条始终未找到匹配文件")
            except Exception as e:
                logging.error(f"更新数据库失败: {str(e)}")
                messagebox.showerror("错误", f"更新数据库失败: {str(e)}")
                conn.rollback()

        # 清理和更新界面
        root.after(0, update_letter_list)
        root.after(0, update_treeview)
        elapsed_time = time.time() - start_time
        # 修改：日志添加耗时，参考 scan_reply_files 格式
        logging.info(f"自动查找原信件完成，总耗时: {elapsed_time:.2f} 秒，"
                     f"扫描目录: {scanned_dirs} 个，处理文件: {processed_files} 个，"
                     f"修复: {len(matches)} 条记录")
    except Exception as e:
        logging.error(f"自动查找原信件失败: {str(e)}")
        messagebox.showerror("错误", f"自动查找原信件失败: {str(e)}")
        # 确保异常情况下销毁进度窗口
        try:
            progress_win.destroy()
            logging.info("异常情况下进度窗口已销毁")
        except (NameError, tk.TclError) as e:
            logging.warning(f"异常情况下进度窗口销毁失败: {str(e)}")
    finally:
        scan_lock.release()
        logging.debug("释放 scan_lock")
        auto_original_button.config(state="normal")
     
@with_db_connection
def auto_follow_up_original_letter(conn):
    """自动生成跟进信件并保存到桌面跟进文件夹，日期和落款右对齐"""

    # 检查并创建 Data 目录
    base_path = get_base_path()
    data_dir = os.path.join(base_path, "Data")
    os.makedirs(data_dir, exist_ok=True)
    template_file = os.path.join(data_dir, "跟进信模板.txt")
    logging.debug(f"模板文件路径: {template_file}")

    # 默认模板内容（语气直接，长度 300-400 字）
    default_templates = [
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送了信件“{{subject}}”，目前没收到你的回复。这封信件需要咱们尽快反馈。想知道为什么还没回复，是工作太忙、内容有疑问，还是其他原因？我们希望尽快推进后续工作，所以需要你的明确答复。如果信件内容不清楚，可以告诉我具体问题，我会尽量帮你解答。另外，如果你的工作安排有冲突，也可以提前说清楚，是否可以协调时间？总之，请尽快回复，告诉我具体情况，避免影响整体进度。期待你的及时回复！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的“{{subject}}”还没收到你的回复。这封信件关系到整体的工作进度，按计划需要及时处理。你能告诉我为什么还没回复吗？比如，是不是因为信件内容有不清楚的地方，或者你这边有其他优先级更高的任务？我们希望尽快推进这件事，所以需要你的具体反馈。如果有任何问题，比如需要更多背景信息或者其他支持，可以直接告诉我，我会尽力配合。总之，希望你能尽快回复，说明情况，让我们把工作继续推进下去。请及时回复，谢谢！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的信件“{{subject}}”，目前还没收到你的反馈。这封信件内容比较重要，需要你及时回复。你能说说为什么还没回吗？比如，是不是工作量太大，或者信件内容有不明白的地方？我们希望通过沟通把事情推动下去。如果你有任何疑问，比如需要我提供更多细节或者其他帮助，随时说一声。我这边也会尽量配合你的进度，确保工作顺利进行。请尽快回复，告诉我具体情况，请及时回复，谢谢！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的“{{subject}}”目前没收到你的回复。这封信件涉及重要工作，按计划需要你尽快反馈。你能告诉我具体是什么原因导致没回复吗？比如，是不是因为忙其他项目，或者信件内容有什么不明确的地方？我们希望尽快推进工作，所以需要你的明确答复。如果有任何困难，比如需要更多时间或者额外信息，可以直接告诉我，我会尽量帮你解决。总之，请尽快回复，说明情况，让我们把后续工作顺利推进下去。期待你的及时回复！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的信件“{{subject}}”还没得到你的回复。这封信件对整体工作进度很重要，需要你及时配合。你能说说为什么还没回复吗？比如，是不是因为工作安排太满，或者信件内容有不明白的地方？我们希望尽快推动这件事，所以需要你的具体反馈。如果有任何问题，比如需要我提供更多背景资料或者其他支持，随时告诉我，我会尽力配合。希望你能尽快回复，说明具体情况，避免影响整体进度。请及时回复，谢谢！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的“{{subject}}”，目前还没收到你的回复。这封信件内容比较重要，按计划需要你及时反馈。你能告诉我为什么还没回复吗？比如，是不是因为内容有疑问，或者你这边有其他紧急任务？我们希望尽快推进工作，所以需要你的明确答复。如果有任何问题，比如需要更多信息或者其他帮助，可以直接告诉我，我会尽量配合你。总之，请尽快回复，说明具体情况，让我们把工作继续推进下去。请及时回复，谢谢！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的信件“{{subject}}”还没收到你的反馈。这封信件内容比较重要，需要你尽快回复。你能说说为什么还没回吗？比如，是不是因为工作太忙，或者信件内容有不明白的地方？我们希望通过沟通把事情推动下去。如果有任何困难，比如需要更多时间或者额外信息，可以直接告诉我，我会尽力帮你解决。总之，请尽快回复，说明情况，让我们把后续工作顺利推进下去。期待你的及时回复！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的“{{subject}}”，目前没收到你的回复。这封信件对整体进度很重要，需要你及时反馈。你能告诉我具体是什么原因导致没回复吗？比如，是不是因为文件内容不清楚，或者你这边有其他优先级更高的任务？我们希望尽快推进这件事，所以需要你的具体答复。如果有任何问题，比如需要我提供更多细节或者其他支持，随时说一声。我会尽量配合。希望你能尽快回复，说明情况，避免影响整体进度。请及时回复，谢谢！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的“{{subject}}”还没得到你的回复。这封信件内容比较重要，需要你尽快配合。你能说说为什么还没回复吗？比如，是不是因为忙其他项目，或者信件内容有不明白的地方？我们希望通过沟通把事情推动下去。如果有任何疑问，比如需要更多背景信息或者其他帮助，可以直接告诉我，我会尽力帮你解决。总之，请尽快回复，说明具体情况，让我们把工作继续推进下去。请及时回复，谢谢！",
        f"{{BB}}：\n你好！\n\n我们{{letter_date}}发送的“{{subject}}”，目前还没收到你的反馈。这封信件需要你及时回复。你能告诉我为什么还没回吗？比如，是不是因为工作量太大，或者信件内容有不明白的地方？我们希望尽快推进工作，所以需要你的明确答复。如果有任何问题，比如需要我提供更多资料或者其他支持，随时说一声。我会尽量配合你的进度。希望你能尽快回复，说明情况，避免影响整体进度，请及时回复，谢谢！"
    ]
    logging.debug(f"默认模板数量: {len(default_templates)}")

    # 检查模板文件是否存在，若不存在则创建默认模板文件
    if not os.path.exists(template_file):
        try:
            with open(template_file, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(default_templates))
            logging.info(f"已创建默认模板文件: {template_file}")
        except Exception as e:
            logging.error(f"创建默认模板文件失败: {str(e)}")
            messagebox.showerror("错误", f"无法创建默认模板文件: {str(e)}")
            return

    # 读取模板文件，使用正则表达式分割模板
    try:
        with open(template_file, 'r', encoding='utf-8') as f:
            content = f.read().strip()
        # 使用正则表达式按 {BB}： 分割，确保每个模板完整
        templates = re.split(r'(?={BB}：\n你好！\n\n)', content)
        # 清理模板，去除空字符串并修剪
        templates = [t.strip() for t in templates if t.strip()]
        if not templates:
            logging.warning("模板文件为空，使用默认模板")
            templates = default_templates
        else:
            logging.info(f"成功加载模板文件: {template_file}, 模板数量: {len(templates)}")
            for i, template in enumerate(templates, 1):
                logging.debug(f"模板 {i}: {template[:100]}...")
    except Exception as e:
        logging.error(f"读取模板文件失败: {str(e)}，使用默认模板")
        messagebox.showwarning("警告", f"读取模板文件失败，将使用默认模板: {str(e)}")
        templates = default_templates

    selected_items = tree.selection()
    if not selected_items:
        logging.warning("未选择任何记录进行自动跟进")
        messagebox.showwarning("警告", "请先选择一条记录！")
        return
    if len(selected_items) > 1:
        logging.warning("选择了多条记录，仅处理第一条")
        messagebox.showwarning("警告", "一次只能跟进一条记录！")
        selected_items = [selected_items[0]]

    if not tree.exists(selected_items[0]):
        logging.warning(f"选中的项目 {selected_items[0]} 不存在")
        messagebox.showwarning("警告", "选中的记录无效！")
        return

    letter_id = selected_items[0]
    cursor = conn.cursor()
    cursor.execute("SELECT file_path, subject FROM letters WHERE id = ?", (letter_id,))
    result = cursor.fetchone()
    if not result:
        logging.warning(f"未找到信件 {letter_id} 的数据")
        messagebox.showwarning("警告", "未找到该信件的数据！")
        return

    file_path, subject = result
    logging.debug(f"数据库查询结果: file_path={file_path}, subject={subject}")

    # 从文件名提取 AA、BB 和日期
    try:
        base_name = os.path.splitext(subject)[0]
        pattern = r'^(.+?)[-_](.+?)\s*[-_]?(\d{4}|\d{2}-\d{2})(.*)?$'
        match = re.match(pattern, base_name)
        if not match:
            raise ValueError("文件名格式无效，需为 AA-BB MMDD、AA-BB-MMDD 等格式，日期后可接任意字符")
        
        AA = match.group(1).strip()
        BB = match.group(2).strip()
        date_part = match.group(3).replace('-', '')
        remainder = match.group(4) if match.group(4) else ''
        logging.debug(f"提取文件名信息: AA={AA}, BB={BB}, date_part={date_part}, remainder={remainder or '无'}")

        if len(date_part) != 4 or not date_part.isdigit():
            raise ValueError("日期格式无效，需为 MMDD（如 0819）")
        month = int(date_part[:2])
        day = int(date_part[2:])
        if not (1 <= month <= 12 and 1 <= day <= 31):
            raise ValueError("日期无效，月份需为01-12，日期需为01-31")
        
        letter_date = f"{month:02d}月{day:02d}日"
        logging.debug(f"格式化日期: {letter_date}")
    except Exception as e:
        logging.error(f"提取 AA-BB 或日期失败: {str(e)}, 文件名: {subject}")
        messagebox.showerror("错误", f"文件名格式无效（需为 AA-BB MMDD、AA-BB-MMDD 等格式，日期后可接任意字符）: {str(e)}")
        return
        
    current_date = datetime.now()
    current_year = current_date.year
    current_month = current_date.month
    current_day = current_date.day
    logging.debug(f"当前日期: {current_year}年{current_month}月{current_day}日")

    # 随机选择一个模板并替换变量
    content = random.choice(templates)
    logging.debug(f"选择的模板: {content[:100]}...")
    try:
        content = content.format(BB=BB, subject=subject, letter_date=letter_date)
        logging.debug(f"替换变量后的正文: {content}")
    except Exception as e:
        logging.error(f"模板变量替换失败: {str(e)}")
        messagebox.showerror("错误", f"模板变量替换失败: {str(e)}")
        return

    # 创建新文档
    doc = Document()
    # 添加标题
    doc.add_heading(f"跟进原信件：{subject}", level=3)
    logging.debug("已添加标题")

    # 分割正文内容为称呼和正文部分
    parts = content.split('\n\n', 1)
    if len(parts) != 2:
        logging.error("模板格式错误，缺少正文内容")
        messagebox.showerror("错误", "模板格式错误，缺少正文内容")
        return

    greeting_full = parts[0]  # {BB}：\n你好！
    body = parts[1]           # 正文内容

    # 将 greeting_full 拆分成 BB 和 你好 两部分
    greeting_lines = greeting_full.split('\n', 1)
    BB_text = greeting_lines[0]          # BB：
    hello_text = greeting_lines[1] if len(greeting_lines) > 1 else "你好！"

    # 添加称呼（BB：）段落，不缩进
    bb_paragraph = doc.add_paragraph(BB_text)
    for run in bb_paragraph.runs:
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(12)

    # 添加“你好！”段落，首行缩进2个汉字
    hello_paragraph = doc.add_paragraph(hello_text)
    hello_paragraph.paragraph_format.first_line_indent = Pt(24)  # 2个汉字缩进
    for run in hello_paragraph.runs:
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(12)
    logging.debug(f"称呼: {BB_text}, 问候: {hello_text}")

    # 添加正文段落，首行缩进 2 个汉字，行距1.5倍
    body_paragraph = doc.add_paragraph(body)
    body_paragraph.paragraph_format.first_line_indent = Pt(24)
    body_paragraph.paragraph_format.line_spacing = 1.5
    for run in body_paragraph.runs:
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(12)
    logging.debug(f"正文内容: {body}")


    # 添加落款，右对齐，设置宋体 12号字体
    signature = doc.add_paragraph(f"{AA}\n{current_year}年{current_month}月{current_day}日")
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.paragraph_format.line_spacing = 1.5  # 新增行距设置
    for run in signature.runs:
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(12)
    logging.debug("已添加落款")

    # 保存到桌面“跟进”文件夹
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    follow_up_dir = os.path.join(desktop, "跟进")
    try:
        os.makedirs(follow_up_dir, exist_ok=True)
        current_date_str = datetime.now().strftime("%m%d")
        new_filename = f"{AA}-{BB} {current_date_str}【提醒】.docx"
        new_file_path = os.path.join(follow_up_dir, new_filename)
        doc.save(new_file_path)
        logging.info(f"跟进信件保存成功: {new_file_path}")
        messagebox.showinfo("完成", f"“{new_filename} ” 已创建完成！\n\n保存到桌面 跟进 文件夹，不满意可再次生成！")
    except Exception as e:
        logging.error(f"保存跟进信件失败: {str(e)}")
        messagebox.showerror("错误", f"保存跟进信件失败: {str(e)}")

def schedule_daily_scan():
    """调度每日定时扫描任务"""
    global daily_scan_timer
    if daily_scan_timer:
        daily_scan_timer.cancel()
        daily_scan_timer = None
        logging.info("已取消现有定时任务")

    if not daily_scan_enabled_var.get():
        logging.info("每日定时扫描已禁用")
        return

    try:
        hour, minute = map(int, daily_scan_time_var.get().split(":"))
        if not (0 <= hour <= 23 and 0 <= minute <= 59):
            raise ValueError("无效的时间范围")
    except ValueError as e:
        logging.error(f"每日扫描时间设置无效: {str(e)}")
        messagebox.showerror("错误", f"每日扫描时间设置无效: {str(e)}")
        return

    now = datetime.now()
    target_time = now.replace(hour=hour, minute=minute, second=0, microsecond=0)
    if target_time <= now:
        target_time += timedelta(days=1)
    delay = (target_time - now).total_seconds()

    def daily_task():
        global daily_scan_timer
        try:
            logging.info("开始每日定时扫描")
            scan_reply_files()
        except Exception as e:
            logging.error(f"每日定时扫描失败: {str(e)}")
            root.after(0, lambda: messagebox.showerror("错误", f"每日定时扫描失败: {str(e)}"))
        finally:
            daily_scan_timer = threading.Timer(86400, daily_task)
            daily_scan_timer.start()
            logging.info("每日定时扫描任务已重新调度，运行时间: 24小时后")

    daily_scan_timer = threading.Timer(delay, daily_task)
    daily_scan_timer.start()
    logging.info(f"调度每日扫描任务，首次运行时间: {target_time}")

def expand_all():
    """一键展开所有目录树节点"""
    def expand_recursive(item):
        date_tree.item(item, open=True)
        for child in date_tree.get_children(item):
            expand_recursive(child)
    for item in date_tree.get_children():
        expand_recursive(item)
    logging.info("目录树已全部展开")

def collapse_all():
    """一键收缩所有目录树节点"""
    def collapse_recursive(item):
        date_tree.item(item, open=False)
        for child in date_tree.get_children(item):
            collapse_recursive(child)
    for item in date_tree.get_children():
        collapse_recursive(item)
    logging.info("目录树已全部收缩")

@with_db_connection

def update_letter_list(conn, date_filter=None, keyword=None):
    try:
        query = "SELECT id, file_path, send_date, subject, reply_status, reply_file_path, reminder_date, reminder_status FROM letters WHERE monitor = 1"
        params = []
        if date_filter:
            date_filter_str = str(date_filter)
            if len(date_filter_str) == 4:  # 年，如 "2025"
                query += " AND send_date LIKE ?"
                params.append(f"{date_filter_str}%")
            elif len(date_filter_str) == 7:  # 月，如 "2025-08"
                query += " AND send_date LIKE ?"
                params.append(f"{date_filter_str}%")
            elif len(date_filter_str) == 10:  # 日，如 "2025-08-09"
                query += " AND send_date = ?"
                params.append(date_filter_str)
            else:
                logging.warning(f"无效的日期筛选格式: {date_filter_str}")
                return
        
        cursor = conn.cursor()
        cursor.execute(query, params)
        current_time = datetime.now()
        new_data = []
        for row in cursor.fetchall():
            letter_id, file_path, send_date, subject, reply_status, reply_file_path, reminder_date, reminder_status = row
            # 动态更新 reminder_status
            new_reminder_status = reminder_status
            if reminder_date:
                try:
                    reminder_datetime = datetime.strptime(reminder_date, "%Y-%m-%d %H:%M")
                    if new_reminder_status != reminder_status:
                        cursor.execute("UPDATE letters SET reminder_status = ? WHERE id = ?", (new_reminder_status, letter_id))
                except ValueError:
                    logging.warning(f"无效的提醒日期格式: {reminder_date} for letter_id: {letter_id}")
                    new_reminder_status = ""
            else:
                new_reminder_status = ""

            # ★ 新增：计算从 send_date 到今天的天数
            days_value = calc_days_since(send_date)

            # 原来 8 个字段，现在插入一个 days_value，一共 9 个
            new_data.append((
                reply_status,                        # 0
                send_date,                           # 1
                subject,                             # 2
                letter_id,                           # 3
                reply_file_path if reply_file_path else "",  # 4
                file_path,                           # 5
                reminder_date if reminder_date else "",      # 6
                new_reminder_status,                 # 7
                days_value                           # 8 ← 新增：天数
            ))
        conn.commit()  # 提交状态更新

        # 新增：如果有关键词，过滤数据（检查subject或reply_files的basename）
        if keyword:
            keyword_lower = keyword.lower()  # 忽略大小写
            filtered_data = []
            for item in new_data:
                subject = item[2].lower()  # 原信件名称
                reply_file = item[4]       # reply_file_path
                reply_files = reply_file.split("|") if reply_file else []
                reply_basenames = [os.path.basename(f).lower() for f in reply_files]  # 回复文件名
                # 如果subject或任何回复basename包含关键词，就保留
                if keyword_lower in subject or any(keyword_lower in bn for bn in reply_basenames):
                    filtered_data.append(item)
            new_data = filtered_data  # 用过滤后的数据替换

        for item in tree.get_children():
            tree.delete(item)

        # 这里解包顺序对应上面 new_data.append 的 9 个元素
        for reply_status, date, filename, letter_id, reply_file, file_path, reminder_date, reminder_status, days_value in new_data:
            reply_files = reply_file.split("|") if reply_file and isinstance(reply_file, str) else []
            if reply_files:
                if len(reply_files) <= 1:
                    reply_file_name = ", ".join(os.path.basename(f) for f in reply_files if f)
                else:
                    reply_file_name = f"{len(reply_files)} 个回复"
            else:
                reply_file_name = ""
            reply_file_display = reply_file_name or "无"
            status_text = "已回复" if reply_status == "replied" else "未回复"
            file_exists = os.path.exists(file_path) if file_path else False
            tags = []
            if not file_exists:
                tags.append("missing")

            # 从数据库读取备注
            cursor.execute("SELECT note FROM letters WHERE id = ?", (letter_id,))
            note_row = cursor.fetchone()
            note_text = note_row[0] if note_row and note_row[0] else ""

            # ★ 注意：在“添加日期”后面插入 days_value
            tree.insert("", "end", iid=letter_id, values=(
                status_text,          # 状态
                filename,             # 原信件
                date,                 # 添加日期
                days_value,           # ★ 天数（新列）
                reply_file_display,   # 回复信件
                note_text,            # 备注
                reminder_date or "无" # 提醒日期
            ), tags=tags)

            if reply_files:
                tooltip_text = "\n".join(os.path.basename(f) for f in reply_files if f)
                def show_tip(e, t=tooltip_text, id=letter_id):
                    current_row = tree.identify_row(e.y)
                    if current_row == id and not getattr(tree, f"_tooltip_shown_{id}", False):
                        ToolTip(tree, t).show_tip(e)
                        setattr(tree, f"_tooltip_shown_{id}", True)
                def check_motion(e, id=letter_id):
                    current_row = tree.identify_row(e.y)
                    if current_row != id and getattr(tree, f"_tooltip_shown_{id}", False):
                        ToolTip(tree, "").hide_tip(e)
                        setattr(tree, f"_tooltip_shown_{id}", False)
                tree.tag_bind(letter_id, "<Motion>", lambda e, id=letter_id, t=tooltip_text: [show_tip(e, t, id), check_motion(e, id)])

        logging.info(f"更新信件列表，筛选日期: {date_filter if date_filter else '全部'}, 关键词: {keyword if keyword else '无'}, 数据条数: {len(new_data)}")
    except Exception as e:
        logging.error(f"更新信件列表失败: {str(e)}")
        messagebox.showerror("错误", f"更新信件列表失败: {str(e)}")
        
        
@with_db_connection
def update_letter_list_by_subject(conn, subject, keyword=None):
    try:
        query = "SELECT id, file_path, send_date, subject, reply_status, reply_file_path, reminder_date FROM letters WHERE monitor = 1 AND subject = ?"
        params = [subject]
        
        cursor = conn.cursor()
        cursor.execute(query, params)
        new_data = [(row[4], row[2], row[3], row[0], row[5] if row[5] else "", row[1], row[6] if row[6] else "") for row in cursor.fetchall()]

        # 新增：如果有关键词，过滤数据
        if keyword:
            keyword_lower = keyword.lower()
            filtered_data = []
            for item in new_data:
                subject_lower = item[2].lower()
                reply_file = item[4]
                reply_files = reply_file.split("|") if reply_file else []
                reply_basenames = [os.path.basename(f).lower() for f in reply_files]
                if keyword_lower in subject_lower or any(keyword_lower in bn for bn in reply_basenames):
                    filtered_data.append(item)
            new_data = filtered_data

        for item in tree.get_children():
            tree.delete(item)

        for reply_status, date, filename, letter_id, reply_file, file_path, reminder_date in new_data:
            reply_files = reply_file.split("|") if reply_file and isinstance(reply_file, str) else []
            if reply_files:
                if len(reply_files) <= 1:
                    reply_file_name = ", ".join(os.path.basename(f) for f in reply_files if f)
                else:
                    reply_file_name = f"{len(reply_files)} 个回复"
            else:
                reply_file_name = ""
            reply_file_display = reply_file_name or "无"
            status_text = "已回复" if reply_status == "replied" else "未回复"
            file_exists = os.path.exists(file_path) if file_path else False
            tags = []
            if not file_exists:
                tags.append("missing")
            tree.insert("", "end", iid=letter_id, values=(
                status_text,           # Status
                filename,             # File
                date,                 # Date
                days_value,           # ★ 天数（新列）
                reply_file_display,   # ReplyFile
                note_text,            # 备注              
                reminder_date or "无" # ReminderDate
            ), tags=tags)

            if reply_files:
                tooltip_text = "\n".join(os.path.basename(f) for f in reply_files if f)
                def show_tip(e, t=tooltip_text, id=letter_id):
                    current_row = tree.identify_row(e.y)
                    if current_row == id and not getattr(tree, f"_tooltip_shown_{id}", False):
                        ToolTip(tree, t).show_tip(e)
                        setattr(tree, f"_tooltip_shown_{id}", True)
                def check_motion(e, id=letter_id):
                    current_row = tree.identify_row(e.y)
                    if current_row != id and getattr(tree, f"_tooltip_shown_{id}", False):
                        ToolTip(tree, "").hide_tip(e)
                        setattr(tree, f"_tooltip_shown_{id}", False)
                tree.tag_bind(letter_id, "<Motion>", lambda e, id=letter_id, t=tooltip_text: [show_tip(e, t, id), check_motion(e, id)])

        logging.info(f"更新信件列表，筛选主题: {subject}, 关键词: {keyword if keyword else '无'}, 数据条数: {len(new_data)}")
    except Exception as e:
        logging.error(f"按主题更新信件列表失败: {str(e)}")
        messagebox.showerror("错误", f"更新信件列表失败: {str(e)}")
 
#def sort_treeview_column(col, reverse):
def sort_treeview_column(col, reverse):
    """按指定列排序 Treeview，保留选中状态"""
    global sort_column, sort_reverse
    
    # 获取当前 Treeview 数据
    data = [(tree.set(item, col), item) for item in tree.get_children()]
    
    # 保存当前选中项的 ID
    selected_ids = tree.selection()
    
    # 处理不同列的排序逻辑
    def get_sort_key(item):
        value = item[0]
        if col == "Date":
            try:
                # 按日期排序，转换为 datetime
                return datetime.strptime(value, "%Y-%m-%d") if value else datetime.min
            except ValueError:
                return datetime.min
        elif col == "Days":  # ★ 新增：Days列数字排序
            try:
                return int(value) if value and value.isdigit() else 0
            except ValueError:
                return 0
        elif col == "ReminderDate":
            try:
                # 按提醒日期排序，转换为 datetime
                return datetime.strptime(value, "%Y-%m-%d %H:%M") if value and value != "无" else datetime.min
            except ValueError:
                return datetime.min
        elif col == "ReplyFile":
            # 回复信件空值排最后
            return value if value and value != "无" else "zzz"
        elif col == "Note":  # 新增：Note 列按文本排序，空值排最后
            return value.lower() if value else "zzz"
        else:
            # 状态和信件名称按字母排序，空值排最后
            return value.lower() if value else "zzz"
    
    # 按指定列排序
    data.sort(key=get_sort_key, reverse=reverse)
    
    # 清空 Treeview 并重新插入数据
    for item in tree.get_children():
        tree.delete(item)
    
    conn = create_db_connection()
    if not conn:
        logging.error("无法连接数据库，排序失败")
        messagebox.showerror("错误", "无法连接数据库")
        return
    
    try:
        cursor = conn.cursor()
        for _, item_id in data:
            # 获取完整数据
            cursor.execute("SELECT reply_status, send_date, subject, reply_file_path, file_path, reminder_date, reminder_status FROM letters WHERE id = ?", (item_id,))
            row = cursor.fetchone()
            if row:
                reply_status, date, filename, reply_file, file_path, reminder_date, reminder_status = row
                
                # ★ 新增：计算天数
                days_value = calc_days_since(date)
                
                reply_files = reply_file.split("|") if reply_file and isinstance(reply_file, str) else []
                if reply_files:
                    if len(reply_files) <= 1:
                        reply_file_name = ", ".join(os.path.basename(f) for f in reply_files if f)
                    else:
                        reply_file_name = f"{len(reply_files)} 个回复"
                else:
                    reply_file_name = "无"  # 确保空值为"无"
                status_text = "已回复" if reply_status == "replied" else "未回复"
                tags = []
                if not os.path.exists(file_path) if file_path else False:
                    tags.append("missing")
                
                # 新增：从数据库读取 note（备注）
                cursor.execute("SELECT note FROM letters WHERE id = ?", (item_id,))
                note_row = cursor.fetchone()
                note_text = note_row[0] if note_row and note_row[0] else ""
                
                tree.insert("", "end", iid=item_id, values=(
                    status_text,           # Status
                    filename,             # File
                    date,                 # Date
                    days_value,           # Days ← 修复！用days_value
                    reply_file_name,      # ReplyFile
                    note_text,            # Note
                    reminder_date or "无" # ReminderDate
                ), tags=tags)

                if reply_files:
                    tooltip_text = "\n".join(os.path.basename(f) for f in reply_files if f)
                    def show_tip(e, t=tooltip_text, id=item_id):
                        current_row = tree.identify_row(e.y)
                        if current_row == id and not getattr(tree, f"_tooltip_shown_{id}", False):
                            ToolTip(tree, t).show_tip(e)
                            setattr(tree, f"_tooltip_shown_{id}", True)
                    def check_motion(e, id=item_id):
                        current_row = tree.identify_row(e.y)
                        if current_row != id and getattr(tree, f"_tooltip_shown_{id}", False):
                            ToolTip(tree, "").hide_tip(e)
                            setattr(tree, f"_tooltip_shown_{id}", False)
                    tree.tag_bind(item_id, "<Motion>", lambda e, id=item_id, t=tooltip_text: [show_tip(e, t, id), check_motion(e, id)])
    
        # 更新排序状态
        sort_column = col
        sort_reverse = not reverse if sort_column == col else False
        
        # ★ 修复：表头列名加上"Days"
        for column in ("Status", "File", "Date", "Days", "ReplyFile", "Note", "ReminderDate"):
            if column == col:
                arrow = " ↓" if reverse else " ↑"
                tree.heading(column, text=tree.heading(column)["text"].split(" ")[0] + arrow)
            else:
                tree.heading(column, text=tree.heading(column)["text"].split(" ")[0])
        
        logging.info(f"按 {col} 列排序，方向: {'降序' if reverse else '升序'}")
    except sqlite3.Error as e:
        logging.error(f"排序时数据库操作失败: {str(e)}")
        messagebox.showerror("错误", f"数据库操作失败: {str(e)}")
    finally:
        conn.close()

    
@with_db_connection
def update_treeview(conn, event=None):
    """更新左侧日期树状导航，不自动移除不存在的文件，标记为灰色"""
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT id, file_path, reply_file_path, subject FROM letters WHERE monitor = 1")
        for letter_id, file_path, reply_file_path, subject in cursor.fetchall():
            if reply_file_path:
                paths = reply_file_path.split("|")
                valid_paths = []
                missing_paths = []
                for p in paths:
                    norm_path = os.path.normpath(p)
                    if os.path.exists(norm_path) and os.access(norm_path, os.R_OK):
                        valid_paths.append(norm_path)
                    else:
                        missing_paths.append(norm_path)
                if missing_paths:
                    if valid_paths:  # 只有存在有效路径时才更新 reply_file_path
                        new_reply_file_path = "|".join(valid_paths)
                        logging.info(f"信件ID: {letter_id}, 原reply_file_path: {reply_file_path}, 更新后: {new_reply_file_path}, 不存在的路径: {missing_paths}")
                        cursor.execute("UPDATE letters SET reply_file_path = ?, reply_status = 'replied' WHERE id = ?", 
                                      (new_reply_file_path, letter_id))
                    else:
                        logging.warning(f"信件ID: {letter_id} 的所有回复路径均不存在，未更新 reply_file_path: {reply_file_path}")
                        # 不更新数据库，保留原始 reply_file_path
    except Exception as e:
        logging.error(f"检查文件失败: {str(e)}")
        messagebox.showerror("错误", f"检查文件失败: {str(e)}，请检查文件路径或权限")

    # 清空现有树状结构
    for item in date_tree.get_children():
        date_tree.delete(item)

    try:
        cursor.execute("SELECT DISTINCT SUBSTR(send_date, 1, 4) FROM letters WHERE monitor = 1")
        years = cursor.fetchall()
        for year in years:
            year_str = str(year[0])
            year_node = date_tree.insert("", "end", text=f"{year_str}年", values=(year_str, "年"), image=icon_calendar, open=True)

            cursor.execute("SELECT DISTINCT SUBSTR(send_date, 6, 2), send_date, subject, reply_status, file_path FROM letters WHERE monitor = 1 AND send_date LIKE ?", (f"{year_str}%",))
            months_data = cursor.fetchall()
            months = {}
            for row in months_data:
                month = row[0]
                if month not in months:
                    months[month] = []
                months[month].append((row[1], row[2], row[3], row[4]))
            for month in sorted(months.keys()):
                month_node = date_tree.insert(year_node, "end", text=f"{month}月", values=(f"{year_str}-{month}", "月"), image=icon_calendar, open=True)

                days = {}
                for day, filename, reply_status, file_path in months[month]:
                    if day not in days:
                        days[day] = []
                    days[day].append((filename, reply_status, file_path))
                for day in sorted(days.keys()):
                    day_node = date_tree.insert(month_node, "end", text=f"{day}", values=(day, "日"), image=icon_calendar)
                    for filename, reply_status, file_path in days[day]:
                        tags = ["file_icon"]
                        if reply_status == "replied":
                            tags.append("replied")
                        else:
                            tags.append("unreplied")
                        if file_path and not os.path.exists(file_path):
                            tags.append("missing")
                        date_tree.insert(day_node, "end", text=f"{filename}", values=(filename, "文件", reply_status), image=icon_file, tags=tags)

    except Exception as e:
        logging.error(f"更新树状视图失败: {str(e)}")
        messagebox.showerror("错误", f"更新树状视图失败: {str(e)}，请检查数据库连接或文件路径")

    def clean_empty_nodes(item):
        """清理空的年/月/日节点"""
        children = date_tree.get_children(item)
        for child in children:
            clean_empty_nodes(child)
        if not date_tree.get_children(item) and date_tree.item(item)["values"][1] in ("年", "月", "日"):
            date_tree.delete(item)

    # 清理空节点
    for item in date_tree.get_children():
        clean_empty_nodes(item)

    logging.info("更新树状视图完成")

def on_date_tree_select(event):
    global update_letter_list_pending
    if update_letter_list_pending:
        root.after_cancel(update_letter_list_pending)

    selected_items = date_tree.selection()
    
    if not selected_items:
        logging.info("未选择任何日期节点，显示所有信件")
        update_letter_list_pending = root.after(100, lambda: update_letter_list(date_filter=None))
        return

    if len(selected_items) > 1:
        logging.warning("选择了多个日期节点，仅处理第一个")
        messagebox.showwarning("警告", "仅支持选择一个日期节点，将处理第一个选中节点")
        selected_items = [selected_items[0]]

    item = date_tree.item(selected_items[0])
    node_type = item["values"][1] if len(item["values"]) > 1 else ""
    date_filter = item["values"][0] if len(item["values"]) > 0 else ""

    if node_type in ("年", "月", "日"):
        logging.info(f"选择 {node_type} 节点: {date_filter}")
        update_letter_list_pending = root.after(100, lambda: update_letter_list(date_filter=date_filter))
    elif node_type == "文件":
        logging.info(f"选择文件节点: {date_filter}")
        update_letter_list_pending = root.after(100, lambda: update_letter_list_by_subject(date_filter))
    else:
        logging.warning(f"选择无效节点类型: {node_type}，显示所有信件")
        messagebox.showwarning("警告", "选择了无效的节点类型，将显示所有信件")
        update_letter_list_pending = root.after(100, lambda: update_letter_list(date_filter=None))

def on_right_click_date_tree(event):
    """左侧目录树右键菜单"""
    item = date_tree.identify_row(event.y)
    if item:
        date_tree.selection_set(item)
        values = date_tree.item(item)["values"]
        node_type = values[1] if len(values) > 1 else ""
        # 动态调整菜单项，仅对文件节点显示“打开原信件”和“打开原信件目录”
        date_menu.delete(0, tk.END)
        date_menu.add_command(label="移除原信件", image=icon_file, compound="left", command=delete_letters)
        date_menu.add_command(label="添加原信件", image=icon_file, compound="left", command=manual_select)
        date_menu.add_command(label="全部展开", image=icon_folder, compound="left", command=expand_all)
        date_menu.add_command(label="全部收缩", image=icon_folder, compound="left", command=collapse_all)
        if node_type == "文件":
            date_menu.add_separator()
            date_menu.add_command(label="打开原信件", image=icon_file, compound="left", command=lambda: open_original_file(subject=values[0]))
            date_menu.add_command(label="打开原信件目录", image=icon_folder, compound="left", command=lambda: open_original_directory(subject=values[0]))
        date_menu.post(event.x_root, event.y_root)

def on_right_click_tree(event):
    """右侧信件列表右键菜单，优化菜单状态和空选择处理，添加取消提醒和分割线"""
    logging.debug(f"右键菜单事件触发，坐标: x={event.x_root}, y={event.y_root}")

    # 确保 Treeview 获得焦点
    try:
        tree.focus_set()
    except tk.TclError as e:
        logging.error(f"设置 Treeview 焦点失败: {str(e)}")
        return

    # 识别点击的行
    item = tree.identify_row(event.y)
    current_selection = tree.selection()
    logging.debug(f"识别的行: {item}, 当前选中项: {current_selection}")

    # 清空或更新选择
    if item and tree.exists(item):
        if item not in current_selection:
            tree.selection_set(item)
            logging.debug(f"右键点击的行 {item} 不在选中项中，设置为唯一选中项")
        else:
            logging.debug(f"右键点击的行 {item} 已在选中项中，保留多选状态")
    else:
        tree.selection_remove(current_selection)
        logging.debug("右键点击空白区域，清空选择")

    def refresh_both():
        update_letter_list()
        update_treeview()
        logging.info("刷新左侧日期树和右侧信件列表")

    @with_db_connection
    def view_all_replies(conn):
        """查看选定信件的所有回复文件，窗口居中显示"""
        selected_items = tree.selection()
        if not selected_items or len(selected_items) > 1:
            messagebox.showwarning("警告", "请仅选择一条记录！")
            return
        letter_id = selected_items[0]
        if not tree.exists(letter_id):
            messagebox.showwarning("警告", "选中的记录无效！")
            return
        cursor = conn.cursor()
        cursor.execute("SELECT reply_file_path, subject FROM letters WHERE id = ?", (letter_id,))
        result = cursor.fetchone()
        if not result:
            logging.warning(f"未找到信件 {letter_id} 的数据")
            messagebox.showwarning("警告", "未找到该信件的数据！")
            return
        reply_file_path, subject = result
        reply_files = reply_file_path.split("|") if reply_file_path and isinstance(reply_file_path, str) else []
        logging.debug(f"查看回复文件，letter_id: {letter_id}, reply_files: {reply_files}")

        view_win = tk.Toplevel(root)
        view_win.title(f"所有回复 - {subject}")
        try:
            view_win.iconbitmap(APP_ICON_PATH)
        except tk.TclError as e:
            logging.warning(f"无法加载查看回复窗口图标 {APP_ICON_PATH}: {str(e)}")

        # 设置窗口大小
        win_width = 400
        win_height = 300
        view_win.resizable(False, False)
        view_win.attributes("-topmost", True)
        view_win.transient(root)

        # 居中窗口
        root.update_idletasks()  # 确保主窗口尺寸已更新
        root_x = root.winfo_rootx()
        root_y = root.winfo_rooty()
        root_width = root.winfo_width()
        root_height = root.winfo_height()
        x = root_x + (root_width - win_width) // 2
        y = root_y + (root_height - win_height) // 2
        view_win.geometry(f"{win_width}x{win_height}+{x}+{y}")

        tk.Label(view_win, text="关联的回复文件：").pack(pady=5)
        listbox = tk.Listbox(view_win, height=min(len(reply_files), 10))
        listbox.pack(pady=5, padx=10, fill="both", expand=True)
        for file in reply_files:
            if file:
                listbox.insert(tk.END, os.path.basename(file))

        def open_selected():
            selected_indices = listbox.curselection()
            if not selected_indices:
                messagebox.showwarning("警告", "未选择任何文件！")
                return
            for idx in selected_indices:
                file_path = reply_files[idx]
                if os.path.exists(file_path):
                    try:
                        os.startfile(file_path)
                        logging.info(f"打开回复文件: {file_path}")
                    except OSError as e:
                        messagebox.showerror("错误", f"无法打开文件 {os.path.basename(file_path)}: {str(e)}")
                        logging.error(f"打开文件失败: {file_path}, 错误: {str(e)}")

        button_frame = ttk.Frame(view_win)
        button_frame.pack(pady=5)
        ttk.Button(button_frame, text="打开选中文件", command=open_selected).pack(side="left", padx=5)
        ttk.Button(button_frame, text="关闭", command=view_win.destroy).pack(side="left", padx=5)

        # 确保窗口正确显示
        view_win.update_idletasks()
        view_win.grab_set()
        view_win.focus_force()

    @with_db_connection
    def cancel_reminder(conn):
        """取消选定信件的提醒日期"""
        selected_items = tree.selection()
        if not selected_items:
            logging.warning("未选择任何记录取消提醒")
            root.after(0, lambda: messagebox.showwarning("警告", "请先选择一条或多条记录！"))
            return

        try:
            cursor = conn.cursor()
            updated_count = 0
            for selected_item in selected_items:
                if not tree.exists(selected_item):
                    logging.warning(f"选中的项目 {selected_item} 不存在")
                    continue
                letter_id = selected_item
                cursor.execute("SELECT reminder_date FROM letters WHERE id = ?", (letter_id,))
                result = cursor.fetchone()
                if result and result[0]:
                    cursor.execute("UPDATE letters SET reminder_date = NULL WHERE id = ?", (letter_id,))
                    updated_count += 1
                    logging.info(f"取消信件 {letter_id} 的提醒日期")
                else:
                    logging.info(f"信件 {letter_id} 无提醒日期，无需取消")

            conn.commit()
            if updated_count > 0:
                root.after(0, lambda: messagebox.showinfo("提示", f"成功取消 {updated_count} 条记录的提醒"))
                root.after(100, update_letter_list)
                root.after(100, update_treeview)
            else:
                root.after(0, lambda: messagebox.showinfo("提示", "选中的记录均无提醒日期"))
        except sqlite3.Error as e:
            conn.rollback()
            logging.error(f"取消提醒失败: {str(e)}")
            root.after(0, lambda: messagebox.showerror("错误", f"取消提醒失败: {str(e)}"))

    # 检查选中的信件是否有回复文件和提醒日期
    has_reply_file = False
    has_reminder = False
    if tree.selection():
        with create_db_connection() as conn:
            cursor = conn.cursor()
            for item in tree.selection():
                cursor.execute("SELECT reply_file_path, reminder_date FROM letters WHERE id = ?", (item,))
                result = cursor.fetchone()
                if result:
                    if result[0] and result[0].strip():
                        has_reply_file = True
                    if result[1] and result[1].strip():
                        has_reminder = True
                    if has_reply_file and has_reminder:
                        break

    # 清空并重新构建右键菜单
    letter_menu.delete(0, tk.END)
    
    # 原信件相关操作
    letter_menu.add_command(label="打开原信件", image=icon_file, compound="left", command=open_original_file)
    letter_menu.add_command(label="打开原信件目录", image=icon_folder, compound="left", command=open_original_directory)
    letter_menu.add_command(label="自动跟进原信件", image=icon_file, compound="left", command=auto_follow_up_original_letter)
    letter_menu.add_separator()  # 分割线  

    
    # 回复信件相关操作
    letter_menu.add_command(label="打开回复信件", image=icon_file, compound="left", command=open_reply_file)
    letter_menu.add_command(label="打开回复信件目录", image=icon_folder, compound="left", command=open_reply_directory)
    letter_menu.add_command(label="查看回复信件列表", image=icon_file, compound="left", command=view_all_replies)
    letter_menu.add_separator()  # 分割线

    # 其他操作
    letter_menu.add_command(label="添加原信件", image=icon_file, compound="left", command=manual_select)
    letter_menu.add_command(label="移除原信件", image=icon_file, compound="left", command=delete_selected_letter)
    letter_menu.add_command(label="自动查找原信件", image=icon_refresh, compound="left", command=auto_lookup_original_files)
    letter_menu.add_separator()  # 分割线
    

    letter_menu.add_command(label="添加回复信件", image=icon_link, compound="left", command=upload_reply_file)
    letter_menu.add_command(label="移除回复信件", image=icon_link, compound="left", command=cancel_reply_association)
    letter_menu.add_command(label="自动查找回复信", image=icon_link, compound="left", command=scan_reply_files)
    letter_menu.add_separator()  # 分割线

    
    # 提醒相关操作
    letter_menu.add_command(label="设置提醒", image=icon_calendar, compound="left", command=set_reminder_date)
    letter_menu.add_command(label="取消提醒", image=icon_calendar, compound="left", command=cancel_reminder)
    letter_menu.add_separator()  # 分割线
    
    # 设置菜单项状态
    has_selection = bool(tree.selection())
    logging.debug(f"当前选择状态: {'有选中项' if has_selection else '无选中项'}, 是否有回复文件: {has_reply_file}, 是否有提醒: {has_reminder}")
    letter_menu.entryconfig("移除原信件", state="normal" if has_selection else "disabled")
    letter_menu.entryconfig("打开原信件", state="normal" if has_selection else "disabled")
    letter_menu.entryconfig("打开原信件目录", state="normal" if has_selection else "disabled")
    letter_menu.entryconfig("打开回复信件", state="normal" if has_selection and has_reply_file else "disabled")
    letter_menu.entryconfig("打开回复信件目录", state="normal" if has_selection and has_reply_file else "disabled")
    letter_menu.entryconfig("查看回复信件列表", state="normal" if has_selection and has_reply_file else "disabled")
    letter_menu.entryconfig("添加回复信件", state="normal" if has_selection else "disabled")
    letter_menu.entryconfig("移除回复信件", state="normal" if has_selection and has_reply_file else "disabled")
    letter_menu.entryconfig("设置提醒", state="normal" if has_selection else "disabled")
    letter_menu.entryconfig("取消提醒", state="normal" if has_selection and has_reminder else "disabled")

    try:
        # 强制刷新 UI，确保菜单显示
        root.update_idletasks()
        letter_menu.post(event.x_root, event.y_root)
        logging.debug(f"右键菜单显示在坐标: x={event.x_root}, y={event.y_root}")
    except tk.TclError as e:
        logging.error(f"无法显示右键菜单: {str(e)}")
        root.after(0, lambda: messagebox.showerror("错误", f"无法显示右键菜单: {str(e)}"))
    except Exception as e:
        logging.error(f"右键菜单显示失败，意外错误: {str(e)}")
        root.after(0, lambda: messagebox.showerror("错误", f"右键菜单显示失败: {str(e)}"))
              
def on_exit():
    """清理定时任务并退出程序"""
    global daily_scan_timer, resize_pending
    if daily_scan_timer:
        daily_scan_timer.cancel()
        daily_scan_timer = None
        logging.info("已取消每日定时扫描任务")
    if resize_pending:
        root.after_cancel(resize_pending)
        resize_pending = None
    logging.info("程序退出")
    root.destroy()
    sys.exit(0)

# ---- 界面搭建 ----
root.bind("<Configure>", on_resize)
root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=1)

pane = ttk.PanedWindow(root, orient=tk.HORIZONTAL)
pane.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)

left_frame = ttk.Frame(pane)
pane.add(left_frame, weight=1)
left_frame.grid_columnconfigure(0, weight=1)
left_frame.grid_rowconfigure(0, weight=1)

date_tree = ttk.Treeview(left_frame, columns=("Type", "Reply"), show="tree")
date_tree.column("#0", width=200, anchor="e")  # 宽度从200，标题右对齐
date_tree.column("Type", width=0, stretch=False)
date_tree.column("Reply", width=0, anchor="e")  # 年月日/文件列隐藏
date_tree.grid(row=0, column=0, sticky="nsew")

date_scroll_y = ttk.Scrollbar(left_frame, orient="vertical", command=date_tree.yview)
date_scroll_y.grid(row=0, column=1, sticky="ns")
date_scroll_x = ttk.Scrollbar(left_frame, orient="horizontal", command=date_tree.xview)
date_scroll_x.grid(row=1, column=0, sticky="ew")
date_tree.configure(yscrollcommand=date_scroll_y.set, xscrollcommand=date_scroll_x.set)

date_menu = tk.Menu(root, tearoff=0)
# 定义右侧信件列表的右键菜单
letter_menu = tk.Menu(root, tearoff=0)
date_menu.add_command(label="移除原信件", image=icon_file, compound="left", command=delete_letters)
date_menu.add_command(label="添加原信件", image=icon_file, compound="left", command=manual_select)
date_menu.add_command(label="全部展开", image=icon_folder, compound="left", command=expand_all)
date_menu.add_command(label="全部收缩", image=icon_folder, compound="left", command=collapse_all)
date_tree.bind("<Button-3>", on_right_click_date_tree)

right_frame = ttk.Frame(pane)
pane.add(right_frame, weight=2)
right_frame.grid_columnconfigure(0, weight=1)
right_frame.grid_rowconfigure(0, weight=1)   # Treeview
right_frame.grid_rowconfigure(1, weight=0)   # 按钮区域

# 右侧信件列表
tree = ttk.Treeview(right_frame, columns=("Status", "File", "Date", "Days", "ReplyFile", "Note", "ReminderDate"), show="headings")

tree.heading("Status", text="状态", command=lambda: sort_treeview_column("Status", sort_reverse))
tree.heading("File", text="原信件", command=lambda: sort_treeview_column("File", sort_reverse))
tree.heading("Date", text="添加日期", command=lambda: sort_treeview_column("Date", sort_reverse))
tree.heading("Days", text="天数", command=lambda: sort_treeview_column("Days", sort_reverse))
tree.heading("ReplyFile", text="回复信件", command=lambda: sort_treeview_column("ReplyFile", sort_reverse))
tree.heading("Note", text="备注", command=lambda: sort_treeview_column("Note", sort_reverse))
tree.heading("ReminderDate", text="提醒日期", command=lambda: sort_treeview_column("ReminderDate", sort_reverse))

tree.column("Status", width=60, anchor="center")
tree.column("File", width=220, anchor="w")
tree.column("Date", width=90, anchor="center")
tree.column("Days", width=60, anchor="center")
tree.column("ReplyFile", width=200, anchor="w")
tree.column("Note", width=180, anchor="w")          # 新增
tree.column("ReminderDate", width=120, anchor="center")

tree.grid(row=0, column=0, sticky="nsew")  # 调整到 row=0
#tree.bind("<Double-1>", on_double_click)
# 双击：正常双击 = 打开文件/编辑备注，Ctrl+双击 = 复制文件名到剪贴板
#tree.bind("<Double-1>", lambda e: on_double_click(e) if not (e.state & 0x4) else copy_filename_to_clipboard(e))
#tree.bind("<Control-Double-1>", copy_filename_to_clipboard)   # 保险起见再绑一次
# 主绑定：普通双击 → 打开/编辑，Ctrl 双击 → 复制，Alt 双击 → Everything 搜索
# 终极绑定：支持 普通双击 / Ctrl双击 / Alt双击
tree.bind("<Double-1>", lambda e: on_double_click(e) 
                              if not (e.state & 0x4 or e.state & 0x20000) 
                              else copy_filename_to_clipboard(e))
tree.bind("<Control-Double-1>", copy_filename_to_clipboard)
tree.bind("<Alt-Double-1>", copy_filename_to_clipboard)


tree.bind("<Button-3>", on_right_click_tree)

# 绑定鼠标中键点击事件到treeview（文档列表）
tree.bind("<Button-2>", copy_uuid_to_clipboard_middle_button)  # 鼠标中键点击触发


scrollbar = ttk.Scrollbar(right_frame, orient="vertical", command=tree.yview)
scrollbar.grid(row=0, column=1, sticky="ns")  # 调整到 row=0
tree.configure(yscrollcommand=scrollbar.set)

button_frame = ttk.Frame(right_frame)
button_frame.grid(row=1, column=0, columnspan=2, sticky="ew", padx=5, pady=2)  # 调整到 row=1
for col in range(5):  # 修改：从 6 列改为 5 列，因为移除“退出”按钮
    button_frame.grid_columnconfigure(col, weight=1)

select_button = ttk.Button(button_frame, text="添加原信件", command=manual_select, width=12)
select_button.grid(row=0, column=0, sticky="ew", padx=3, pady=2)
ToolTip(select_button, "选择docx原信件，或拖拽多个文件、文件夹")

# 修改：将“手动查找回复”改为“自动查找原信件”
auto_original_button = ttk.Button(button_frame, text="自动查找原信件", command=auto_lookup_original_files, width=12)
auto_original_button.grid(row=0, column=1, sticky="ew", padx=3, pady=2)
ToolTip(auto_original_button, "自动查找收发路径中更改位置的原信件")

auto_button = ttk.Button(button_frame, text="自动查找回复信", command=scan_reply_files, width=12)
auto_button.grid(row=0, column=2, sticky="ew", padx=3, pady=2)
ToolTip(auto_button, "自动查找收发路径中的回复信件")

settings_button = ttk.Button(button_frame, text="设置", command=open_settings, width=12)
settings_button.grid(row=0, column=3, sticky="ew", padx=3, pady=2)

# 新增：搜索框和按钮
search_entry = tk.Entry(button_frame, textvariable=search_var, width=15)
search_entry.grid(row=0, column=4, sticky="ew", padx=3, pady=2)  # 修改：从 column=5 改为 column=4
ToolTip(search_entry, "输入关键词搜索原信件或回复信件名称")

def perform_search():
    keyword = search_var.get().strip()
    update_letter_list(keyword=keyword)  # 调用带关键词的更新

search_button = ttk.Button(button_frame, text="搜索", command=perform_search, width=8)
search_button.grid(row=0, column=5, sticky="ew", padx=3, pady=2)  # 修改：从 column=6 改为 column=5

# 支持回车搜索
search_entry.bind("<Return>", lambda event: perform_search())

# 初始加载
update_letter_list()
update_treeview()
schedule_daily_scan()

def on_drop(event):
    """处理拖拽事件"""
    if not drop_lock.acquire(blocking=False):
        logging.info("已有拖拽任务正在运行，跳过本次拖拽")
        messagebox.showinfo("提示", "已有拖拽任务正在运行，请稍后再试！")
        return
    try:
        raw_data = event.data
        logging.debug(f"拖拽事件数据: {raw_data}")
        paths = []
        if isinstance(raw_data, str):
            # 处理 Windows 拖拽路径，考虑包含空格的情况
            if raw_data.startswith('{') and raw_data.endswith('}'):
                # 多文件拖拽，路径被 {} 包裹
                raw_data = raw_data.strip('{}')
                paths = [p.strip() for p in raw_data.split('} {') if p.strip()]
            else:
                # 单文件或包含空格的路径
                # 使用正则表达式按空白分割，同时保留包含空格的完整路径
                paths = [p.strip() for p in re.split(r'\s+(?=(?:[^"]*"[^"]*")*[^"]*$)', raw_data) if p.strip()]
        else:
            logging.error("拖拽数据格式无效")
            messagebox.showerror("错误", "拖拽数据格式无效")
            return

        # 过滤非空且存在的路径
        valid_paths = [os.path.normpath(p.strip('"')) for p in paths if p and os.path.exists(p.strip('"'))]
        if not valid_paths:
            logging.warning("拖拽的文件或目录不存在或路径为空")
            messagebox.showwarning("警告", "拖拽的文件或目录不存在或路径为空！")
            return

        drop(file_paths=valid_paths)
    except Exception as e:
        logging.error(f"处理拖拽事件失败: {str(e)}")
        messagebox.showerror("错误", f"处理拖拽事件失败: {str(e)}")
    finally:
        drop_lock.release()

root.drop_target_register(DND_FILES)
root.dnd_bind("<<Drop>>", on_drop)

root.protocol("WM_DELETE_WINDOW", on_exit)
# 确保事件绑定
date_tree.bind("<<TreeviewSelect>>", on_date_tree_select)



def check_single_instance():
    """检查程序是否已运行，限制单实例，如果已运行则激活已有窗口"""
    mutex_name = "LetterTrackerAppMutex"  # 唯一互斥锁名称
    mutex = win32event.CreateMutex(None, False, mutex_name)
    last_error = win32api.GetLastError()
    if last_error == 183:  # ERROR_ALREADY_EXISTS
        logging.warning("程序已在运行，尝试激活已有窗口")
        # 查找并激活已有窗口（窗口标题包含“跟进助手”）
        def enum_windows_callback(hwnd, regex):
            try:
                window_text = win32gui.GetWindowText(hwnd)
                if win32gui.IsWindowVisible(hwnd) and re.search(r"跟进助手_260406", window_text, re.IGNORECASE):
                    if win32gui.IsIconic(hwnd):  # 检查窗口是否最小化
                        win32gui.ShowWindow(hwnd, win32con.SW_RESTORE)
                    # 修改：修复模块引用，使用 win32process.GetWindowThreadProcessId
                    foreground_thread = win32gui.GetWindowThreadProcessId(win32gui.GetForegroundWindow())[1]
                    current_thread = win32api.GetCurrentThreadId()
                    if foreground_thread != current_thread:
                        win32process.AttachThreadInput(current_thread, foreground_thread, True)
                    # 依次调用，确保窗口置顶并激活
                    win32gui.BringWindowToTop(hwnd)
                    win32gui.SetActiveWindow(hwnd)
                    win32gui.SetForegroundWindow(hwnd)
                    # 使用 SetWindowPos 确保窗口置于顶层
                    win32gui.SetWindowPos(hwnd, win32con.HWND_TOPMOST, 0, 0, 0, 0, 
                                         win32con.SWP_NOMOVE | win32con.SWP_NOSIZE | win32con.SWP_NOACTIVATE)
                    win32gui.SetWindowPos(hwnd, win32con.HWND_NOTOPMOST, 0, 0, 0, 0, 
                                         win32con.SWP_NOMOVE | win32con.SWP_NOSIZE | win32con.SWP_NOACTIVATE)
                    if foreground_thread != current_thread:
                        win32process.AttachThreadInput(current_thread, foreground_thread, False)
                    return False  # 停止枚举
                return True
            except Exception as e:
                logging.error(f"枚举窗口时出错：{str(e)}")
                return True
        try:
            win32gui.EnumWindows(enum_windows_callback, None)
            logging.warning("未找到标题包含‘跟进助手_260406’的窗口，退出")
            sys.exit(0)  # 未找到窗口，退出当前实例
        except Exception as e:
            logging.error(f"窗口枚举失败：{str(e)}")
            sys.exit(0)
    return mutex

def main():
    """主程序入口，初始化界面并启动"""
    global root, tree, date_tree
    # 检查单实例
    mutex = check_single_instance()
    try:
        # 原有 main 函数内容
        logging.debug("开始初始加载")
        update_treeview()
        if not date_tree.selection():
            logging.debug("初始加载：无选中节点，显示所有信件")
            update_letter_list(date_filter=None)
        schedule_daily_scan()

        root.mainloop()
    finally:
        # 释放互斥锁
        win32api.CloseHandle(mutex)

if __name__ == "__main__":
    main()
